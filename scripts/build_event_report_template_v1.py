"""Build the compact Event-only DOCX report template.

This follows the supplied ``Incident_Report_27-8-69.odt``: a pale-blue Event
banner, one general-information section, and no Incident-only investigation,
containment, sign-off, or statutory appendix pages.

Run from the repository root::

    python scripts/build_event_report_template_v1.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR))

from scripts.build_report_template_v2 import (  # noqa: E402
    BODY_FONT,
    BORDER,
    LABEL_GRAY,
    LOGO_PATH,
    MUTED,
    SYMBOL_FONT,
    TEXT,
    add_runs,
    checkbox_segments,
    clear_table_borders,
    set_cell_margins,
    set_cell_shading,
    set_table_borders,
)
from apps.incidents.report_content import (  # noqa: E402
    EVENT_FOOTER_RIGHT,
    EVENT_SECTION1_ROWS,
    FOOTER_LEFT,
    SECTION_TITLES,
)


OUTPUT_PATH = (
    BASE_DIR / 'apps' / 'incidents' / 'report_templates'
    / 'event_report_template_v1.docx'
)
EVENT_BLUE = 'D9EAF7'
EVENT_SECTION_BLUE = 'EAF3FA'


def _first_paragraph(cell):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    return paragraph


def _set_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement('w:cantSplit'))
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=28, start=80, bottom=28, end=80)


def _style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.right_margin = Inches(0.7)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.7)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.2)

    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn('w:ascii'), BODY_FONT)
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), BODY_FONT)
    normal._element.rPr.rFonts.set(qn('w:cs'), BODY_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1


def _build_header(doc):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))
    else:
        add_runs(paragraph, [('NT', BODY_FONT, 16, TEXT, True)])


def _build_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(7.1))
    table.autofit = False
    clear_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(3.9)
    right.width = Inches(3.2)
    lp = _first_paragraph(left)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(lp, [(FOOTER_LEFT, BODY_FONT, 8.5, MUTED, False)])
    rp = _first_paragraph(right)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_runs(rp, [(EVENT_FOOTER_RIGHT, BODY_FONT, 8.5, MUTED, False)])
    empty = footer.paragraphs[0]
    empty._p.getparent().remove(empty._p)


def _add_title(doc):
    table = doc.add_table(rows=1, cols=1)
    _set_widths(table, [7.1])
    clear_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, EVENT_BLUE)
    set_cell_margins(cell, top=75, bottom=75)
    paragraph = _first_paragraph(cell)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(paragraph, [('Alert Event REPORT', BODY_FONT, 19, '1F2933', True)])
    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0)
    subtitle.paragraph_format.line_spacing = 1
    add_runs(
        subtitle,
        [('แบบฟอร์มแจ้งเหตุการณ์ผิดปกติ', BODY_FONT, 14, '1F2933', True)],
    )


def _add_section_band(doc):
    table = doc.add_table(rows=1, cols=1)
    _set_widths(table, [7.1])
    clear_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, EVENT_SECTION_BLUE)
    set_cell_margins(cell, top=34, bottom=34)
    paragraph = _first_paragraph(cell)
    paragraph.paragraph_format.keep_with_next = True
    add_runs(
        paragraph,
        [(f'1. {SECTION_TITLES["1"]}', BODY_FONT, 13, '1F2933', True)],
    )


def _event_rows():
    rows = []
    for kind, label, spec in EVENT_SECTION1_ROWS:
        if kind == 'kv':
            value = [(f'{{{{{spec}}}}}', BODY_FONT, 11.5, TEXT, False)]
        else:
            value = [
                (text, font, 10.5 if font == SYMBOL_FONT else 11.5, color, bold)
                for text, font, _size, color, bold in checkbox_segments(spec)
            ]
        rows.append((label, value))
    return rows


def _add_event_table(doc):
    rows = _event_rows()
    table = doc.add_table(rows=len(rows), cols=2)
    _set_widths(table, [2.65, 4.45])
    set_table_borders(table, color=BORDER, size='5')
    for index, (label, value_segments) in enumerate(rows):
        left, right = table.rows[index].cells
        set_cell_shading(left, LABEL_GRAY)
        add_runs(
            _first_paragraph(left),
            [(label, BODY_FONT, 11.5, TEXT, True)],
        )
        add_runs(_first_paragraph(right), value_segments)


def build(output_path=OUTPUT_PATH):
    doc = Document()
    _style_document(doc)
    _build_header(doc)
    _build_footer(doc)
    _add_title(doc)

    _add_section_band(doc)
    _add_event_table(doc)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)


if __name__ == '__main__':
    build()
