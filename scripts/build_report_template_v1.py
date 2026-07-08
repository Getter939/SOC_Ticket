from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = BASE_DIR / 'apps' / 'incidents' / 'report_templates' / 'report_template_v1.docx'

BLUE = '1F4D78'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
BORDER = 'B8C2CC'
TEXT = '1F2933'
MUTED = '5B6775'


def set_run_font(run, name='Calibri', size=11, color=TEXT, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_borders(table, color=BORDER, size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def set_table_indent(table, dxa=120):
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.find(qn('w:tblInd'))
    if ind is None:
        ind = OxmlElement('w:tblInd')
        tbl_pr.append(ind)
    ind.set(qn('w:w'), str(dxa))
    ind.set(qn('w:type'), 'dxa')


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ('Heading 1', 16, BLUE),
        ('Heading 2', 13, BLUE),
        ('Heading 3', 12, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run('SOC Ticket Report Template v1')
    set_run_font(run, size=9, color=MUTED)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run('INCIDENT REPORT')
    set_run_font(run, size=18, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run('แบบฟอร์มรายงานเหตุการณ์ผิดปกติ')
    set_run_font(run, size=14, color=BLUE, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run('Template {{template_version}} | Generated {{generated_at}}')
    set_run_font(run, size=9, color=MUTED)


def section_heading(doc, text):
    paragraph = doc.add_paragraph(style='Heading 1')
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_widths(table, [1.55, 1.7, 1.55, 1.7])
    set_table_borders(table)
    set_table_indent(table)
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            if col_idx in (0, 2):
                set_cell_shading(cell, LIGHT_GRAY)
                bold = True
            else:
                bold = False
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, bold=bold)
    doc.add_paragraph()
    return table


def add_single_table(doc, rows, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_widths(table, [1.75, 4.75])
    set_table_borders(table)
    set_table_indent(table)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, header_fill)
        for cell, text, bold in ((left, label, True), (right, value, False)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, bold=bold)
    doc.add_paragraph()
    return table


def add_signature_table(doc):
    table = doc.add_table(rows=4, cols=3)
    set_table_widths(table, [2.15, 2.15, 2.2])
    set_table_borders(table)
    set_table_indent(table)
    headers = ('ผู้รายงาน / Created by', 'ผู้ตรวจสอบ / Verified by', 'ผู้อนุมัติ / Approved by')
    values = (
        ('{{created_by}}', '{{verified_by}}', '{{approved_by}}'),
        ('{{created_at}}', '{{verified_at}}', '{{approved_at}}'),
        ('..........................................', '..........................................', '..........................................'),
    )
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, bold=True)
    for row_idx, row_values in enumerate(values, start=1):
        for col, text in enumerate(row_values):
            paragraph = table.rows[row_idx].cells[col].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5)


def add_appendix(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_heading(doc, 'ภาคผนวก: หมวดหมู่ของภัยคุกคามทางไซเบอร์')
    intro = doc.add_paragraph()
    intro.add_run(
        'อ้างอิงตามภาคผนวกท้ายประกาศคณะกรรมการการรักษาความมั่นคงปลอดภัยไซเบอร์แห่งชาติ '
        'เรื่อง ลักษณะภัยคุกคามทางไซเบอร์ มาตรการป้องกัน รับมือ ประเมิน ปราบปราม '
        'และระงับภัยคุกคามทางไซเบอร์แต่ละระดับ พ.ศ. 2564'
    )
    categories = [
        ('Training and Exercises', 'เหตุการณ์จำลอง และ การฝึกจู่โจม ของหน่วยงาน'),
        ('Unsuccessful Activity Attempt', 'การพยายามเข้าถึงระบบที่ไม่สำเร็จ'),
        ('Reconnaissance', 'การพยายามบุกรุกเพื่อสำรวจข้อมูลองค์กรเพื่อโจมตี'),
        ('Non-Compliance Activity', 'การดำเนินการที่ไม่เป็นไปตามมาตรฐานความปลอดภัยที่หน่วยงานกำหนด'),
        ('Malicious Logic', 'การบุกรุกโดยการใช้มัลแวร์'),
        ('User Level Intrusion', 'การบุกรุกในระดับผู้ใช้งาน'),
        ('Root Level Intrusion', 'การบุกรุกในระดับผู้ควบคุมระบบ'),
        ('Denial of Service', 'การบุกรุกที่ทำให้ไม่สามารถเข้าใช้บริการได้'),
        ('Investigating', 'เหตุการณ์ที่อยู่ระหว่างการวิเคราะห์สอบสวน'),
        ('Explained Anomaly', 'เหตุการณ์ผิดปกติที่ได้รับการวิเคราะห์แล้วไม่ใช่เหตุการณ์ที่เป็นภัยคุกคาม'),
    ]
    add_single_table(doc, categories, header_fill=LIGHT_GRAY)


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    section_heading(doc, '1. ข้อมูลทั่วไป (General Information)')
    add_kv_table(doc, [
        ('หมายเลข Incident', '{{ticket_id}}', 'Reference', '{{reference_id}}'),
        ('วันที่/เวลาที่พบเหตุ', '{{incident_datetime}}', 'ชื่อ Incident/Event', '{{incident_name}}'),
        ('ประเภท', '{{classification}}', 'สถานะปัจจุบัน', '{{status}}'),
        ('ระดับความรุนแรง (SIEM)', '{{siem_severity}}', 'ระดับความรุนแรง (สกมช.)', '{{ncsa_severity}}'),
        ('หมวดหมู่ภัยคุกคาม', '{{category}}', 'รายละเอียดหมวดหมู่', '{{category_detail}}'),
        ('ผู้รายงาน', '{{reporter}}', 'แหล่งข้อมูล (Log Source)', '{{log_source}}'),
        ('ช่องทางที่มา', '{{source_channel}}', 'เรื่องที่ดำเนินการแล้ว', '{{actions_taken_summary}}'),
        ('การดำเนินการลำดับถัดไป', '{{next_steps_summary}}', 'Template', '{{template_version}}'),
    ])

    section_heading(doc, '2. รายละเอียดเหตุการณ์ (Incident Description)')
    add_single_table(doc, [('รายละเอียดเหตุการณ์', '{{incident_description}}')])

    section_heading(doc, '3. Scope ทรัพย์สินที่ได้รับผลกระทบ')
    add_kv_table(doc, [
        ('ระบบ/บริการ', '{{system_name}}', 'หน่วยงานเจ้าของทรัพย์สิน', '{{asset_owner}}'),
        ('Host Name', '{{host_name}}', 'IP Address', '{{ip_address}}'),
        ('MAC Address', '{{mac_address}}', 'Operating System', '{{operating_system}}'),
        ('ประเภททรัพย์สิน', '{{asset_type}}', 'มีการกระจายไปยังจุดอื่น', '{{spread_to_others}}'),
    ])

    section_heading(doc, '4. Indicators of Compromise หรือหลักฐานที่พบ')
    add_single_table(doc, [
        ('Destination IP', '{{destination_ip}}'),
        ('IoC Details', '{{ioc_details}}'),
        ('Evidence / Attachments', '{{evidence_files}}'),
        ('MITRE ATT&CK Phases', '{{mitre_phases}}'),
    ])

    section_heading(doc, '5. สิ่งที่ต้องดำเนินการ (Containment)')
    add_single_table(doc, [
        ('สิ่งที่ต้องดำเนินการ', '{{action_required}}'),
        ('ข้อควรระวัง', '{{action_precautions}}'),
    ])

    section_heading(doc, '6. สรุปผลการดำเนินการแก้ไข')
    add_single_table(doc, [
        ('ผลการตรวจสอบ', '{{remediation_summary}}'),
        ('มาตรการควบคุม', '{{containment_report}}'),
    ])

    section_heading(doc, '7. Sign-off')
    add_signature_table(doc)
    add_appendix(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == '__main__':
    build()
