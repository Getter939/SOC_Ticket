"""Build the v2 incident-report DOCX template in NT (National Telecom) house style.

The layout mirrors the official NT SOC incident-report form: a gold title
banner, light-yellow section bands, gray label columns, the NT logo in every
page header, a fixed footer, and checkbox form fields. The document is a
*template*: values and checkbox states are filled per ticket by
``apps.incidents.reports`` via ``{{placeholder}}`` substitution.

Every ``{{placeholder}}`` is authored as its own run so run-level substitution
can preserve per-run fonts — checkbox glyphs render in DejaVu Sans (which has
☐/☑), Thai and Latin text in TH Sarabun New.

Run from the repo root:  python scripts/build_report_template_v2.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR))
# Shared, Django-free static content — single source of truth with the HTML view.
from apps.incidents.report_content import (  # noqa: E402
    APPENDIX_CATEGORIES,
    APPENDIX_INTRO,
    FOOTER_LEFT,
    FOOTER_RIGHT,
    SECTION1_ROWS,
    SECTION3_ROWS,
    SECTION4_ROWS,
    SECTION_TITLES,
)
TEMPLATES_DIR = BASE_DIR / 'apps' / 'incidents' / 'report_templates'
OUTPUT_PATH = TEMPLATES_DIR / 'report_template_v2.docx'
LOGO_PATH = TEMPLATES_DIR / 'assets' / 'nt_logo.png'

# NT palette (sampled from the reference form)
GOLD = 'FFD100'          # title banner
SECTION_YELLOW = 'FFF9C4'  # section heading bands
LABEL_GRAY = 'F2F2F2'    # table label columns
TEXT = '404040'          # body text
BORDER = 'BFBFBF'
MUTED = '5B6775'

BODY_FONT = 'TH Sarabun New'
SYMBOL_FONT = 'DejaVu Sans'   # supplies ☐/☑ glyphs

# ── low-level helpers ──────────────────────────────────────────────────── #

def set_run_font(run, name=BODY_FONT, size=14, color=TEXT, bold=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_borders(table, color=BORDER, size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tbl_pr.append(borders)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def _first_paragraph(cell):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_runs(paragraph, segments):
    """segments: list of (text, font, size, color, bold)."""
    for text, font, size, color, bold in segments:
        run = paragraph.add_run(text)
        set_run_font(run, name=font, size=size, color=color, bold=bold)


def checkbox_segments(options):
    """Build (glyph-placeholder, label) run segments for a checkbox row.

    options: list of (placeholder_key, label). Each glyph run is DejaVu Sans so
    the substituted ☐/☑ renders; each label run is the body font.
    """
    segments = []
    for i, (key, label) in enumerate(options):
        if i:
            segments.append(('   ', BODY_FONT, 14, TEXT, False))
        segments.append((f'{{{{{key}}}}}', SYMBOL_FONT, 14, TEXT, False))
        segments.append((f' {label}', BODY_FONT, 14, TEXT, False))
    return segments


def docx_rows(table):
    """Render a shared row table (report_content) into add_kv_table() input.

    The same tables drive apps.incidents.reports for the HTML/PDF preview, so a
    label or an option order is written once and both outputs stay in step.
    Here a 'kv' row becomes a {{placeholder}} and a 'checks' row becomes
    checkbox runs; there they become values and booleans.
    """
    rendered = []
    for kind, label, spec in table:
        if kind == 'kv':
            rendered.append((label, f'{{{{{spec}}}}}'))
        else:
            rendered.append((label, checkbox_segments(spec)))
    return rendered


# ── page furniture ─────────────────────────────────────────────────────── #

def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.3)

    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn('w:ascii'), BODY_FONT)
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), BODY_FONT)
    normal._element.rPr.rFonts.set(qn('w:cs'), BODY_FONT)
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1


def build_header(doc):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.7))
    else:
        run = p.add_run('NT')
        set_run_font(run, size=18, color=TEXT, bold=True)


def build_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.6))
    table.autofit = False
    clear_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(3.6)
    right.width = Inches(3.0)
    lp = _first_paragraph(left)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(lp, [(FOOTER_LEFT, BODY_FONT, 10, MUTED, False)])
    rp = _first_paragraph(right)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_runs(rp, [(FOOTER_RIGHT, BODY_FONT, 10, MUTED, False)])
    # Remove the stray empty paragraph the footer starts with.
    empty = footer.paragraphs[0]
    empty._p.getparent().remove(empty._p)


# ── content blocks ─────────────────────────────────────────────────────── #

def add_title_banner(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.6])
    clear_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, GOLD)
    set_cell_margins(cell, top=140, bottom=140)
    p = _first_paragraph(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, [('INCIDENT REPORT: Containment', BODY_FONT, 22, '1F2933', True)])
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    add_runs(p2, [('แบบฟอร์มรายงานเหตุการณ์ผิดปกติ', BODY_FONT, 16, '1F2933', True)])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_band(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.6])
    clear_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SECTION_YELLOW)
    set_cell_margins(cell, top=70, bottom=70)
    p = _first_paragraph(cell)
    p.paragraph_format.keep_with_next = True
    add_runs(p, [(text, BODY_FONT, 15, '1F2933', True)])


def add_kv_table(doc, rows):
    """rows: list of (label, value_segments). value_segments is either a
    placeholder string (single body run) or a list of run segments."""
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_widths(table, [2.3, 4.3])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, LABEL_GRAY)
        lp = _first_paragraph(left)
        add_runs(lp, [(label, BODY_FONT, 14, TEXT, True)])
        rp = _first_paragraph(right)
        if isinstance(value, str):
            add_runs(rp, [(value, BODY_FONT, 14, TEXT, False)])
        else:
            add_runs(rp, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_freetext_box(doc, placeholder, min_height_pt=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.6])
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_margins(cell, top=100, bottom=100)
    p = _first_paragraph(cell)
    add_runs(p, [(placeholder, BODY_FONT, 14, TEXT, False)])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_remediation_section(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.6])
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_margins(cell, top=90, bottom=90)
    # A fixed 15-item remediation checklist used to be printed here. It was
    # removed — it was never ticked by the system and section 6 already carries
    # the real, data-driven containment checklist. Section 8 is now just the
    # two free-text results below.
    gap = _first_paragraph(cell)
    add_runs(gap, [('ผลการตรวจสอบ / Investigation Findings:', BODY_FONT, 14, TEXT, True)])
    p = cell.add_paragraph()
    add_runs(p, [('{{remediation_summary}}', BODY_FONT, 14, TEXT, False)])
    p2 = cell.add_paragraph()
    add_runs(p2, [('มาตรการควบคุม / Countermeasure:', BODY_FONT, 14, TEXT, True)])
    p3 = cell.add_paragraph()
    add_runs(p3, [('{{containment_report}}', BODY_FONT, 14, TEXT, False)])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_signoff(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [3.3, 3.3])
    clear_table_borders(table)
    blocks = [
        ('{{signoff_admin}}', 'ผู้ดำเนินการแก้ไข'),
        ('{{signoff_approver}}', 'ผู้อนุมัติ'),
    ]
    for col, (name_ph, role) in enumerate(blocks):
        cell = table.rows[0].cells[col]
        lines = [
            '..........................................',
            name_ph,
            f'( {role} )',
            'วันที่ .......... / .......... / ..........',
        ]
        for i, text in enumerate(lines):
            p = _first_paragraph(cell) if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, [(text, BODY_FONT, 14, TEXT, False)])
    # Keep the whole sign-off block together on one page.
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def add_appendix(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading = doc.add_paragraph()
    add_runs(heading, [('*หมวดหมู่ของภัยคุกคามทางไซเบอร์', BODY_FONT, 15, '1F2933', True)])
    intro = doc.add_paragraph()
    add_runs(intro, [(APPENDIX_INTRO, BODY_FONT, 13, TEXT, False)])
    sub = doc.add_paragraph()
    add_runs(sub, [('ข้อ ๑ การจำแนกหมวดหมู่ของภัยคุกคามทางไซเบอร์', BODY_FONT, 14, TEXT, True)])

    table = doc.add_table(rows=len(APPENDIX_CATEGORIES) + 1, cols=2)
    set_table_widths(table, [1.0, 5.6])
    set_table_borders(table)
    head_no, head_desc = table.rows[0].cells
    set_cell_shading(head_no, LABEL_GRAY)
    set_cell_shading(head_desc, LABEL_GRAY)
    add_runs(_first_paragraph(head_no), [('หมวดหมู่', BODY_FONT, 14, TEXT, True)])
    add_runs(_first_paragraph(head_desc), [('คำอธิบาย', BODY_FONT, 14, TEXT, True)])
    for i, (num, desc) in enumerate(APPENDIX_CATEGORIES, start=1):
        no_cell, desc_cell = table.rows[i].cells
        pno = _first_paragraph(no_cell)
        pno.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(pno, [(num, BODY_FONT, 14, TEXT, False)])
        add_runs(_first_paragraph(desc_cell), [(desc, BODY_FONT, 14, TEXT, False)])


# ── document ───────────────────────────────────────────────────────────── #

def build(output_path=OUTPUT_PATH):
    doc = Document()
    style_doc(doc)
    build_header(doc)
    build_footer(doc)

    add_title_banner(doc)

    add_section_band(doc, f'1. {SECTION_TITLES["1"]}')
    add_kv_table(doc, docx_rows(SECTION1_ROWS))

    add_section_band(doc, f'2. {SECTION_TITLES["2"]}')
    add_freetext_box(doc, '{{incident_description}}')

    add_section_band(doc, f'3. {SECTION_TITLES["3"]}')
    add_kv_table(doc, docx_rows(SECTION3_ROWS))

    add_section_band(doc, f'4. {SECTION_TITLES["4"]}')
    add_kv_table(doc, docx_rows(SECTION4_ROWS))

    add_section_band(doc, f'5. {SECTION_TITLES["5"]}')
    add_freetext_box(doc, '{{evidence_log}}')

    add_section_band(doc, '6. สิ่งที่ต้องดำเนินการ (Containment)')
    add_freetext_box(doc, '{{action_required}}')

    add_section_band(doc, '7. ข้อควรระวังในการดำเนินการ')
    add_freetext_box(doc, '{{action_precautions}}')

    add_section_band(doc, '8. สรุปผลการดำเนินการแก้ไข')
    add_remediation_section(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_signoff(doc)

    add_appendix(doc)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)


if __name__ == '__main__':
    build()
