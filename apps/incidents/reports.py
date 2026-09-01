import base64
import hashlib
import logging
import re
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote, urlparse

from django.template.loader import render_to_string
from django.utils import timezone
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.fonts import addMapping
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFError, TTFont
from .models import Ticket
from .report_content import (
    APPENDIX_CATEGORIES,
    APPENDIX_INTRO,
    EVENT_FOOTER_RIGHT,
    EVENT_SECTION1_ROWS,
    FOOTER_LEFT,
    FOOTER_RIGHT,
    SECTION1_ROWS,
    SECTION3_ROWS,
    SECTION4_ROWS,
    SECTION_TITLES,
)


logger = logging.getLogger(__name__)

REPORT_TEMPLATE_VERSION = 'v2'
REPORT_TEMPLATE_NAME = f'report_template_{REPORT_TEMPLATE_VERSION}.docx'
EVENT_REPORT_TEMPLATE_VERSION = 'event-v1'
EVENT_REPORT_TEMPLATE_NAME = 'event_report_template_v1.docx'
REPORT_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
PDF_CONTENT_TYPE = 'application/pdf'
REPORT_TEMPLATE_PATH = Path(__file__).resolve().parent / 'report_templates' / REPORT_TEMPLATE_NAME
EVENT_REPORT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / 'report_templates' / EVENT_REPORT_TEMPLATE_NAME
)
REPORT_PREVIEW_TEMPLATE = 'incidents/report_preview.html'
REPORT_FONT_NAME = 'ReportUnicode'
REPORT_FONT_DIR = Path(__file__).resolve().parent / 'report_templates' / 'fonts'
# TH Sarabun New is bundled under report_templates/fonts/ (see NOTICE.md) so the
# PDF renders Thai on any host. Each entry: (bold, italic, filename, reportlab
# font name) — real weights instead of faux-mapping everything to the regular.
REPORT_FONT_VARIANTS = (
    (0, 0, 'THSarabunNew.ttf', REPORT_FONT_NAME),
    (1, 0, 'THSarabunNew-Bold.ttf', f'{REPORT_FONT_NAME}-Bold'),
    (0, 1, 'THSarabunNew-Italic.ttf', f'{REPORT_FONT_NAME}-Italic'),
    (1, 1, 'THSarabunNew-BoldItalic.ttf', f'{REPORT_FONT_NAME}-BoldItalic'),
)
# Last-resort system paths, only consulted if the bundled regular face is
# missing — keeps Thai rendering alive on a host where the bundle was stripped.
REPORT_FONT_FALLBACKS = (
    Path('C:/Windows/Fonts/THSarabunNew.ttf'),
    Path('/usr/share/fonts/truetype/thai/THSarabunNew.ttf'),
    Path('/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf'),
    Path('/usr/share/fonts/truetype/thai/Garuda.ttf'),
)
# DejaVu Sans supplies the ballot-box glyphs (☐/☑) that TH Sarabun New lacks;
# used only for the checkbox marks in the PDF path.
SYMBOL_FONT_NAME = 'ReportSymbol'
SYMBOL_FONT_VARIANTS = (
    (0, 'DejaVuSans.ttf', SYMBOL_FONT_NAME),
    (1, 'DejaVuSans-Bold.ttf', f'{SYMBOL_FONT_NAME}-Bold'),
)

CHECKED = '☑'
UNCHECKED = '☐'
# Font Word uses for the ballot glyphs in a filled DOCX (TH Sarabun lacks them).
DOCX_SYMBOL_FONT = 'DejaVu Sans'
_BALLOT_CHARS = frozenset('☑☐☒')


def _chk(flag):
    return CHECKED if flag else UNCHECKED

REPORT_LOGO_PATH = Path(__file__).resolve().parent / 'report_templates' / 'assets' / 'nt_logo.png'


@dataclass(frozen=True)
class GeneratedTicketReport:
    filename: str
    content: bytes
    content_type: str = REPORT_CONTENT_TYPE

    def as_file(self):
        return BytesIO(self.content)


def generate_ticket_report(ticket_id, generated_by=None, hide_empty=True):
    ticket = _load_ticket(ticket_id)
    generated_at = timezone.now()
    context = build_ticket_report_context(ticket, generated_at=generated_at)
    template_path = (
        EVENT_REPORT_TEMPLATE_PATH if _is_event_report(ticket) else REPORT_TEMPLATE_PATH
    )
    doc = Document(template_path)
    if hide_empty:
        _remove_empty_docx_fields(doc, context)
        _renumber_docx_section_one(doc)
    _replace_placeholders(doc, context)

    output = BytesIO()
    doc.save(output)
    content = output.getvalue()
    digest = hashlib.sha256(content).hexdigest()

    template_version = _report_template_version(ticket)
    _record_export_metadata(
        ticket, generated_by, generated_at, digest,
        report_format='docx', template_version=template_version,
    )

    filename = f'report_{ticket.ticket_id}_{template_version}.docx'
    return GeneratedTicketReport(filename=filename, content=content)


def generate_ticket_report_pdf(
    ticket_id, generated_by=None, base_url=None, hide_empty=True,
):
    ticket = _load_ticket(ticket_id)
    generated_at = timezone.now()
    context = build_ticket_report_render_context(
        ticket,
        generated_at=generated_at,
        show_report_actions=False,
        hide_empty=hide_empty,
    )
    html = render_to_string(REPORT_PREVIEW_TEMPLATE, context)
    content = _render_pdf_from_html(html, base_url=base_url)
    digest = hashlib.sha256(content).hexdigest()

    template_version = _report_template_version(ticket)
    _record_export_metadata(
        ticket, generated_by, generated_at, digest,
        report_format='pdf', template_version=template_version,
    )

    filename = f'report_{ticket.ticket_id}_{template_version}.pdf'
    return GeneratedTicketReport(
        filename=filename,
        content=content,
        content_type=PDF_CONTENT_TYPE,
    )


def build_ticket_report_render_context(
    ticket, generated_at=None, show_report_actions=True, hide_empty=True,
):
    is_event_report = _is_event_report(ticket)
    report = build_ticket_report_context(ticket, generated_at=generated_at)
    return {
        'ticket': ticket,
        'report': report,
        'sections': build_ticket_report_sections(
            report, ticket, hide_empty=hide_empty,
        ),
        'appendix_categories': APPENDIX_CATEGORIES,
        'appendix_intro': APPENDIX_INTRO,
        'footer_left': FOOTER_LEFT,
        'footer_right': EVENT_FOOTER_RIGHT if is_event_report else FOOTER_RIGHT,
        'is_event_report': is_event_report,
        'browser_title': 'Event Report' if is_event_report else 'Incident Report',
        'report_heading': (
            'Alert Event REPORT' if is_event_report else 'INCIDENT REPORT: Containment'
        ),
        'report_subtitle': (
            'แบบฟอร์มแจ้งเหตุการณ์ผิดปกติ'
            if is_event_report else 'แบบฟอร์มรายงานเหตุการณ์ผิดปกติ'
        ),
        'nt_logo': _logo_data_uri(),
        'show_report_actions': show_report_actions,
        'hide_empty': hide_empty,
    }


def build_ticket_report_context(ticket, generated_at=None):
    generated_at = generated_at or timezone.now()
    asset = ticket.asset_type
    asset_known = asset in {'Computer', 'Server', 'Network Device'}
    is_event_report = _is_event_report(ticket)
    context = {
        # The official report number carries a presentation-only classification
        # suffix.  The Ticket Reference itself remains immutable in the database.
        'ticket_id': _report_ticket_id(ticket),
        # Section 1 prints this in the Thai style used on the paper form.
        'incident_datetime': _format_dt_thai(ticket.incident_datetime),
        'incident_name': _value(ticket.incident_name),
        'category': _value(ticket.get_detailed_issue_display()),
        'reporter': _user_label(ticket.created_by, include_phone=True),
        'log_source': _value(ticket.log_source),
        'status': _value(ticket.get_status_display()),
        'actions_taken_summary': _value(ticket.actions_taken_summary),
        'next_steps_summary': _value(ticket.next_steps_summary),
        'incident_description': _value(ticket.issue_description),
        'host_ip': _host_ip(ticket),
        'system_name': _value(ticket.device_name),
        'asset_owner': _value(ticket.asset_owner),
        'asset_owner_name': _value(ticket.asset_owner_name),
        'host_name': _value(ticket.device_name),
        'ip_address': _value(ticket.ip_address),
        'operating_system': _value(ticket.operating_system),
        'ioc_process': _value(ticket.ioc_details),
        'ioc_command': '-',
        'ioc_hash': '-',
        'ioc_ip': _value(ticket.destination_ip),
        'ioc_user': _value(ticket.ioc_user),
        'evidence_log': _evidence_log(ticket),
        'action_required': _containment_checklist_flat(ticket),
        'action_precautions': _value(ticket.action_precautions),
        'remediation_summary': _value(ticket.remediation_summary),
        'containment_report': _value(ticket.containment_report),
        'signoff_admin': _signoff_name(ticket.assigned_admin),
        'signoff_approver': _signoff_name(ticket.approved_by),
        'template_version': _report_template_version(ticket),
        'generated_at': _format_dt(generated_at),
        # Checkbox states (☑/☐) driven by the ticket's actual values.
        'chk_class_event': _chk(ticket.classification == Ticket.CLASSIFICATION_EVENT),
        'chk_class_incident': _chk(ticket.classification == Ticket.CLASSIFICATION_INCIDENT),
        'chk_sev_critical': _chk(ticket.severity == 'Critical'),
        'chk_sev_high': _chk(ticket.severity == 'High'),
        'chk_sev_medium': _chk(ticket.severity == 'Medium'),
        'chk_sev_low': _chk(ticket.severity == 'Low'),
        'chk_imp_high': _chk(ticket.is_emergency),
        'chk_imp_normal': _chk(not is_event_report and not ticket.is_emergency),
        'chk_spread_yes': _chk(ticket.spread_to_others is True),
        'chk_spread_no': _chk(ticket.spread_to_others is False),
        'chk_ncsa_critical': _chk(ticket.ncsa_severity == Ticket.NCSA_SEVERITY_CRITICAL),
        'chk_ncsa_severe': _chk(ticket.ncsa_severity == Ticket.NCSA_SEVERITY_SEVERE),
        'chk_ncsa_nonsevere': _chk(ticket.ncsa_severity == Ticket.NCSA_SEVERITY_NON_SEVERE),
        'chk_asset_computer': _chk(asset == 'Computer'),
        'chk_asset_server': _chk(asset == 'Server'),
        'chk_asset_network': _chk(asset == 'Network Device'),
        'chk_asset_unknown': _chk(not asset_known),
    }
    if is_event_report:
        context['chk_imp_general'] = _chk(not ticket.is_emergency)
    return context


def _is_event_report(ticket):
    return ticket.classification == Ticket.CLASSIFICATION_EVENT


def _report_template_version(ticket):
    return EVENT_REPORT_TEMPLATE_VERSION if _is_event_report(ticket) else REPORT_TEMPLATE_VERSION


def _report_ticket_id(ticket):
    suffix = {
        Ticket.CLASSIFICATION_EVENT: 'E',
        Ticket.CLASSIFICATION_INCIDENT: 'I',
    }.get(ticket.classification)
    ticket_id = _value(ticket.ticket_id)
    return f'{ticket_id}-{suffix}' if suffix else ticket_id


def build_ticket_report_sections(report, ticket, hide_empty=False):
    """Structured sections for the HTML/PDF preview, mirroring the v2 DOCX form.

    Row shapes consumed by report_preview.html:
      {'type': 'kv', 'label', 'value'}
      {'type': 'checks', 'label', 'options': [{'label', 'checked'}, ...]}
      {'type': 'text', 'value'}                     — full-width free-text box
    """
    def kv(label, value):
        return {'type': 'kv', 'label': label, 'value': value}

    def checks(label, options):
        return {'type': 'checks', 'label': label,
                'options': [{'label': lbl, 'checked': flag} for lbl, flag in options]}

    def text(value):
        return {'type': 'text', 'value': value}

    def rows_from(table):
        """Render a shared row table (report_content) into preview rows.

        Checkbox state is read back out of the ``chk_*`` keys the context
        already carries, rather than recomputing the same comparisons here —
        one source for the ticked state, and the DOCX builder walks the same
        table with the same keys.
        """
        built = []
        for kind, label, spec in table:
            if kind == 'kv':
                built.append(kv(label, report[spec]))
            else:
                built.append(checks(label, [
                    (option_label, report[chk_key] == CHECKED)
                    for chk_key, option_label in spec
                ]))
        return built

    if _is_event_report(ticket):
        sections = [{
            'number': '1',
            'title': SECTION_TITLES['1'],
            'rows': rows_from(EVENT_SECTION1_ROWS),
        }]
    else:
        sections = [
            {'number': '1', 'title': SECTION_TITLES['1'],
             'rows': rows_from(SECTION1_ROWS)},
            {'number': '2', 'title': SECTION_TITLES['2'], 'rows': [
                text(report['incident_description'])]},
            {'number': '3', 'title': SECTION_TITLES['3'],
             'rows': rows_from(SECTION3_ROWS)},
            {'number': '4', 'title': SECTION_TITLES['4'],
             'rows': rows_from(SECTION4_ROWS)},
            {'number': '5', 'title': SECTION_TITLES['5'],
             'rows': [text(report['evidence_log'])]},
            {'number': '6', 'title': SECTION_TITLES['6'], 'rows': [
                _containment_checklist_row(ticket) or text(report['action_required'])]},
            {'number': '7', 'title': SECTION_TITLES['7'], 'rows': [
                text(report['action_precautions'])]},
            {'number': '8', 'title': SECTION_TITLES['8'], 'rows': [
                kv('ผลการตรวจสอบ / Investigation Findings', report['remediation_summary']),
                kv('มาตรการควบคุม / Countermeasure', report['containment_report']),
            ]},
        ]
    if not hide_empty:
        return sections

    for section in sections:
        section['rows'] = [
            row for row in section['rows']
            if row['type'] not in {'kv', 'text'} or row['value'] != '-'
        ]
        if section['number'] == '1':
            for index, row in enumerate(section['rows'], start=1):
                row['label'] = re.sub(
                    r'^1\.\d+(?=\s)', f'1.{index}', row['label'], count=1,
                )
    return [section for section in sections if section['rows']]


def _record_export_metadata(
    ticket, generated_by, generated_at, digest, report_format, template_version,
):
    generated_by_id = getattr(generated_by, 'pk', None)
    Ticket.objects.filter(pk=ticket.pk).update(
        report_template_version=template_version,
        report_format=report_format,
        report_generated_by_id=generated_by_id,
        report_generated_at=generated_at,
        report_ticket_updated_at=ticket.updated_at,
        report_sha256=digest,
    )


def _render_pdf_from_html(html, base_url=None):
    xhtml2pdf_default, pisa = _load_pdf_dependencies()
    _register_pdf_font(xhtml2pdf_default)
    output = BytesIO()
    result = pisa.CreatePDF(
        src=html,
        dest=output,
        encoding='utf-8',
        path=base_url,
        link_callback=_resolve_pdf_resource,
    )
    if result.err:
        raise ValueError('Unable to render incident report PDF')
    return output.getvalue()


def _load_pdf_dependencies():
    """Load PDF support only when an export is requested.

    ``python-bidi`` ships a compiled extension on Windows. Some managed
    workstations block that extension through application-control policy;
    importing it while Django resolves URLs would otherwise prevent the whole
    site from starting, including features that do not create PDFs.
    """
    try:
        from xhtml2pdf import default as xhtml2pdf_default
        from xhtml2pdf import pisa
    except ImportError as exc:
        raise RuntimeError(
            'PDF export is unavailable because its xhtml2pdf dependency could '
            'not be loaded. Ask IT to allow the python-bidi package, then retry.'
        ) from exc
    return xhtml2pdf_default, pisa


def _register_pdf_font(xhtml2pdf_default=None):
    if xhtml2pdf_default is None:
        xhtml2pdf_default, _ = _load_pdf_dependencies()
    xhtml2pdf_default.DEFAULT_FONT[REPORT_FONT_NAME.lower()] = REPORT_FONT_NAME
    if REPORT_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return

    registered = {}  # (bold, italic) -> reportlab font name
    for bold, italic, filename, font_name in REPORT_FONT_VARIANTS:
        font_path = REPORT_FONT_DIR / filename
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        except (OSError, TTFError):
            logger.warning('Could not register bundled report font %s', font_path)
            continue
        registered[(bold, italic)] = font_name

    if (0, 0) not in registered:
        # Bundled regular face missing/unreadable — fall back to a system font
        # so Thai text still renders instead of blank boxes.
        for font_path in REPORT_FONT_FALLBACKS:
            if not font_path.exists():
                continue
            try:
                pdfmetrics.registerFont(TTFont(REPORT_FONT_NAME, str(font_path)))
            except (OSError, TTFError):
                continue
            registered[(0, 0)] = REPORT_FONT_NAME
            logger.warning(
                'Bundled report font missing; using system fallback %s', font_path,
            )
            break

    if (0, 0) not in registered:
        logger.error(
            'No Thai-capable report font found (looked in %s and system paths); '
            'PDF Thai text will render as blank boxes',
            REPORT_FONT_DIR,
        )
        return

    regular = registered[(0, 0)]
    for bold in (0, 1):
        for italic in (0, 1):
            addMapping(REPORT_FONT_NAME, bold, italic, registered.get((bold, italic), regular))

    _register_symbol_font(xhtml2pdf_default)


def _register_symbol_font(xhtml2pdf_default=None):
    """Register DejaVu Sans for the checkbox glyphs the PDF path needs."""
    if xhtml2pdf_default is None:
        xhtml2pdf_default, _ = _load_pdf_dependencies()
    xhtml2pdf_default.DEFAULT_FONT[SYMBOL_FONT_NAME.lower()] = SYMBOL_FONT_NAME
    if SYMBOL_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return
    symbol_regular = None
    for bold, filename, font_name in SYMBOL_FONT_VARIANTS:
        font_path = REPORT_FONT_DIR / filename
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        except (OSError, TTFError):
            continue
        if bold == 0:
            symbol_regular = font_name
    if symbol_regular is None:
        logger.warning(
            'Checkbox symbol font (DejaVu Sans) missing in %s; PDF checkboxes '
            'may render as blank boxes', REPORT_FONT_DIR,
        )


def _resolve_pdf_resource(uri, rel):
    parsed = urlparse(uri)
    if parsed.scheme == 'file':
        path = unquote(parsed.path)
        if re.match(r'^/[A-Za-z]:/', path):
            path = path[1:]
        return path
    return uri


def _load_ticket(ticket_id):
    return (
        Ticket.objects
        .select_related(
            'project_incident', 'created_by', 'created_by__profile',
            'verified_by', 'approved_by', 'assigned_admin',
        )
        .prefetch_related('attachments')
        .get(pk=ticket_id)
    )


_DOCX_OPTIONAL_SECTION_FIELDS = {
    f'2. {SECTION_TITLES["2"]}': ('incident_description',),
    f'4. {SECTION_TITLES["4"]}': (
        'ioc_process', 'ioc_command', 'ioc_hash', 'ioc_ip', 'ioc_user',
    ),
    f'5. {SECTION_TITLES["5"]}': ('evidence_log',),
    f'6. {SECTION_TITLES["6"]}': ('action_required',),
    f'7. {SECTION_TITLES["7"]}': ('action_precautions',),
    f'8. {SECTION_TITLES["8"]}': (
        'remediation_summary', 'containment_report',
    ),
}


def _remove_empty_docx_fields(doc, context):
    """Remove report rows whose rendered value is only the ``-`` placeholder.

    Most fields occupy one table row. Section 8 has two fields in one row, so
    an empty field there is represented by a label paragraph followed by its
    value paragraph; remove that pair while preserving its populated sibling.
    Section headings are removed when every field in that section is empty.
    """
    empty_headings = {
        heading
        for heading, keys in _DOCX_OPTIONAL_SECTION_FIELDS.items()
        if all(context.get(key) == '-' for key in keys)
    }

    for table in _iter_docx_tables(doc):
        for row in list(table.rows):
            row_text = '\n'.join(cell.text for cell in row.cells)
            if row_text.strip() in empty_headings:
                _remove_docx_row(row)
                continue

            keys = set(re.findall(r'\{\{([^}]+)\}\}', row_text))
            if keys and all(context.get(key) == '-' for key in keys):
                _remove_docx_row(row)
                continue

            _remove_empty_docx_paragraph_pairs(row, context)


def _iter_docx_tables(doc):
    def walk(table):
        yield table
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    yield from walk(nested)

    for table in doc.tables:
        yield from walk(table)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for table in header_footer.tables:
                yield from walk(table)


def _remove_docx_row(row):
    row._tr.getparent().remove(row._tr)


def _remove_empty_docx_paragraph_pairs(row, context):
    seen_cells = set()
    for cell in row.cells:
        if id(cell._tc) in seen_cells:
            continue
        seen_cells.add(id(cell._tc))
        paragraphs = list(cell.paragraphs)
        for index in range(len(paragraphs) - 1, -1, -1):
            paragraph = paragraphs[index]
            keys = set(re.findall(r'\{\{([^}]+)\}\}', paragraph.text))
            if not keys or not all(context.get(key) == '-' for key in keys):
                continue

            element = paragraph._element
            element.getparent().remove(element)
            if index == 0:
                continue
            label = paragraphs[index - 1]
            if '{{' not in label.text and label._element.getparent() is not None:
                label._element.getparent().remove(label._element)


def _renumber_docx_section_one(doc):
    index = 0
    for table in _iter_docx_tables(doc):
        for row in table.rows:
            matched = False
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    match = re.match(r'^1\.\d+(?=\s)', paragraph.text)
                    if match is None:
                        continue
                    index += 1
                    _replace_docx_paragraph_prefix(
                        paragraph, match.group(0), f'1.{index}',
                    )
                    matched = True
                    break
                if matched:
                    break


def _replace_docx_paragraph_prefix(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return

    # The current template keeps each label in one run. This fallback handles
    # a future template that splits the numeric prefix across styled runs.
    if not paragraph.runs:
        return
    text = paragraph.text
    paragraph.runs[0].text = f'{new}{text[len(old):]}'
    for run in paragraph.runs[1:]:
        run.text = ''


def _replace_placeholders(doc, context):
    # Replace within individual runs (not whole paragraphs) so each run keeps
    # its own font — the template authors every {{placeholder}} as its own run,
    # letting checkbox glyphs (DejaVu Sans) and labels (TH Sarabun New) coexist
    # in one cell. run.text's setter turns \n into <w:br>, so multi-line values
    # keep their line breaks. A placeholder split across runs is left in place
    # and caught by the unresolved-placeholder check below.
    replacements = {f'{{{{{key}}}}}': str(value) for key, value in context.items()}
    for paragraph in _iter_paragraphs(doc):
        for run in paragraph.runs:
            text = run.text
            if '{{' not in text:
                continue
            replaced = text
            for placeholder, value in replacements.items():
                if placeholder in replaced:
                    replaced = replaced.replace(placeholder, value)
            if replaced != text:
                _fill_run(run, replaced)
    remaining = sorted({
        match
        for paragraph in _iter_paragraphs(doc)
        for match in re.findall(r'\{\{[^}]+\}\}', paragraph.text)
    })
    if remaining:
        raise ValueError(f'Unresolved report template placeholders: {", ".join(remaining)}')


def _fill_run(run, text):
    """Assign ``text`` to a run. Ballot glyphs (☑/☐/☒) are emitted in DejaVu Sans
    — TH Sarabun New has no such glyphs — while the rest keeps the run's own
    font. \\n becomes a Word line break. The common no-ballot case is a plain
    ``run.text = text`` (which itself converts \\n to <w:br>)."""
    if not any(ch in _BALLOT_CHARS for ch in text):
        run.text = text
        return

    r = run._r
    rpr = r.find(qn('w:rPr'))
    parent = r.getparent()
    index = parent.index(r)
    parent.remove(r)
    for offset, (is_ballot, chunk) in enumerate(_segment_ballots(text)):
        new_r = OxmlElement('w:r')
        if rpr is not None:
            new_rpr = deepcopy(rpr)
            if is_ballot:
                _force_run_font(new_rpr, DOCX_SYMBOL_FONT)
            new_r.append(new_rpr)
        _append_run_text(new_r, chunk)
        parent.insert(index + offset, new_r)


def _segment_ballots(text):
    """Split text into alternating (is_ballot, chunk) runs."""
    segments = []
    buf, buf_is_ballot = '', None
    for ch in text:
        is_ballot = ch in _BALLOT_CHARS
        if buf_is_ballot is None:
            buf, buf_is_ballot = ch, is_ballot
        elif is_ballot == buf_is_ballot:
            buf += ch
        else:
            segments.append((buf_is_ballot, buf))
            buf, buf_is_ballot = ch, is_ballot
    if buf:
        segments.append((buf_is_ballot, buf))
    return segments


def _append_run_text(r_el, chunk):
    """Append <w:t>/<w:br> children to a run element, turning \\n into breaks."""
    for i, part in enumerate(chunk.split('\n')):
        if i:
            r_el.append(OxmlElement('w:br'))
        if part:
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = part
            r_el.append(t)


def _force_run_font(rpr, name):
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), name)


def _iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                yield paragraph
            for table in header_footer.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _format_dt(value):
    if not value:
        return '-'
    if timezone.is_aware(value):
        value = timezone.localtime(value)
    return value.strftime('%d/%m/%Y %H:%M')


# Abbreviated Thai months, index 1-12, as written on the NT incident-report form.
_THAI_MONTHS_ABBR = (
    None, 'ม.ค.', 'ก.พ.', 'มี.ค.', 'เม.ย.', 'พ.ค.', 'มิ.ย.',
    'ก.ค.', 'ส.ค.', 'ก.ย.', 'ต.ค.', 'พ.ย.', 'ธ.ค.',
)


def _format_dt_thai(value):
    """Section 1's date style: ``2 / ก.ค. / 69  11:00 น.``

    Thai Buddhist year, abbreviated Thai month, two-digit year — matching the
    paper form the report is filed alongside. Deliberately limited to the
    section 1 date rows; the 'generated at' meta line and the report filename
    stay Gregorian so exports remain sortable and unambiguous.
    """
    if not value:
        return '-'
    if timezone.is_aware(value):
        value = timezone.localtime(value)
    buddhist_year = (value.year + 543) % 100
    return (
        f'{value.day} / {_THAI_MONTHS_ABBR[value.month]} / '
        f'{buddhist_year:02d}  {value:%H:%M} น.'
    )


def _value(value):
    if value is None:
        return '-'
    text = str(value).strip()
    return text or '-'


def _user_label(user, include_phone=False):
    if not user:
        return '-'
    label = user.get_full_name() or user.username
    if include_phone:
        phone = getattr(getattr(user, 'profile', None), 'phone', '')
        if phone:
            label = f'{label}, {phone}'
    return label


def _attachment_summary(ticket):
    parts = []
    for attachment in ticket.attachments.all():
        label = attachment.original_name
        if attachment.description:
            label = f'{label} - {attachment.description}'
        parts.append(label)
    return '\n'.join(parts) if parts else '-'


def _host_ip(ticket):
    parts = [p for p in (ticket.device_name, ticket.ip_address) if p]
    return ' / '.join(parts) if parts else '-'


def _evidence_log(ticket):
    parts = []
    attachments = _attachment_summary(ticket)
    if attachments != '-':
        parts.append(attachments)
    mitre = ', '.join(ticket.mitre_tactic_labels)
    if mitre:
        parts.append(f'MITRE ATT&CK: {mitre}')
    return '\n'.join(parts) if parts else '-'


def _signoff_name(user):
    if not user:
        return '(........................................................)'
    return f'( {user.get_full_name() or user.username} )'


def _logo_data_uri():
    if not REPORT_LOGO_PATH.exists():
        return ''
    encoded = base64.b64encode(REPORT_LOGO_PATH.read_bytes()).decode('ascii')
    return f'data:image/png;base64,{encoded}'


def _containment_checklist_flat(ticket):
    """Section-6 value for the DOCX: ☑/☐-prefixed items + trailing text, or the
    plain action_required when nothing is itemized."""
    items, trailing = ticket.containment_checklist_display()
    if not items:
        return _value(ticket.action_required)
    lines = [f'{_chk(item["done"])} {item["text"]}' for item in items]
    text = '\n'.join(lines)
    if trailing:
        text = f'{text}\n{trailing}'
    return text


def _containment_checklist_row(ticket):
    """Structured section-6 row for the HTML/PDF preview, or None to fall back to
    plain text when action_required has no checklist items."""
    items, trailing = ticket.containment_checklist_display()
    if not items:
        return None
    return {'type': 'containment_checklist', 'items': items, 'trailing': trailing}
