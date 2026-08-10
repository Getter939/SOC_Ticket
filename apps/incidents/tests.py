"""
Tests for the redesigned SOC ticketing workflow.

Test classes
────────────
1.  TicketVisibilityQuerysetTest  — Ticket.objects.visible_to() queryset scoping
2.  TicketVisibilityViewTest      — HTTP-level visibility enforcement
3.  WorkflowTransitionTest        — Every legal state-machine edge, every illegal edge
4.  WorkflowPermissionTest        — Per-transition role/tier permissions (positive + negative)
5.  T1ClassificationCreateTest    — Tier 1 Event/Incident disposition at creation
6.  Tier2EscalationTest           — Tier 2 return-only constraint (never assign / never create)
7.  ManagerRoutingTest            — requires_manager_verification (emergency flag only)
8.  EmergencyFlagTest             — emergency-flag permissions + audit
9.  AdminFieldAccessTest          — System Admin write access to containment/remediation fields
10. SignOffFieldsTest             — verified_by/at and approved_by/at are write-once
11. NotificationEmailTest         — Email notifications on AWAITING_CONTAINMENT transitions
12. WazuhTriageActionTest         — 2-action Tier 1 triage + required release reason
13. TriageWorkflowIntegrityTest   — manual-triage + wazuh-alert ticket creation
14. SuperuserAccessTest           — superuser bypass across the redesigned flow
15. AttachmentDownloadSecurityTest / AttachmentUploadLimitTest
16. TicketReportExportTest       — preview, DOCX/PDF generation + metadata

Run with:  py manage.py test apps.incidents --settings=config.settings_local
"""

import hashlib
import importlib.util
import json
import re
import shutil
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from datetime import datetime, timedelta

from django.conf import settings
from django.contrib.auth.models import User
from django.core import mail
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management import call_command
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from docx import Document
from pypdf import PdfReader

from apps.accounts.models import UserProfile
from apps.incidents import history
from apps.incidents import ola as ola_buckets
from apps.incidents.forms import (
    AdminAssignmentForm, AttachmentForm, ProjectIncidentTargetForm,
    ResponseRequestForm, SubtaskForm, TicketForm, TriageForm,
)
from apps.incidents.models import (
    ProjectIncident, ProjectIncidentAttachment, ProjectIncidentLog,
    StagedAttachment, ThreatGuidance, Ticket, TicketAttachment, TicketLog,
    TicketAlertLink, TicketSubtask, TriageRecord, bundle_suffix_for_index,
)
from apps.incidents.notifications import (
    notify_containment_required,
    notify_response_request_created,
    notify_response_request_completed,
)
from apps.incidents.views import (
    _can_create_ticket_from_triage, _can_delete_ticket_attachment,
    _can_upload_ticket_attachment,
)
from apps.incidents.reports import (
    REPORT_TEMPLATE_PATH, REPORT_TEMPLATE_VERSION,
    build_ticket_report_context, build_ticket_report_sections,
    generate_ticket_report, _iter_paragraphs,
)
from apps.wazuh_ingest.models import WazuhAlert


# ──────────────────────────────────────────────────────────────────────────── #
# Shared helpers                                                               #
# ──────────────────────────────────────────────────────────────────────────── #

def _make_user(username, role, department='Test', phone='000', **kwargs):
    """Create a User + UserProfile in one call. Pass tier='T1'/'T2' via kwargs."""
    user = User.objects.create_user(username=username, password='testpass123')
    UserProfile.objects.create(
        user=user, role=role, department=department, phone=phone, **kwargs
    )
    return user


def _make_t1(username='t1', **kwargs):
    return _make_user(username, UserProfile.ROLE_SOC_STAFF, tier=UserProfile.TIER_T1, **kwargs)


def _make_t2(username='t2', **kwargs):
    return _make_user(username, UserProfile.ROLE_SOC_STAFF, tier=UserProfile.TIER_T2, **kwargs)


def _make_forensic(username='forensic', **kwargs):
    return _make_user(username, UserProfile.ROLE_FORENSIC, **kwargs)


def _make_redteam_manager(username='redteam', **kwargs):
    return _make_user(username, UserProfile.ROLE_REDTEAM_MANAGER, **kwargs)


def _make_ticket(**kwargs):
    """Create a Ticket with sensible defaults (bypasses the state machine)."""
    defaults = dict(
        device_name='10.0.0.1',
        ip_address='192.168.0.1',
        issue_description='Test ticket',
    )
    defaults.update(kwargs)
    return Ticket.objects.create(**defaults)


def _ticket_post_data(**overrides):
    """A valid create_ticket POST payload."""
    data = {
        'classification': Ticket.CLASSIFICATION_INCIDENT,
        't1_route': TicketForm.ROUTE_ESCALATE_T2,
        'severity': 'High',
        'ncsa_severity': Ticket.NCSA_SEVERITY_SEVERE,
        # Required since the form stopped defaulting it to "now" client-side.
        'incident_datetime': timezone.localtime().strftime('%Y-%m-%dT%H:%M'),
        'log_source': 'Wazuh',
        'issue_type': 'SIEM',
        'detailed_issue': 'Investigating',
        'detailed_issue2': 'Investigating Other',
        'device_name': 'TEST-ENDPOINT-01',
        'issue_description': 'Confirmed suspicious activity.',
        'ip_address': '192.0.2.10',
    }
    data.update(overrides)
    return data


def _advance_helper_manager():
    """A shared SOC Manager for driving the pre-containment forward step in
    tests that don't otherwise need one. Get-or-create keeps it unique per DB."""
    return _make_user('adv_helper_mgr', UserProfile.ROLE_SOC_MANAGER) \
        if not User.objects.filter(username='adv_helper_mgr').exists() \
        else User.objects.get(username='adv_helper_mgr')


def _advance_to(ticket, target_status, t1, admin=None, mgr=None, t2=None):
    """
    Drive a ticket from its current status to target_status along the
    Incident → assign-admin happy path. The ticket first passes the SOC Manager
    pre-containment review (PENDING_MGR_TRIAGE), then the admin contains it and
    Tier 2 verifies; the final manager approval fires when the emergency flag
    requires it.
    """
    if ticket.created_by_id is None:
        ticket.created_by = t1
    if ticket.classification != Ticket.CLASSIFICATION_INCIDENT:
        ticket.classification = Ticket.CLASSIFICATION_INCIDENT
    # The admin lane is the path this helper drives; remember it so the manager
    # forward is legal.
    ticket.t1_route = Ticket.T1_ROUTE_ADMIN
    ticket.save(update_fields=['created_by', 'classification', 't1_route'])

    # A manager is always needed for the pre-containment forward; fall back to a
    # shared helper manager when the caller didn't supply one.
    forward_mgr = mgr or _advance_helper_manager()

    path = [
        Ticket.STATUS_NEW,
        Ticket.STATUS_PENDING_MGR_TRIAGE,
        Ticket.STATUS_AWAITING_CONTAINMENT,
        Ticket.STATUS_CONTAINMENT_REPORTED,
    ]
    if ticket.requires_manager_verification:
        path += [Ticket.STATUS_PENDING_MANAGER, Ticket.STATUS_APPROVED]
    else:
        path += [Ticket.STATUS_APPROVED]

    i = path.index(ticket.status)
    j = path.index(target_status)
    for step in path[i + 1: j + 1]:
        if step == Ticket.STATUS_PENDING_MGR_TRIAGE:
            ticket.transition_to(step, t1, 'route to SOC Manager review')
        elif step == Ticket.STATUS_AWAITING_CONTAINMENT:
            ticket.transition_to(step, forward_mgr, 'manager forwards to admin')
        elif step == Ticket.STATUS_CONTAINMENT_REPORTED:
            ticket.containment_report = 'Contained.'
            ticket.transition_to(step, admin, 'containment note')
        elif step == Ticket.STATUS_PENDING_MANAGER:
            ticket.transition_to(step, t2, 'T2 verified — route to manager')
        elif step == Ticket.STATUS_APPROVED:
            actor = mgr if ticket.requires_manager_verification else t2
            ticket.transition_to(step, actor, 'close')


def _docx_text(content):
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    parts.extend(
        p.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
    )
    return '\n'.join(parts)


class TicketReferenceTest(TestCase):
    def test_new_tickets_receive_monthly_soc_references(self):
        when = timezone.make_aware(datetime(2026, 7, 10, 8, 30))
        with patch('apps.incidents.models.timezone.now', return_value=when):
            first = _make_ticket()
            second = _make_ticket()

        self.assertEqual(first.ticket_id, 'SOC-202607-0001')
        self.assertEqual(second.ticket_id, 'SOC-202607-0002')


# ──────────────────────────────────────────────────────────────────────────── #
# 1. Visibility queryset tests                                                 #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketVisibilityQuerysetTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.soc_staff   = _make_t1('soc_staff')
        cls.soc_manager = _make_user('soc_manager', UserProfile.ROLE_SOC_MANAGER)
        cls.admin_a     = _make_user('admin_a',     UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin_b     = _make_user('admin_b',     UserProfile.ROLE_SYSTEM_ADMIN)
        cls.forensic    = _make_forensic('vis_forensic')
        cls.redteam     = _make_redteam_manager('vis_redteam')
        cls.no_profile  = User.objects.create_user(username='noprofile', password='testpass123')

        cls.ticket_a = _make_ticket(assigned_admin=cls.admin_a)
        cls.ticket_b = _make_ticket(assigned_admin=cls.admin_b)
        cls.ticket_unassigned = _make_ticket()

        # A forensic request on ticket_a makes it (and only it) visible to the
        # forensic analyst; a VA/PT request on ticket_b to the red-team manager.
        TicketSubtask.objects.create(
            ticket=cls.ticket_a, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=cls.forensic,
        )
        TicketSubtask.objects.create(
            ticket=cls.ticket_b, subtask_type=TicketSubtask.TYPE_VA_PT,
            title='Pentest', assigned_to=cls.redteam,
        )

    def test_soc_staff_sees_all_tickets(self):
        self.assertEqual(Ticket.objects.visible_to(self.soc_staff).count(), 3)

    def test_soc_manager_sees_all_tickets(self):
        self.assertEqual(Ticket.objects.visible_to(self.soc_manager).count(), 3)

    def test_system_admin_sees_only_own_ticket(self):
        qs = Ticket.objects.visible_to(self.admin_a)
        self.assertEqual(qs.count(), 1)
        self.assertEqual(qs.first(), self.ticket_a)

    def test_system_admin_cannot_see_other_admins_ticket(self):
        self.assertNotIn(self.ticket_b, Ticket.objects.visible_to(self.admin_a))

    def test_no_profile_sees_no_tickets(self):
        self.assertEqual(Ticket.objects.visible_to(self.no_profile).count(), 0)

    def test_forensic_sees_only_ticket_with_their_request(self):
        qs = Ticket.objects.visible_to(self.forensic)
        self.assertEqual(list(qs), [self.ticket_a])

    def test_redteam_manager_sees_only_ticket_with_their_request(self):
        qs = Ticket.objects.visible_to(self.redteam)
        self.assertEqual(list(qs), [self.ticket_b])

    def test_response_team_visibility_is_not_duplicated(self):
        # A second request assigned to the same responder must not double-count.
        # It has to be a type their role actually receives, or visible_to()
        # filters it out and the distinct() guard never gets exercised.
        TicketSubtask.objects.create(
            ticket=self.ticket_a, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='Second RCA', assigned_to=self.forensic,
        )
        self.assertEqual(self.ticket_a.subtasks.filter(assigned_to=self.forensic).count(), 2)
        self.assertEqual(Ticket.objects.visible_to(self.forensic).count(), 1)

    def test_mismatched_response_type_does_not_expose_ticket(self):
        # K9. Assignment alone must not be a key to the ticket: the request also
        # has to be one this role receives. Written with objects.create() on
        # purpose — that bypasses TicketSubtask.clean(), reproducing exactly what
        # a seed command, a data migration, or the Django admin could leave
        # behind. Read-time filtering is what has to catch it.
        TicketSubtask.objects.create(
            ticket=self.ticket_unassigned,
            subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='Misrouted RCA', assigned_to=self.redteam,
        )
        visible = Ticket.objects.visible_to(self.redteam)
        self.assertNotIn(self.ticket_unassigned, visible)
        self.assertEqual(list(visible), [self.ticket_b])  # only their VA/PT one

    def test_types_for_role_is_the_inverse_of_response_routing(self):
        self.assertEqual(
            TicketSubtask.types_for_role(UserProfile.ROLE_REDTEAM_MANAGER),
            frozenset({TicketSubtask.TYPE_VA_PT, TicketSubtask.TYPE_INFRA_SEC}),
        )
        self.assertEqual(
            TicketSubtask.types_for_role(UserProfile.ROLE_FORENSIC),
            frozenset({TicketSubtask.TYPE_FORENSIC_RCA}),
        )
        # Every response type is routed to exactly one role, and no type is lost.
        routed = set()
        for role in (UserProfile.ROLE_REDTEAM_MANAGER, UserProfile.ROLE_FORENSIC):
            routed |= TicketSubtask.types_for_role(role)
        self.assertEqual(routed, set(TicketSubtask.RESPONSE_TYPES))

    def test_types_for_role_is_empty_for_non_response_roles(self):
        # Safe to drop into a subtask_type__in= filter: an unrecognised role
        # matches nothing rather than everything.
        for role in (
            UserProfile.ROLE_SOC_STAFF, UserProfile.ROLE_SOC_MANAGER,
            UserProfile.ROLE_SYSTEM_ADMIN, UserProfile.ROLE_SYSTEM_OWNER,
            UserProfile.ROLE_EXECUTIVE, '',
        ):
            self.assertEqual(TicketSubtask.types_for_role(role), frozenset(), role)

    def test_response_team_legacy_subtask_does_not_expose_ticket(self):
        # Response-only access: being handed an ordinary Investigation subtask on
        # an unrelated ticket must NOT expose that ticket. Regression for the
        # over-broad visible_to filter (was: any assigned subtask).
        TicketSubtask.objects.create(
            ticket=self.ticket_unassigned,
            subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Dig into logs', assigned_to=self.forensic,
        )
        visible = Ticket.objects.visible_to(self.forensic)
        self.assertNotIn(self.ticket_unassigned, visible)
        self.assertEqual(list(visible), [self.ticket_a])  # only the response one


# ──────────────────────────────────────────────────────────────────────────── #
# 2. Visibility view tests                                                     #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketVisibilityViewTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.soc_staff = _make_t1('v_soc_staff')
        cls.admin_a   = _make_user('v_admin_a', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin_b   = _make_user('v_admin_b', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.ticket_a = _make_ticket(assigned_admin=cls.admin_a)
        cls.ticket_b = _make_ticket(assigned_admin=cls.admin_b)

    def test_admin_a_can_view_own_ticket(self):
        self.client.login(username='v_admin_a', password='testpass123')
        url = reverse('ticket_detail', kwargs={'pk': self.ticket_a.pk})
        self.assertEqual(self.client.get(url).status_code, 200)

    def test_admin_a_gets_404_on_admin_b_ticket(self):
        self.client.login(username='v_admin_a', password='testpass123')
        url = reverse('ticket_detail', kwargs={'pk': self.ticket_b.pk})
        self.assertEqual(self.client.get(url).status_code, 404)

    def test_unauthenticated_user_redirected(self):
        url = reverse('ticket_detail', kwargs={'pk': self.ticket_a.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 302)
        self.assertIn('/login/', response['Location'])


# ──────────────────────────────────────────────────────────────────────────── #
# 3. Workflow transition tests                                                  #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketReportExportTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('report_t1', phone='02-574-8209')
        cls.admin = _make_user('report_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.other_admin = _make_user('report_other_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.ticket = _make_ticket(
            created_by=cls.t1,
            assigned_admin=cls.admin,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            incident_name='Suspicious SoftEther Signed File',
            incident_datetime=timezone.now(),
            reference_id='INC-2026-0001',
            log_source='Wazuh',
            severity='High',
            ncsa_severity=Ticket.NCSA_SEVERITY_SEVERE,
            issue_type='SIEM',
            detailed_issue='Malicious Logic',
            detailed_issue2='Malware EDR',
            device_name='SRV-SQL-01',
            ip_address='192.0.2.10',
            mac_address='AA:BB:CC:DD:EE:FF',
            asset_type='Server',
            operating_system='Windows Server 2019',
            asset_owner='IT Operations',
            spread_to_others=False,
            destination_ip='203.0.113.50',
            ioc_details='203.0.113.50\nsoftether.example',
            mitre_phase='Initial Access,Execution',
            action_required='Block IoC and inspect persistence.',
            action_precautions='Preserve memory and logs before reboot.',
            actions_taken_summary='SOC contacted the owner and blocked the IP.',
            next_steps_summary='Monitor endpoint telemetry for 24 hours.',
            remediation_summary='Unauthorized service removed.',
            containment_report='Host isolated and C2 destination blocked.',
        )
        TicketAttachment.objects.create(
            ticket=cls.ticket,
            file='ticket_attachments/report/evidence.log',
            original_name='evidence.log',
            uploaded_by=cls.t1,
        )

    def test_generate_ticket_report_renders_docx_and_updates_metadata(self):
        snapshot_updated_at = self.ticket.updated_at

        report = generate_ticket_report(self.ticket.pk, generated_by=self.t1)
        content = report.content
        text = _docx_text(content)

        self.assertEqual(report.filename, f'report_{self.ticket.ticket_id}_{REPORT_TEMPLATE_VERSION}.docx')
        self.assertIn('Suspicious SoftEther Signed File', text)
        self.assertIn('SOC contacted the owner and blocked the IP.', text)
        self.assertIn('Host isolated and C2 destination blocked.', text)
        self.assertIn('Initial Access, Execution', text)
        self.assertIn('evidence.log', text)
        self.assertNotIn('{{ticket_id}}', text)
        # Checkbox states reflect the ticket: INCIDENT / High / SEVERE / Server.
        self.assertIn('☑ Incident', text)
        self.assertIn('☐ Event', text)
        self.assertIn('☑ High', text)
        self.assertIn('☑ ร้ายแรง', text)   # NCSA severe
        self.assertIn('☑ Server', text)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.report_template_version, REPORT_TEMPLATE_VERSION)
        self.assertEqual(self.ticket.report_format, 'docx')
        self.assertEqual(self.ticket.report_generated_by, self.t1)
        self.assertEqual(self.ticket.report_ticket_updated_at, snapshot_updated_at)
        self.assertEqual(self.ticket.report_sha256, hashlib.sha256(content).hexdigest())
        self.assertIsNotNone(self.ticket.report_generated_at)

    def test_ticket_report_docx_endpoint_streams_authorized_download(self):
        self.client.force_login(self.t1)
        response = self.client.post(reverse('ticket_report_docx', args=[self.ticket.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn(f'report_{self.ticket.ticket_id}_{REPORT_TEMPLATE_VERSION}.docx', response['Content-Disposition'])
        content = b''.join(response.streaming_content)
        self.assertIn('Suspicious SoftEther Signed File', _docx_text(content))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.report_format, 'docx')
        self.assertEqual(self.ticket.report_generated_by, self.t1)
        self.assertEqual(self.ticket.report_sha256, hashlib.sha256(content).hexdigest())

    def test_ticket_report_export_rejects_get(self):
        self.client.force_login(self.t1)
        snapshot = self.ticket.report_generated_at
        for url_name in ('ticket_report_docx', 'ticket_report_pdf'):
            response = self.client.get(reverse(url_name, args=[self.ticket.pk]))
            self.assertEqual(response.status_code, 405)
        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.report_generated_at, snapshot)

    def test_ticket_report_export_failure_redirects_with_message(self):
        self.client.force_login(self.t1)
        with patch(
            'apps.incidents.views.generate_ticket_report',
            side_effect=ValueError('Unresolved report template placeholders: {{bogus}}'),
        ):
            response = self.client.post(
                reverse('ticket_report_docx', args=[self.ticket.pk]), follow=True,
            )
        self.assertRedirects(response, reverse('ticket_detail', args=[self.ticket.pk]))
        message_texts = [str(m) for m in response.context['messages']]
        self.assertTrue(any('ไม่สามารถสร้างรายงาน DOCX ได้' in m for m in message_texts))

    def test_ticket_report_preview_returns_read_only_html(self):
        self.client.force_login(self.t1)
        response = self.client.get(reverse('ticket_report_preview', args=[self.ticket.pk]))

        self.assertEqual(response.status_code, 200)
        # NT-form section titles.
        self.assertContains(response, 'ข้อมูลทั่วไป (General Information)')
        self.assertContains(response, 'รายละเอียดเหตุการณ์ (Incident Description)')
        self.assertContains(response, 'Scope ทรัพย์สินที่ได้รับผลกระทบ')
        self.assertContains(response, 'Indicators of Compromise หรือหลักฐานที่พบ')
        self.assertContains(response, 'สิ่งที่ต้องดำเนินการ (Containment)')
        self.assertContains(response, 'สรุปผลการดำเนินการแก้ไข')
        self.assertContains(response, 'หมวดหมู่ของภัยคุกคามทางไซเบอร์')
        # Ticket data filled in.
        self.assertContains(response, 'Suspicious SoftEther Signed File')
        self.assertContains(response, 'SOC contacted the owner and blocked the IP.')
        self.assertContains(response, 'Host isolated and C2 destination blocked.')
        self.assertContains(response, 'กลับไปแก้ไข Ticket')
        # Data-driven checkbox: INCIDENT is checked, Event is not (☑=&#9745;, ☐=&#9744;).
        self.assertContains(response, '&#9745;</span>&#160;Incident')
        self.assertContains(response, '&#9744;</span>&#160;Event')
        # NT logo embedded as a data URI.
        self.assertContains(response, 'data:image/png;base64,')

    def test_ticket_report_pdf_endpoint_streams_valid_pdf_and_updates_metadata(self):
        self.client.force_login(self.t1)
        response = self.client.post(reverse('ticket_report_pdf', args=[self.ticket.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(f'report_{self.ticket.ticket_id}_{REPORT_TEMPLATE_VERSION}.pdf', response['Content-Disposition'])

        content = b''.join(response.streaming_content)
        self.assertTrue(content.startswith(b'%PDF'))
        pdf = PdfReader(BytesIO(content))
        self.assertGreaterEqual(len(pdf.pages), 1)
        text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        normalized_text = ' '.join(text.split())
        self.assertIn('Suspicious SoftEther Signed File', normalized_text)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.report_template_version, REPORT_TEMPLATE_VERSION)
        self.assertEqual(self.ticket.report_format, 'pdf')
        self.assertEqual(self.ticket.report_generated_by, self.t1)
        self.assertEqual(self.ticket.report_sha256, hashlib.sha256(content).hexdigest())
        self.assertIsNotNone(self.ticket.report_generated_at)

    def test_pdf_export_embeds_bundled_thai_font(self):
        from reportlab.pdfbase import pdfmetrics
        from apps.incidents.reports import REPORT_FONT_NAME, _register_pdf_font

        _register_pdf_font()
        registered = pdfmetrics.getRegisteredFontNames()
        self.assertIn(REPORT_FONT_NAME, registered)
        self.assertIn(f'{REPORT_FONT_NAME}-Bold', registered)

        self.client.force_login(self.t1)
        response = self.client.post(reverse('ticket_report_pdf', args=[self.ticket.pk]))
        content = b''.join(response.streaming_content)
        # The bundled TH Sarabun New faces must be embedded so Thai headings and
        # values render as real glyphs instead of blank boxes on any host.
        self.assertIn(b'THSarabunNew', content)
        self.assertIn(b'THSarabunNew-Bold', content)

    def test_ticket_report_docx_endpoint_respects_ticket_visibility(self):
        self.client.force_login(self.other_admin)
        response = self.client.post(reverse('ticket_report_docx', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 404)

    @staticmethod
    def _docx_placeholders(path):
        doc = Document(str(path))
        found = set()
        for paragraph in _iter_paragraphs(doc):
            found.update(re.findall(r'\{\{([^}]+)\}\}', paragraph.text))
        return found

    def test_template_placeholders_match_context_keys(self):
        # Both directions: an orphan placeholder in the .docx would raise at
        # export time, and an unused context key is silent drift.
        context_keys = set(build_ticket_report_context(self.ticket))
        self.assertEqual(self._docx_placeholders(REPORT_TEMPLATE_PATH), context_keys)

    def test_build_script_matches_committed_template(self):
        script_path = Path(settings.BASE_DIR) / 'scripts' / 'build_report_template_v2.py'
        spec = importlib.util.spec_from_file_location('build_report_template_v2', script_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)

        with tempfile.TemporaryDirectory() as tmp:
            rebuilt_path = Path(tmp) / 'rebuilt_template.docx'
            module.build(rebuilt_path)
            rebuilt = self._docx_placeholders(rebuilt_path)

        self.assertEqual(rebuilt, self._docx_placeholders(REPORT_TEMPLATE_PATH))

    def _checkbox_options(self, ticket, label):
        """Return the {label: checked} options for a named checkbox row.

        Matched on containment, not equality: section 1 labels carry a "1.N "
        prefix to mirror the paper form, and these tests care about the row's
        identity rather than its position in the numbering.
        """
        report = build_ticket_report_context(ticket)
        for section in build_ticket_report_sections(report, ticket):
            for row in section['rows']:
                if row.get('type') == 'checks' and label in row['label']:
                    return {opt['label']: opt['checked'] for opt in row['options']}
        raise AssertionError(f'checkbox row {label!r} not found')

    def test_report_sections_reflect_ticket_checkbox_state(self):
        opts = self._checkbox_options(self.ticket, 'ประเภท: event หรือ incident')
        self.assertTrue(opts['Incident'])
        self.assertFalse(opts['Event'])

        sev = self._checkbox_options(self.ticket, 'ระดับความรุนแรง (อ้างอิงตามระบบ SIEM)')
        self.assertEqual([k for k, v in sev.items() if v], ['High'])

        ncsa = self._checkbox_options(self.ticket, 'ระดับความรุนแรง (อ้างอิงตาม สกมช.)')
        self.assertEqual([k for k, v in ncsa.items() if v], ['ร้ายแรง'])

        # is_emergency drives ระดับความสำคัญ (สำคัญ vs สำคัญมาก).
        imp = self._checkbox_options(self.ticket, 'ระดับความสำคัญ')
        self.assertEqual([k for k, v in imp.items() if v], ['สำคัญ'])

    def test_emergency_flag_flips_importance_checkbox(self):
        self.ticket.is_emergency = True
        imp = self._checkbox_options(self.ticket, 'ระดับความสำคัญ')
        self.assertEqual([k for k, v in imp.items() if v], ['สำคัญมาก'])

    def test_pdf_repeats_footer_on_every_page(self):
        self.client.force_login(self.t1)
        response = self.client.post(reverse('ticket_report_pdf', args=[self.ticket.pk]))
        pdf = PdfReader(BytesIO(b''.join(response.streaming_content)))
        pages_with_footer = sum(
            1 for page in pdf.pages
            if 'INCIDENT REPORT CONTAINMENT' in (page.extract_text() or '')
        )
        self.assertEqual(pages_with_footer, len(pdf.pages))

    def test_docx_renders_multiline_values_as_line_breaks(self):
        report = generate_ticket_report(self.ticket.pk, generated_by=self.t1)
        doc = Document(BytesIO(report.content))

        target = next(
            (p for p in _iter_paragraphs(doc)
             if '203.0.113.50' in p.text and 'softether.example' in p.text),
            None,
        )
        self.assertIsNotNone(target, 'ioc_details paragraph not found in DOCX')
        # python-docx turns \n into <w:br/> (older versions used <w:cr/>);
        # either way the two IoC lines must not collapse into one.
        breaks = target._p.findall(f'.//{{{target._p.nsmap["w"]}}}br')
        carriage = target._p.findall(f'.//{{{target._p.nsmap["w"]}}}cr')
        self.assertTrue(breaks or carriage, 'multiline value lost its line break')

    def test_ticket_detail_shows_stale_report_badge(self):
        stale_marker = 'หลังสร้างรายงานล่าสุด'
        self.client.force_login(self.t1)

        # No report generated yet — no badge.
        response = self.client.get(reverse('ticket_detail', args=[self.ticket.pk]))
        self.assertNotContains(response, stale_marker)

        # Fresh report — still no badge.
        self.client.post(reverse('ticket_report_docx', args=[self.ticket.pk]))
        response = self.client.get(reverse('ticket_detail', args=[self.ticket.pk]))
        self.assertNotContains(response, stale_marker)

        # Ticket modified after the export — badge appears.
        Ticket.objects.filter(pk=self.ticket.pk).update(
            updated_at=timezone.now() + timedelta(seconds=5),
        )
        response = self.client.get(reverse('ticket_detail', args=[self.ticket.pk]))
        self.assertContains(response, stale_marker)
        self.assertContains(response, 'ล่าสุด (DOCX')


class WorkflowTransitionTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('wf_t1')
        cls.t2    = _make_t2('wf_t2')
        cls.mgr   = _make_user('wf_mgr',   UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('wf_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _incident(self, severity='High'):
        return _make_ticket(
            assigned_admin=self.admin, created_by=self.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT, severity=severity,
        )

    # ── Happy path ──────────────────────────────────────────────────────── #

    def test_new_incident_to_pending_mgr_triage(self):
        t = self._incident()
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route to manager')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_TRIAGE)

    def test_manager_forwards_to_admin_lane(self):
        t = self._incident()
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route')
        t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.mgr, 'forward')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    def test_manager_forwards_to_owner_lane(self):
        t = self._incident()
        t.t1_route = Ticket.T1_ROUTE_OWNER
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route')
        t.transition_to(Ticket.STATUS_AWAITING_OWNER, self.mgr, 'forward')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_OWNER)

    def test_manager_cannot_forward_to_wrong_lane(self):
        """t1_route=ADMIN forbids forwarding to the owner lane, and vice versa."""
        t = self._incident()
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route')
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_OWNER, self.mgr, 'wrong lane')

    def test_new_incident_to_escalated_t2_stamps_escalation(self):
        t = self._incident()
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_ESCALATED_T2)
        self.assertIsNotNone(t.escalated_to_t2_at)
        self.assertTrue(t.was_escalated_to_t2)

    def test_new_event_escalates_then_t2_closes_without_manager(self):
        """Tier 1 can no longer close an Event directly — Tier 2 confirms it."""
        t = _make_ticket(created_by=self.t1, classification=Ticket.CLASSIFICATION_EVENT)
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'event → T2 confirm')
        self.assertEqual(t.status, Ticket.STATUS_ESCALATED_T2)
        t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'confirmed event')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)

    def test_t1_cannot_close_event_directly(self):
        t = _make_ticket(created_by=self.t1, classification=Ticket.CLASSIFICATION_EVENT)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t1, 'benign')

    def test_transition_stamps_status_changed_at(self):
        t = self._incident()
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        before = t.status_changed_at
        self.assertIsNotNone(before)  # seeded on creation
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route')
        t.refresh_from_db()
        self.assertGreater(t.status_changed_at, before)

    def test_same_status_note_does_not_bump_status_changed_at(self):
        t = self._incident()
        stamp = t.status_changed_at
        # Same-status, note-only update — not a lifecycle move.
        t.transition_to(Ticket.STATUS_NEW, self.t1, 'just a note')
        t.refresh_from_db()
        self.assertEqual(t.status_changed_at, stamp)

    def test_escalated_incident_returns_to_t1_review(self):
        t = self._incident()
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        t.transition_to(Ticket.STATUS_T1_REVIEW, self.t2, 'confirm incident')
        self.assertEqual(t.status, Ticket.STATUS_T1_REVIEW)

    def test_t1_review_to_pending_mgr_triage(self):
        t = self._incident()
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        t.transition_to(Ticket.STATUS_T1_REVIEW, self.t2, 'confirm')
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route to manager')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_TRIAGE)

    def test_full_happy_path_t2_closes_without_manager(self):
        t = self._incident(severity='High')
        _advance_to(t, Ticket.STATUS_APPROVED, self.t1, self.admin, t2=self.t2)
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_full_happy_path_emergency_via_manager(self):
        t = self._incident(severity='Critical')
        t.is_emergency = True
        t.save(update_fields=['is_emergency'])
        _advance_to(t, Ticket.STATUS_APPROVED, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_containment_rejection_loop(self):
        t = self._incident()
        _advance_to(t, Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, self.admin)
        # Tier 2 (not Tier 1) judges the containment report and sends it back.
        t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t2, 'not contained')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    # ── Illegal transitions ─────────────────────────────────────────────── #

    def test_cannot_skip_states(self):
        t = self._incident()
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, 'skip')

    def test_approved_is_terminal(self):
        t = self._incident()
        _advance_to(t, Ticket.STATUS_APPROVED, self.t1, self.admin, t2=self.t2)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t1, 'reopen')

    def test_closed_event_is_terminal(self):
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_EVENT,
            status=Ticket.STATUS_CLOSED_EVENT,
        )
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t1, 'reopen')

    def test_event_cannot_take_incident_path(self):
        """A ticket classified EVENT cannot be routed into the manager review."""
        t = _make_ticket(
            created_by=self.t1, assigned_admin=self.admin,
            classification=Ticket.CLASSIFICATION_EVENT, t1_route=Ticket.T1_ROUTE_ADMIN,
        )
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'mismatch')

    def test_incident_cannot_be_closed_as_event_before_reclassify(self):
        """The mid-containment Event-close edge requires classification == EVENT."""
        t = self._incident()
        _advance_to(t, Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, self.admin)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'still incident')

    def test_t2_reclassifies_containment_as_event_closes_without_manager(self):
        """Even an emergency ticket closes directly once T2 reclassifies to Event."""
        t = self._incident()
        t.is_emergency = True
        t.save(update_fields=['is_emergency'])
        _advance_to(t, Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, self.admin)
        # T2 decides it was benign after all: flip classification, then close.
        t.classification = Ticket.CLASSIFICATION_EVENT
        t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'reclassified as event')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)

    def test_invalid_status_code_raises(self):
        t = self._incident()
        with self.assertRaises(ValidationError):
            t.transition_to('BOGUS', self.t1, 'bad')


# ──────────────────────────────────────────────────────────────────────────── #
# 4. Permission matrix tests                                                   #
# ──────────────────────────────────────────────────────────────────────────── #

class WorkflowPermissionTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1       = _make_t1('pm_t1')
        cls.other_t1 = _make_t1('pm_t1_other')
        cls.t2       = _make_t2('pm_t2')
        cls.mgr      = _make_user('pm_mgr',     UserProfile.ROLE_SOC_MANAGER)
        cls.admin_a  = _make_user('pm_admin_a', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin_b  = _make_user('pm_admin_b', UserProfile.ROLE_SYSTEM_ADMIN)

    def _ticket_at(self, status, severity='High',
                   classification=Ticket.CLASSIFICATION_INCIDENT, **kwargs):
        opts = dict(
            status=status, severity=severity, classification=classification,
            assigned_admin=self.admin_a, created_by=self.t1,
        )
        opts.update(kwargs)
        return _make_ticket(**opts)

    # NEW → PENDING_MGR_TRIAGE  requires TIER1_CREATOR ─────────────────────

    def test_creator_t1_can_dispatch(self):
        t = self._ticket_at(Ticket.STATUS_NEW, t1_route=Ticket.T1_ROUTE_ADMIN)
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'ok')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_TRIAGE)

    def test_non_creator_t1_cannot_dispatch(self):
        t = self._ticket_at(Ticket.STATUS_NEW, t1_route=Ticket.T1_ROUTE_ADMIN)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.other_t1, 'denied')

    def test_t2_cannot_dispatch(self):
        t = self._ticket_at(Ticket.STATUS_NEW, t1_route=Ticket.T1_ROUTE_ADMIN)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t2, 'denied')

    def test_manager_cannot_dispatch(self):
        """Managers are not Tier 1 and never open/route a fresh ticket."""
        t = self._ticket_at(Ticket.STATUS_NEW, t1_route=Ticket.T1_ROUTE_ADMIN)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.mgr, 'denied')

    # PENDING_MGR_TRIAGE → lane  requires MANAGER ──────────────────────────

    def test_manager_can_forward(self):
        t = self._ticket_at(Ticket.STATUS_PENDING_MGR_TRIAGE, t1_route=Ticket.T1_ROUTE_ADMIN)
        t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.mgr, 'forward')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    def test_t1_cannot_forward_from_mgr_triage(self):
        t = self._ticket_at(Ticket.STATUS_PENDING_MGR_TRIAGE, t1_route=Ticket.T1_ROUTE_ADMIN)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t1, 'denied')

    def test_t2_cannot_forward_from_mgr_triage(self):
        t = self._ticket_at(Ticket.STATUS_PENDING_MGR_TRIAGE, t1_route=Ticket.T1_ROUTE_ADMIN)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t2, 'denied')

    # ESCALATED_T2 → T1_REVIEW  requires TIER2 ─────────────────────────────

    def test_t2_can_return_to_t1(self):
        t = self._ticket_at(Ticket.STATUS_ESCALATED_T2)
        t.transition_to(Ticket.STATUS_T1_REVIEW, self.t2, 'ok')
        self.assertEqual(t.status, Ticket.STATUS_T1_REVIEW)

    def test_t1_cannot_return_to_t1(self):
        t = self._ticket_at(Ticket.STATUS_ESCALATED_T2)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_T1_REVIEW, self.t1, 'denied')

    # AWAITING_CONTAINMENT → CONTAINMENT_REPORTED  requires ASSIGNED_ADMIN ─

    def test_assigned_admin_can_report(self):
        t = self._ticket_at(Ticket.STATUS_AWAITING_CONTAINMENT, assigned_admin=self.admin_a)
        t.containment_report = 'contained'
        t.transition_to(Ticket.STATUS_CONTAINMENT_REPORTED, self.admin_a, 'ok')
        self.assertEqual(t.status, Ticket.STATUS_CONTAINMENT_REPORTED)

    def test_other_admin_cannot_report(self):
        t = self._ticket_at(Ticket.STATUS_AWAITING_CONTAINMENT, assigned_admin=self.admin_a)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CONTAINMENT_REPORTED, self.admin_b, 'denied')

    def test_t1_cannot_report_containment(self):
        t = self._ticket_at(Ticket.STATUS_AWAITING_CONTAINMENT)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, 'denied')

    # CONTAINMENT_REPORTED close  requires TIER2 ──────────────────────────

    def test_t2_can_verify_and_close_when_no_manager(self):
        t = self._ticket_at(Ticket.STATUS_CONTAINMENT_REPORTED, severity='High')
        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'verified — close')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_creator_t1_cannot_close_containment(self):
        """Containment verification moved to Tier 2 — even the creator may not close."""
        t = self._ticket_at(Ticket.STATUS_CONTAINMENT_REPORTED, severity='High')
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t1, 'denied')

    def test_t2_must_route_emergency_to_manager(self):
        t = self._ticket_at(
            Ticket.STATUS_CONTAINMENT_REPORTED, severity='High', is_emergency=True,
        )
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'denied — emergency')
        t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'to manager')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MANAGER)

    def test_t2_can_reject_containment_back_to_admin(self):
        t = self._ticket_at(Ticket.STATUS_CONTAINMENT_REPORTED, severity='High')
        t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t2, 'not contained')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    def test_t1_cannot_reject_containment(self):
        t = self._ticket_at(Ticket.STATUS_CONTAINMENT_REPORTED, severity='High')
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t1, 'denied')

    # PENDING_MANAGER → APPROVED  requires MANAGER ────────────────────────

    def test_manager_can_approve(self):
        t = self._ticket_at(Ticket.STATUS_PENDING_MANAGER, severity='Critical')
        t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'approved')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_t1_cannot_approve_pending_manager(self):
        t = self._ticket_at(Ticket.STATUS_PENDING_MANAGER, severity='Critical')
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t1, 'denied')


# ──────────────────────────────────────────────────────────────────────────── #
# 5. Tier 1 Event/Incident disposition at creation                             #
# ──────────────────────────────────────────────────────────────────────────── #

class T1ClassificationCreateTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('cc_t1')
        cls.t2    = _make_t2('cc_t2')
        cls.admin = _make_user('cc_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def test_event_creation_escalates_to_t2(self):
        """Tier 1 Event no longer closes directly — it escalates to Tier 2."""
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification=Ticket.CLASSIFICATION_EVENT, t1_route='',
        ))
        self.assertEqual(resp.status_code, 302)
        ticket = Ticket.objects.latest('id')
        self.assertEqual(ticket.classification, Ticket.CLASSIFICATION_EVENT)
        self.assertEqual(ticket.status, Ticket.STATUS_ESCALATED_T2)

    def test_incident_assign_admin_routes_to_mgr_triage(self):
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ASSIGN_ADMIN,
            assigned_admin=self.admin.pk,
        ))
        self.assertEqual(resp.status_code, 302)
        ticket = Ticket.objects.latest('id')
        self.assertEqual(ticket.status, Ticket.STATUS_PENDING_MGR_TRIAGE)
        self.assertEqual(ticket.t1_route, Ticket.T1_ROUTE_ADMIN)
        self.assertEqual(ticket.assigned_admin, self.admin)

    def test_incident_direct_owner_routes_to_mgr_triage(self):
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_DIRECT_OWNER,
        ))
        self.assertEqual(resp.status_code, 302)
        ticket = Ticket.objects.latest('id')
        self.assertEqual(ticket.status, Ticket.STATUS_PENDING_MGR_TRIAGE)
        self.assertEqual(ticket.t1_route, Ticket.T1_ROUTE_OWNER)

    def test_incident_escalate_routes_to_t2(self):
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ESCALATE_T2,
        ))
        self.assertEqual(resp.status_code, 302)
        ticket = Ticket.objects.latest('id')
        self.assertEqual(ticket.status, Ticket.STATUS_ESCALATED_T2)
        self.assertIsNotNone(ticket.escalated_to_t2_at)

    def test_incident_assign_admin_without_admin_is_invalid(self):
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ASSIGN_ADMIN,
        ))
        self.assertEqual(resp.status_code, 200)  # form re-rendered with errors
        self.assertFalse(Ticket.objects.exists())

    def test_missing_classification_is_invalid(self):
        self.client.login(username='cc_t1', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            classification='', t1_route='',
        ))
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(Ticket.objects.exists())

    def test_t2_cannot_open_create_ticket_page(self):
        self.client.login(username='cc_t2', password='testpass123')
        resp = self.client.get(reverse('create_ticket'))
        self.assertEqual(resp.status_code, 302)  # redirected — Tier 1 only

    def test_admin_cannot_open_create_ticket_page(self):
        self.client.login(username='cc_admin', password='testpass123')
        resp = self.client.get(reverse('create_ticket'))
        self.assertEqual(resp.status_code, 302)


# ──────────────────────────────────────────────────────────────────────────── #
# 6. Tier 2 return-only constraint                                             #
# ──────────────────────────────────────────────────────────────────────────── #

class Tier2EscalationTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('t2t_t1')
        cls.t2    = _make_t2('t2t_t2')
        cls.admin = _make_user('t2t_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _escalated(self):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_ESCALATED_T2, assigned_admin=self.admin,
            escalated_to_t2_at=timezone.now(),
        )

    def test_t2_confirms_incident_returns_to_t1(self):
        t = self._escalated()
        t.transition_to(Ticket.STATUS_T1_REVIEW, self.t2, 'confirmed incident')
        self.assertEqual(t.status, Ticket.STATUS_T1_REVIEW)

    def test_t2_reclassifies_event_and_closes(self):
        t = self._escalated()
        t.classification = Ticket.CLASSIFICATION_EVENT  # T2 may revise classification
        t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'benign on review')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)

    def test_t2_event_downgrade_needs_manager_verification(self):
        """Counter-measure: Tier 2 cannot dispose of an escalated Incident by
        relabelling it an Event — the SOC Manager verifies that call first."""
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_NEW,
        )
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        self.assertEqual(
            t.classification_at_escalation, Ticket.CLASSIFICATION_INCIDENT)

        t.classification = Ticket.CLASSIFICATION_EVENT
        self.assertTrue(t.is_t2_event_downgrade)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'benign')

        t.transition_to(Ticket.STATUS_PENDING_MGR_EVENT_REVIEW, self.t2, 'benign')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_EVENT_REVIEW)

    def test_manager_confirming_the_downgrade_closes_the_event(self):
        manager = _make_user('t2t_mgr_ok', UserProfile.ROLE_SOC_MANAGER)
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_NEW,
        )
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        t.classification = Ticket.CLASSIFICATION_EVENT
        t.transition_to(Ticket.STATUS_PENDING_MGR_EVENT_REVIEW, self.t2, 'benign')

        t.transition_to(Ticket.STATUS_CLOSED_EVENT, manager, 'agreed, benign')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)

    def test_manager_rejecting_the_downgrade_returns_it_as_an_incident(self):
        manager = _make_user('t2t_mgr_no', UserProfile.ROLE_SOC_MANAGER)
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_NEW,
        )
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'escalate')
        t.classification = Ticket.CLASSIFICATION_EVENT
        t.transition_to(Ticket.STATUS_PENDING_MGR_EVENT_REVIEW, self.t2, 'benign')

        t.transition_to(Ticket.STATUS_ESCALATED_T2, manager, 'this is real')
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_ESCALATED_T2)
        # Flipped back, so Tier 2 must handle it rather than re-propose a close.
        self.assertEqual(t.classification, Ticket.CLASSIFICATION_INCIDENT)
        self.assertEqual(
            t.classification_at_escalation, Ticket.CLASSIFICATION_INCIDENT)
        self.assertFalse(t.is_t2_event_downgrade)

    def test_event_classified_by_tier1_still_closes_directly(self):
        """Not a downgrade — Tier 2 is confirming what Tier 1 already called."""
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_EVENT,
            status=Ticket.STATUS_NEW,
        )
        t.transition_to(Ticket.STATUS_ESCALATED_T2, self.t1, 'event → T2 confirm')
        self.assertFalse(t.is_t2_event_downgrade)

        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MGR_EVENT_REVIEW, self.t2, 'x')

        t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'confirmed benign')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)

    def test_t2_cannot_assign_to_admin(self):
        """No ESCALATED_T2 → AWAITING_CONTAINMENT edge exists at all."""
        t = self._escalated()
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.t2, 'forbidden')

    def test_t2_cannot_create_ticket(self):
        self.client.login(username='t2t_t2', password='testpass123')
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data())
        self.assertEqual(resp.status_code, 302)  # redirected away, Tier 1 only
        self.assertFalse(Ticket.objects.filter(device_name='TEST-ENDPOINT-01').exists())

    def test_t1_review_then_route_to_mgr_triage(self):
        t = self._escalated()
        t.transition_to(Ticket.STATUS_T1_REVIEW, self.t2, 'confirm')
        t.t1_route = Ticket.T1_ROUTE_ADMIN
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'assign admin')
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_TRIAGE)


# ──────────────────────────────────────────────────────────────────────────── #
# 7. Manager routing (requires_manager_verification)                           #
# ──────────────────────────────────────────────────────────────────────────── #

class ManagerRoutingTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('mr_t1')
        cls.t2    = _make_t2('mr_t2')
        cls.mgr   = _make_user('mr_mgr',   UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('mr_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _contained(self, severity='High', is_emergency=False):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            assigned_admin=self.admin, status=Ticket.STATUS_CONTAINMENT_REPORTED,
            severity=severity, is_emergency=is_emergency, containment_report='done',
        )

    def test_critical_does_not_require_manager(self):
        """Severity alone never routes to the manager — only the emergency flag."""
        t = self._contained(severity='Critical')
        self.assertFalse(t.requires_manager_verification)

    def test_high_does_not_require_manager(self):
        t = self._contained(severity='High')
        self.assertFalse(t.requires_manager_verification)

    def test_emergency_forces_manager_even_on_high(self):
        t = self._contained(severity='High', is_emergency=True)
        self.assertTrue(t.requires_manager_verification)

    def test_non_emergency_ticket_t2_closes_directly(self):
        t = self._contained(severity='High')
        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'verified — closed')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_critical_non_emergency_t2_closes_directly(self):
        t = self._contained(severity='Critical')
        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'verified — closed')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_non_emergency_ticket_cannot_route_to_manager(self):
        t = self._contained(severity='High')
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'no need')

    def test_emergency_ticket_t2_cannot_close_directly(self):
        t = self._contained(severity='High', is_emergency=True)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'must go to manager')

    def test_emergency_ticket_routes_to_manager_then_closes(self):
        t = self._contained(severity='High', is_emergency=True)
        t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'T2 verified')
        t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'approved')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)


# ──────────────────────────────────────────────────────────────────────────── #
# 7b. Direct-to-Owner fast path (Low/Medium)                                    #
# ──────────────────────────────────────────────────────────────────────────── #

def _owner_payload(severity='Low', **overrides):
    """A valid create_ticket POST payload for the direct-to-owner route.

    Extends _ticket_post_data with the direct-owner route and a lower default
    statutory severity suitable for the owner-remediation fast path.
    """
    data = _ticket_post_data(
        classification=Ticket.CLASSIFICATION_INCIDENT,
        t1_route=TicketForm.ROUTE_DIRECT_OWNER,
        severity=severity,
        ncsa_severity=Ticket.NCSA_SEVERITY_NON_SEVERE,
        log_source='Windows Security Event Log',
    )
    data.update(overrides)
    return data


class DirectToOwnerPathTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('do_t1')
        cls.other = _make_t1('do_other')   # a different Tier 1 (not the creator)
        cls.t2    = _make_t2('do_t2')
        cls.mgr   = _make_user('do_mgr', UserProfile.ROLE_SOC_MANAGER)

    def _owner_case(self, severity='Low', is_emergency=False,
                    status=Ticket.STATUS_AWAITING_OWNER):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            severity=severity, is_emergency=is_emergency, status=status,
        )

    # ── Model FSM: happy path (Low, non-emergency → Tier 2 review) ──────── #
    def test_full_owner_path_low_severity_closes_via_tier2(self):
        t = self._owner_case(status=Ticket.STATUS_NEW)
        # Owner lane now passes the SOC Manager pre-containment review first.
        t.t1_route = Ticket.T1_ROUTE_OWNER
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route to manager')
        t.transition_to(Ticket.STATUS_AWAITING_OWNER, self.mgr, 'manager forwards to owner')
        self.assertTrue(t.direct_owner_remediation)
        self.assertIsNotNone(t.owner_contacted_at)

        t.transition_to(Ticket.STATUS_OWNER_REMEDIATED, self.t1, 'owner fixed')
        t.transition_to(Ticket.STATUS_PENDING_T2_REVIEW, self.t1, 'to review')
        self.assertIsNone(t.verified_by)           # verification is T2's act now

        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'reviewed & closed')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)
        self.assertEqual(t.verified_by, self.t2)   # T2 sign-off stamped at close
        self.assertEqual(t.approved_by, self.t2)
        self.assertIsNotNone(t.closed_at)

    # ── Review split: non-emergency → Tier 2 only (never the manager) ───── #
    def test_non_emergency_routes_to_tier2_not_manager(self):
        t = self._owner_case(status=Ticket.STATUS_OWNER_REMEDIATED)
        self.assertTrue(t.can_transition_to(Ticket.STATUS_PENDING_T2_REVIEW))
        self.assertFalse(t.can_transition_to(Ticket.STATUS_PENDING_MANAGER))
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t1, 'no')

    # ── Review split: emergency passes Tier 2 first, then the Manager ───── #
    def test_emergency_owner_path_passes_t2_then_manager(self):
        t = self._owner_case(is_emergency=True,
                             status=Ticket.STATUS_OWNER_REMEDIATED)
        # Every owner case goes to Tier 2 review — including emergencies.
        t.transition_to(Ticket.STATUS_PENDING_T2_REVIEW, self.t1, 'to review')
        # Tier 2 may not close an emergency directly; it must go to the manager.
        self.assertFalse(t.can_transition_to(Ticket.STATUS_APPROVED))
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'no')
        t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'to manager')
        t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'approved')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    # ── Tier 2 reject loops back to the owner ──────────────────────────── #
    def test_tier2_can_reject_back_to_owner(self):
        t = self._owner_case(status=Ticket.STATUS_PENDING_T2_REVIEW)
        t.transition_to(Ticket.STATUS_AWAITING_OWNER, self.t2, 'not actually fixed')
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_OWNER)

    # ── Permissions: T1 side is creator-gated; review close is Tier 2 ───── #
    def test_non_creator_t1_cannot_confirm(self):
        t = self._owner_case(status=Ticket.STATUS_AWAITING_OWNER)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_OWNER_REMEDIATED, self.other, 'nope')

    def test_tier2_review_close_requires_tier2(self):
        t = self._owner_case(status=Ticket.STATUS_PENDING_T2_REVIEW)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t1, 'T1 cannot close a T2 review')

    # ── Critical severity alone no longer routes to the SOC Manager ─────── #
    def test_critical_severity_still_closes_via_tier2(self):
        t = self._owner_case(severity='Critical',
                             status=Ticket.STATUS_OWNER_REMEDIATED)
        self.assertTrue(t.can_transition_to(Ticket.STATUS_PENDING_T2_REVIEW))
        t.transition_to(Ticket.STATUS_PENDING_T2_REVIEW, self.t1, 'to review')
        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'verified — closed')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    # ── Form gating: route valid at any severity ───────────────────────── #
    def test_form_accepts_direct_owner_for_low_severity(self):
        form = TicketForm(data=_owner_payload(severity='Low'), user=self.t1)
        self.assertTrue(form.is_valid(), form.errors)

    def test_form_accepts_direct_owner_for_high_severity(self):
        form = TicketForm(data=_owner_payload(severity='High'), user=self.t1)
        self.assertTrue(form.is_valid(), form.errors)

    # ── Create-flow view: routes to the manager review, sends no admin email ─ #
    def test_create_view_routes_to_mgr_triage_without_email(self):
        self.client.login(username='do_t1', password='testpass123')
        mail.outbox = []
        resp = self.client.post(reverse('create_ticket'), _owner_payload(severity='Low'))
        self.assertEqual(resp.status_code, 302)
        t = Ticket.objects.latest('id')
        # The owner lane is remembered; the ticket waits for the SOC Manager.
        self.assertEqual(t.status, Ticket.STATUS_PENDING_MGR_TRIAGE)
        self.assertEqual(t.t1_route, Ticket.T1_ROUTE_OWNER)
        self.assertIsNone(t.assigned_admin)
        # No admin email (owner lane); test managers have no email either.
        self.assertEqual(len(mail.outbox), 0)


# ──────────────────────────────────────────────────────────────────────────── #
# 8. Emergency flag permissions + audit                                        #
# ──────────────────────────────────────────────────────────────────────────── #

class EmergencyFlagTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('em_t1')
        cls.t2    = _make_t2('em_t2')
        cls.mgr   = _make_user('em_mgr',   UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('em_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.owner = _make_user('em_owner', UserProfile.ROLE_SYSTEM_OWNER)
        cls.superuser = User.objects.create_superuser('em_super', 'em@x.com', 'testpass123')

    def _direct_admin_ticket(self):
        """Incident handed directly to admin — never escalated to T2."""
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            assigned_admin=self.admin, status=Ticket.STATUS_AWAITING_CONTAINMENT,
        )

    def _escalated_ticket(self):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_ESCALATED_T2, escalated_to_t2_at=timezone.now(),
        )

    # ── SOC Manager only — no other role may reassess the flag ──────────── #

    def test_t1_cannot_reassess_emergency(self):
        t = self._direct_admin_ticket()
        self.assertFalse(t.can_reassess_emergency(self.t1))
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.t1, 'reason')

    def test_t1_cannot_reassess_emergency_even_on_escalated_ticket(self):
        """The old escalation exception is gone — the manager decides, full stop."""
        t = self._escalated_ticket()
        self.assertFalse(t.can_reassess_emergency(self.t1))
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.t1, 'reason')

    def test_t2_cannot_reassess_emergency(self):
        t = self._direct_admin_ticket()
        self.assertFalse(t.can_reassess_emergency(self.t2))
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.t2, 'reason')

    def test_system_admin_cannot_reassess_emergency(self):
        t = self._direct_admin_ticket()
        self.assertFalse(t.can_reassess_emergency(self.admin))
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.admin, 'reason')

    def test_owner_cannot_reassess_emergency(self):
        self.assertFalse(self._direct_admin_ticket().can_reassess_emergency(self.owner))

    def test_manager_can_reassess_emergency(self):
        t = self._direct_admin_ticket()
        self.assertTrue(t.can_reassess_emergency(self.mgr))
        t.reassess_emergency(True, self.mgr, 'situation escalated')
        t.refresh_from_db()
        self.assertTrue(t.is_emergency)

    def test_superuser_can_reassess_emergency(self):
        t = self._direct_admin_ticket()
        self.assertTrue(t.can_reassess_emergency(self.superuser))

    # ── Initial assessment (pre-containment review) ─────────────────────── #

    def test_initial_assessment_stamps_metadata_even_for_normal(self):
        t = self._escalated_ticket()
        t.assess_emergency_initial(False, self.mgr)
        t.save()
        t.refresh_from_db()
        self.assertFalse(t.is_emergency)
        self.assertEqual(t.emergency_decided_by, self.mgr)
        self.assertIsNotNone(t.emergency_decided_at)

    def test_initial_assessment_emergency_sets_flag_and_metadata(self):
        t = self._escalated_ticket()
        t.assess_emergency_initial(True, self.mgr)
        t.save()
        t.refresh_from_db()
        self.assertTrue(t.is_emergency)
        self.assertEqual(t.emergency_decided_by, self.mgr)

    def test_initial_assessment_metadata_is_write_once(self):
        t = self._escalated_ticket()
        t.assess_emergency_initial(False, self.mgr)
        t.save()
        first_at = t.emergency_decided_at
        # A later reassessment must not overwrite the initial-decision metadata.
        t.status = Ticket.STATUS_AWAITING_CONTAINMENT
        t.save()
        t.reassess_emergency(True, self.superuser, 'later change')
        t.refresh_from_db()
        self.assertEqual(t.emergency_decided_by, self.mgr)
        self.assertEqual(t.emergency_decided_at, first_at)

    def test_non_manager_cannot_make_initial_assessment(self):
        t = self._escalated_ticket()
        with self.assertRaises(ValidationError):
            t.assess_emergency_initial(True, self.t1)

    # ── Reassessment: terminal block + reason + audit ───────────────────── #

    def test_reassessment_forbidden_after_closure(self):
        for terminal in (Ticket.STATUS_APPROVED, Ticket.STATUS_CLOSED_EVENT):
            with self.subTest(status=terminal):
                t = _make_ticket(
                    created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
                    status=terminal,
                )
                self.assertFalse(t.can_reassess_emergency(self.mgr))
                with self.assertRaises(ValidationError):
                    t.reassess_emergency(True, self.mgr, 'too late')

    def test_reassessment_forbidden_at_pending_mgr_triage(self):
        t = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_PENDING_MGR_TRIAGE, t1_route=Ticket.T1_ROUTE_ADMIN,
        )
        self.assertFalse(t.can_reassess_emergency(self.mgr))
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.mgr, 'use the review form')

    def test_reassessment_requires_a_reason(self):
        t = self._direct_admin_ticket()
        with self.assertRaises(ValidationError):
            t.reassess_emergency(True, self.mgr, '   ')
        t.refresh_from_db()
        self.assertFalse(t.is_emergency)

    def test_reassessment_rejects_no_change(self):
        t = self._direct_admin_ticket()  # already Normal
        with self.assertRaises(ValidationError):
            t.reassess_emergency(False, self.mgr, 'no actual change')

    def test_reassessment_writes_audit_log_with_reason(self):
        t = self._direct_admin_ticket()
        before = t.logs.count()
        t.reassess_emergency(True, self.mgr, 'attacker moved laterally')
        self.assertEqual(t.logs.count(), before + 1)
        log = t.logs.first()
        self.assertEqual(log.author, self.mgr)
        self.assertIn('Emergency', log.note)
        self.assertIn('attacker moved laterally', log.note)


# ──────────────────────────────────────────────────────────────────────────── #
# 9. System Admin field access                                                 #
# ──────────────────────────────────────────────────────────────────────────── #

class AdminFieldAccessTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('af_t1')
        cls.admin = _make_user('af_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _awaiting(self):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            assigned_admin=self.admin, status=Ticket.STATUS_AWAITING_CONTAINMENT,
        )

    def test_admin_writes_containment_and_remediation(self):
        t = self._awaiting()
        self.client.login(username='af_admin', password='testpass123')
        resp = self.client.post(reverse('ticket_detail', args=[t.pk]), {
            'action': 'containment',
            'containment_report': 'Isolated the host and blocked the C2 IP.',
            'remediation_summary': 'Root cause: phishing. Reimaged endpoint.',
            'note': 'done',
        })
        self.assertRedirects(resp, reverse('ticket_detail', args=[t.pk]))
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_CONTAINMENT_REPORTED)
        self.assertIn('blocked the C2 IP', t.containment_report)
        self.assertIn('Reimaged endpoint', t.remediation_summary)

    def test_admin_containment_does_not_set_classification(self):
        t = self._awaiting()
        self.client.login(username='af_admin', password='testpass123')
        self.client.post(reverse('ticket_detail', args=[t.pk]), {
            'action': 'containment',
            'containment_report': 'Contained.',
            'note': 'done',
        })
        t.refresh_from_db()
        # classification stays whatever T1 set — admin never touches it.
        self.assertEqual(t.classification, Ticket.CLASSIFICATION_INCIDENT)


# ──────────────────────────────────────────────────────────────────────────── #
# 10. Sign-off field tests                                                     #
# ──────────────────────────────────────────────────────────────────────────── #

class SignOffFieldsTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('sf_t1')
        cls.t2    = _make_t2('sf_t2')
        cls.t2b   = _make_t2('sf_t2b')
        cls.mgr   = _make_user('sf_mgr',   UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('sf_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _incident(self, severity='Critical', is_emergency=True):
        # Manager routing now keys off the emergency flag, not severity.
        return _make_ticket(
            assigned_admin=self.admin, created_by=self.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT, severity=severity,
            is_emergency=is_emergency,
        )

    def test_verified_by_set_when_t2_marks_contained(self):
        t = self._incident()
        _advance_to(t, Ticket.STATUS_PENDING_MANAGER, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        t.refresh_from_db()
        self.assertEqual(t.verified_by, self.t2)
        self.assertIsNotNone(t.verified_at)

    def test_approved_by_set_to_manager(self):
        t = self._incident()
        _advance_to(t, Ticket.STATUS_APPROVED, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        t.refresh_from_db()
        self.assertEqual(t.approved_by, self.mgr)

    def test_direct_close_sets_both_signoffs_to_t2(self):
        t = self._incident(severity='High', is_emergency=False)  # no manager needed
        _advance_to(t, Ticket.STATUS_APPROVED, self.t1, self.admin, t2=self.t2)
        t.refresh_from_db()
        self.assertEqual(t.verified_by, self.t2)
        self.assertEqual(t.approved_by, self.t2)

    def test_verified_by_write_once(self):
        t = self._incident()
        _advance_to(t, Ticket.STATUS_PENDING_MANAGER, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        t.refresh_from_db()
        # Force back to CONTAINMENT_REPORTED; a different Tier 2 re-verifies —
        # the original verified_by must hold (write-once).
        Ticket.objects.filter(pk=t.pk).update(
            status=Ticket.STATUS_CONTAINMENT_REPORTED,
        )
        t.refresh_from_db()
        t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2b, 'again')
        t.refresh_from_db()
        self.assertEqual(t.verified_by, self.t2)


# ──────────────────────────────────────────────────────────────────────────── #
# 11. Email notification tests                                                  #
# ──────────────────────────────────────────────────────────────────────────── #

class NotificationEmailTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1  = _make_t1('ne_t1')
        cls.mgr = _make_user('ne_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('ne_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin.email = 'sysadmin@example.com'
        cls.admin.save()
        cls.admin_no_email = _make_user('ne_admin_noemail', UserProfile.ROLE_SYSTEM_ADMIN)

    def setUp(self):
        mail.outbox = []

    def _routed_ticket(self, admin=None):
        """A ticket driven to AWAITING_CONTAINMENT via the manager review."""
        t = _make_ticket(
            assigned_admin=admin or self.admin, created_by=self.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=Ticket.T1_ROUTE_ADMIN,
        )
        t.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.t1, 'route')
        t.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.mgr, 'forward')
        return t

    def test_routing_sends_one_email_to_admin(self):
        t = self._routed_ticket()
        self.assertTrue(notify_containment_required(t))
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn(self.admin.email, mail.outbox[0].to)

    def test_routing_subject_contains_ticket_id(self):
        t = self._routed_ticket()
        notify_containment_required(t)
        self.assertIn(t.ticket_id, mail.outbox[0].subject)

    def test_rejection_loop_body_contains_reason(self):
        t = self._routed_ticket()
        reason = 'Patch description is missing — include the CVE reference.'
        notify_containment_required(t, reason=reason)
        self.assertIn(reason, mail.outbox[0].body)

    def test_no_email_when_admin_has_no_email(self):
        t = _make_ticket(
            assigned_admin=self.admin_no_email, created_by=self.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT,
        )
        self.assertFalse(notify_containment_required(t))
        self.assertEqual(len(mail.outbox), 0)

    def test_transition_succeeds_without_email(self):
        t = self._routed_ticket(admin=self.admin_no_email)
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_AWAITING_CONTAINMENT)


# ──────────────────────────────────────────────────────────────────────────── #
# 12. Wazuh triage: 2-action + required release reason                          #
# ──────────────────────────────────────────────────────────────────────────── #

class WazuhTriageActionTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('wz_t1')
        cls.t2 = _make_t2('wz_t2')

    def _claimed_alert(self, claimer):
        return WazuhAlert.objects.create(
            opensearch_id=f'os-{claimer.username}-{timezone.now().timestamp()}',
            timestamp=timezone.now(), rule_level=12,
            rule_description='Suspicious activity',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=claimer, claimed_at=timezone.now(),
        )

    def test_create_ticket_action_redirects_to_create_form(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='wz_t1', password='testpass123')
        resp = self.client.post(reverse('triage_action'), {
            'alert_id': alert.pk, 'action': 'create_ticket',
            'note': 'looks real', 'category': WazuhAlert.CATEGORY_MALWARE,
        })
        self.assertEqual(resp.status_code, 302)
        self.assertIn(reverse('create_ticket'), resp.url)

    def test_close_fp_action_is_rejected(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='wz_t1', password='testpass123')
        resp = self.client.post(reverse('triage_action'), {
            'alert_id': alert.pk, 'action': 'close_fp', 'note': 'fp',
        })
        self.assertEqual(resp.status_code, 302)
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRIAGING)  # unchanged

    def test_escalate_action_is_rejected(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='wz_t1', password='testpass123')
        resp = self.client.post(reverse('triage_action'), {
            'alert_id': alert.pk, 'action': 'escalate',
            'note': 'unsure', 'escalate_to': WazuhAlert.TIER_T2,
        })
        self.assertEqual(resp.status_code, 302)
        alert.refresh_from_db()
        self.assertNotEqual(alert.triage_status, WazuhAlert.TRIAGE_ESCALATED)

    def test_release_requires_reason(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='wz_t1', password='testpass123')
        resp = self.client.post(reverse('release_alert'), {'alert_id': alert.pk})
        self.assertEqual(resp.status_code, 302)
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRIAGING)  # not released
        self.assertEqual(alert.release_reason, '')

    def test_release_with_reason_returns_to_pending(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='wz_t1', password='testpass123')
        resp = self.client.post(reverse('release_alert'), {
            'alert_id': alert.pk, 'release_reason': 'Need more context from the host owner.',
        })
        self.assertEqual(resp.status_code, 302)
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_PENDING)
        self.assertIsNone(alert.claimed_by)
        self.assertIn('more context', alert.release_reason)

    def test_tier2_cannot_claim(self):
        alert = WazuhAlert.objects.create(
            opensearch_id='os-pending', timestamp=timezone.now(), rule_level=10,
            triage_status=WazuhAlert.TRIAGE_PENDING,
        )
        self.client.login(username='wz_t2', password='testpass123')
        self.client.post(reverse('claim_alert'), {'alert_id': alert.pk})
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_PENDING)
        self.assertIsNone(alert.claimed_by)


# ──────────────────────────────────────────────────────────────────────────── #
# 13. Triage / wazuh-alert ticket creation integrity                           #
# ──────────────────────────────────────────────────────────────────────────── #

class TriageWorkflowIntegrityTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1       = _make_t1('manual_t1')
        cls.other_t1 = _make_t1('manual_t1_other')
        cls.t2       = _make_t2('manual_t2')

    def test_manual_triage_form_has_no_pre_ticket_decision_or_escalation(self):
        form = TriageForm(user=self.t1)
        self.assertNotIn('decision', form.fields)
        self.assertNotIn('escalated_to', form.fields)

    def test_non_owner_cannot_create_ticket_from_manual_triage(self):
        triage = TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Reported suspicious login.',
            decision=TriageRecord.DECISION_TP, notes='Confirmed by T1.',
        )
        self.client.login(username='manual_t1_other', password='testpass123')
        response = self.client.get(reverse('create_ticket'), {'triage_id': triage.pk})
        self.assertRedirects(response, reverse('triage_list'))
        self.assertFalse(Ticket.objects.exists())

    def test_create_ticket_from_manual_triage_prefills_source(self):
        """The triage source channel auto-fills the ticket's Source (issue_type).

        issue_type and TriageRecord.source share the SOURCE_CHOICES vocabulary,
        so the value carries straight over on the create form's GET.
        """
        triage = TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Reported suspicious login.', source_ip='192.0.2.50',
            decision=TriageRecord.DECISION_TP, notes='Confirmed by T1.',
        )
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'), {'triage_id': triage.pk})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context['form'].initial.get('issue_type'),
            TriageRecord.SOURCE_PHONE,
        )

    def test_wazuh_alert_becomes_true_positive_after_ticket_save(self):
        alert = WazuhAlert.objects.create(
            opensearch_id='ticket-finalize-alert', timestamp=timezone.now(),
            rule_level=14, rule_description='Confirmed ransomware behavior',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
            triage_note='Confirmed malicious.', incident_category=WazuhAlert.CATEGORY_MALWARE,
        )
        self.client.login(username='manual_t1', password='testpass123')
        response = self.client.post(reverse('create_ticket'), _ticket_post_data(
            wazuh_alert=alert.pk,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ESCALATE_T2,
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(wazuh_alert=alert)
        alert.refresh_from_db()
        self.assertEqual(ticket.created_by, self.t1)
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRUE_POSITIVE)
        self.assertIsNone(alert.claimed_by)
        self.assertEqual(ticket.alert_links.get().role, TicketAlertLink.ROLE_PRIMARY)

    def test_alert_bundle_creates_one_incident_with_supporting_alerts(self):
        primary = WazuhAlert.objects.create(
            opensearch_id='ticket-bundle-primary',
            timestamp=timezone.now() - timedelta(minutes=8), rule_level=14,
            rule_description='Ransomware activity on file server',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        supporting = WazuhAlert.objects.create(
            opensearch_id='ticket-bundle-supporting',
            timestamp=timezone.now() - timedelta(minutes=4), rule_level=12,
            rule_description='Lateral movement from file server',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)

        form_response = self.client.get(reverse('create_ticket'), {
            'alert_bundle': [primary.pk, supporting.pk],
        })
        self.assertEqual(form_response.status_code, 200)
        self.assertContains(form_response, 'Alert Bundle (2 Alerts)')
        self.assertEqual(
            form_response.context['form'].initial['wazuh_alert'], primary.pk
        )

        data = _ticket_post_data(
            wazuh_alert=primary.pk,
            alert_bundle=[str(primary.pk), str(supporting.pk)],
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ESCALATE_T2,
        )
        response = self.client.post(reverse('create_ticket'), data)
        self.assertEqual(response.status_code, 302)

        ticket = Ticket.objects.get(wazuh_alert=primary)
        links = {link.alert_id: link.role for link in ticket.alert_links.all()}
        self.assertEqual(links, {
            primary.pk: TicketAlertLink.ROLE_PRIMARY,
            supporting.pk: TicketAlertLink.ROLE_SUPPORTING,
        })
        primary.refresh_from_db()
        supporting.refresh_from_db()
        self.assertEqual(primary.triage_status, WazuhAlert.TRIAGE_TRUE_POSITIVE)
        self.assertEqual(supporting.triage_status, WazuhAlert.TRIAGE_TRUE_POSITIVE)
        self.assertTrue(ticket.logs.filter(note__contains='Alert Bundle').exists())

    def test_alert_bundle_rejects_alert_claimed_by_another_analyst(self):
        other_alert = WazuhAlert.objects.create(
            opensearch_id='ticket-bundle-other-owner', timestamp=timezone.now(),
            rule_level=12, triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.other_t1, claimed_at=timezone.now(),
        )
        owned_alert = WazuhAlert.objects.create(
            opensearch_id='ticket-bundle-owned', timestamp=timezone.now(),
            rule_level=12, triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'), {
            'alert_bundle': [owned_alert.pk, other_alert.pk],
        })
        self.assertRedirects(response, reverse('triage_queue'))
        self.assertFalse(Ticket.objects.exists())

    def test_failed_submit_keeps_the_alert_on_the_case_mode_switch(self):
        # Regression: alert_pk was only assigned on the GET branch, so a POST
        # that came back with errors rendered the single↔multi switch link
        # without the alert and stranded it.
        alert = WazuhAlert.objects.create(
            opensearch_id='ticket-switch-keeps-alert', timestamp=timezone.now(),
            rule_level=12, triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)

        data = _ticket_post_data(wazuh_alert=alert.pk)
        data['device_name'] = ''          # force a validation error
        response = self.client.post(reverse('create_ticket'), data)

        self.assertEqual(response.status_code, 200)
        self.assertFalse(Ticket.objects.exists())
        self.assertIn(
            f'wazuh_alert={alert.pk}', response.context['case_switch_qs'])

    def test_primary_alert_link_must_match_the_ticket_wazuh_alert(self):
        # The primary alert is stored twice (Ticket.wazuh_alert + the PRIMARY
        # link). Nothing in the database can tie them together, so clean() is
        # what stops a ticket from claiming two different primary alerts.
        primary = WazuhAlert.objects.create(
            opensearch_id='link-clean-primary', timestamp=timezone.now(),
            rule_level=12, triage_status=WazuhAlert.TRIAGE_TRIAGING,
        )
        other = WazuhAlert.objects.create(
            opensearch_id='link-clean-other', timestamp=timezone.now(),
            rule_level=12, triage_status=WazuhAlert.TRIAGE_TRIAGING,
        )
        ticket = _make_ticket(created_by=self.t1)
        ticket.wazuh_alert = primary
        ticket.save(update_fields=['wazuh_alert'])

        with self.assertRaises(ValidationError) as ctx:
            TicketAlertLink(
                ticket=ticket, alert=other,
                role=TicketAlertLink.ROLE_PRIMARY,
            ).full_clean()
        self.assertIn('alert', ctx.exception.error_dict)

        # The matching primary, and any supporting alert, validate fine.
        TicketAlertLink(
            ticket=ticket, alert=primary,
            role=TicketAlertLink.ROLE_PRIMARY,
        ).full_clean()
        TicketAlertLink(
            ticket=ticket, alert=other,
            role=TicketAlertLink.ROLE_SUPPORTING,
        ).full_clean()

    def test_alert_links_order_primary_first_regardless_of_role_spelling(self):
        # Ordering must come from an explicit weight, not the alphabetical
        # accident that 'PRIMARY' < 'SUPPORTING'.
        primary = WazuhAlert.objects.create(
            opensearch_id='link-order-primary',
            timestamp=timezone.now(), rule_level=12,
        )
        supporting = WazuhAlert.objects.create(
            opensearch_id='link-order-supporting',
            # Earlier timestamp: it would sort first on any tie-breaker alone.
            timestamp=timezone.now() - timedelta(hours=3), rule_level=12,
        )
        ticket = _make_ticket(created_by=self.t1)
        ticket.wazuh_alert = primary
        ticket.save(update_fields=['wazuh_alert'])
        TicketAlertLink.objects.create(
            ticket=ticket, alert=supporting,
            role=TicketAlertLink.ROLE_SUPPORTING,
        )
        TicketAlertLink.objects.create(
            ticket=ticket, alert=primary, role=TicketAlertLink.ROLE_PRIMARY,
        )

        roles = [link.role for link in ticket.alert_links.all()]
        self.assertEqual(
            roles,
            [TicketAlertLink.ROLE_PRIMARY, TicketAlertLink.ROLE_SUPPORTING],
        )

    def test_wazuh_event_ticket_is_recorded_as_event_history(self):
        alert = WazuhAlert.objects.create(
            opensearch_id='ticket-event-alert', timestamp=timezone.now(),
            rule_level=10, rule_description='Benign scheduled activity',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            wazuh_alert=alert.pk,
            classification=Ticket.CLASSIFICATION_EVENT,
            t1_route='',
        ))
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_FALSE_POSITIVE)

    def test_manual_triage_claim_and_reason_required_release(self):
        triage = TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE,
            analyst=self.t1,
            alert_description='Manual intake awaiting claim.',
            notes='Caller reported unusual behavior.',
        )
        self.client.force_login(self.t1)
        self.client.post(reverse('claim_manual_triage', args=[triage.pk]))
        triage.refresh_from_db()
        self.assertEqual(triage.claimed_by, self.t1)

        self.client.post(reverse('release_manual_triage', args=[triage.pk]), {
            'release_reason': '   ',
        })
        triage.refresh_from_db()
        self.assertEqual(triage.claimed_by, self.t1)

        self.client.post(reverse('release_manual_triage', args=[triage.pk]), {
            'release_reason': 'Shift handoff.',
        })
        triage.refresh_from_db()
        self.assertIsNone(triage.claimed_by)
        self.assertEqual(triage.release_reason, 'Shift handoff.')

    def test_my_queue_shows_manual_reports_and_returned_tickets(self):
        """My Queue = the manual-intake queue plus the analyst's own-court
        tickets — above all a case Tier 2 returned (T1_REVIEW)."""
        TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Caller reported odd VPN logins.',
            notes='Awaiting triage.',
        )
        returned = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_T1_REVIEW,
        )
        # Another analyst's returned case must NOT appear — T1_REVIEW is
        # creator-gated, so it is not this analyst's work.
        _make_ticket(
            created_by=self.other_t1,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_T1_REVIEW,
        )
        # A ticket parked with Tier 2 is not own-court work either.
        _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            status=Ticket.STATUS_ESCALATED_T2,
        )

        self.client.force_login(self.t1)
        response = self.client.get(reverse('my_queue'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Caller reported odd VPN logins.')
        my_ids = {t.pk for t in response.context['my_tickets']}
        self.assertEqual(my_ids, {returned.pk})
        self.assertEqual(response.context['returned_count'], 1)

    def test_my_queue_and_triage_list_are_the_same_page(self):
        """The historical URL name must keep landing on the merged page, so
        every manual-triage redirect and deep link still works."""
        self.client.force_login(self.t1)
        for name in ('my_queue', 'triage_list'):
            with self.subTest(url_name=name):
                response = self.client.get(reverse(name))
                self.assertEqual(response.status_code, 200)
                self.assertTemplateUsed(response, 'incidents/my_queue.html')

    def test_my_queue_badge_counts_pickable_reports_and_own_tickets(self):
        TriageRecord.objects.create(
            source=TriageRecord.SOURCE_EMAIL, analyst=self.t1,
            alert_description='Unclaimed report.', notes='n',
        )
        TriageRecord.objects.create(
            source=TriageRecord.SOURCE_EMAIL, analyst=self.t1,
            alert_description='Claimed by someone else.', notes='n',
            claimed_by=self.other_t1, claimed_at=timezone.now(),
        )
        _make_ticket(created_by=self.t1, status=Ticket.STATUS_T1_REVIEW,
                     classification=Ticket.CLASSIFICATION_INCIDENT)

        self.client.force_login(self.t1)
        response = self.client.get(reverse('my_queue'))
        # 1 unclaimed report + 1 own ticket; the peer-claimed report excluded.
        self.assertEqual(response.context['my_queue_count'], 2)

    def test_case_mode_switch_offers_both_scopes_from_either_form(self):
        """The two creation forms share one menu entry, so each must expose the
        switch to the other scope."""
        self.client.force_login(self.t1)
        for url_name, mode in (
            ('create_ticket', 'single'), ('create_project_incident', 'multi'),
        ):
            with self.subTest(form=url_name):
                response = self.client.get(reverse(url_name))
                self.assertEqual(response.status_code, 200)
                self.assertEqual(response.context['case_mode'], mode)
                self.assertContains(response, reverse('create_ticket'))
                self.assertContains(response, reverse('create_project_incident'))

    def test_switching_scope_keeps_the_originating_triage_record(self):
        """Switching single ↔ multi must not orphan the case from its source,
        or the new ticket would come back unlinked to the intake record."""
        triage = TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Multi-host beaconing reported.', notes='n',
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('create_ticket'), {'triage_id': triage.pk})
        self.assertIn(f'triage_id={triage.pk}', response.context['case_switch_qs'])

        response = self.client.get(
            reverse('create_project_incident'), {'triage_id': triage.pk})
        self.assertIn(f'triage_id={triage.pk}', response.context['case_switch_qs'])

    def test_switching_scope_keeps_the_originating_wazuh_alert(self):
        alert = WazuhAlert.objects.create(
            opensearch_id='case-switch-alert', timestamp=timezone.now(),
            rule_level=12, rule_description='Lateral movement detected',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('create_ticket'), {'wazuh_alert': alert.pk})
        self.assertIn(f'wazuh_alert={alert.pk}', response.context['case_switch_qs'])

    def test_case_switch_query_string_is_empty_for_a_blank_form(self):
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'))
        self.assertEqual(response.context['case_switch_qs'], '')

    def _claimed_report(self, claimer):
        return TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Prank call, nothing security-relevant.',
            notes='Initial intake.', claimed_by=claimer, claimed_at=timezone.now(),
        )

    def test_dismiss_requires_a_reason(self):
        triage = self._claimed_report(self.t1)
        self.client.force_login(self.t1)
        self.client.post(reverse('dismiss_manual_triage', args=[triage.pk]), {
            'dismiss_reason': '   ',
        })
        triage.refresh_from_db()
        self.assertEqual(triage.decision, '')
        self.assertEqual(triage.claimed_by, self.t1)

    def test_dismiss_records_who_disposed_of_the_report(self):
        """The claim is cleared on dismissal, so resolved_by is the only durable
        record of who threw the report away."""
        triage = self._claimed_report(self.t1)
        self.client.force_login(self.t1)
        self.client.post(reverse('dismiss_manual_triage', args=[triage.pk]), {
            'dismiss_reason': 'spam',
        })
        triage.refresh_from_db()
        self.assertEqual(triage.resolved_by, self.t1)
        self.assertIsNotNone(triage.resolved_at)
        self.assertIsNone(triage.claimed_by)

    def test_conversion_records_who_handled_the_report(self):
        triage = self._claimed_report(self.t1)
        self.client.force_login(self.t1)
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            triage_id=triage.pk,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ESCALATE_T2,
        ))
        triage.refresh_from_db()
        self.assertIsNotNone(triage.ticket)
        self.assertEqual(triage.resolved_by, self.t1)
        self.assertIsNotNone(triage.resolved_at)

    def test_my_queue_history_is_scoped_to_the_current_analyst(self):
        mine = self._claimed_report(self.t1)
        theirs = self._claimed_report(self.other_t1)

        self.client.force_login(self.t1)
        self.client.post(reverse('dismiss_manual_triage', args=[mine.pk]), {
            'dismiss_reason': 'mine',
        })
        self.client.force_login(self.other_t1)
        self.client.post(reverse('dismiss_manual_triage', args=[theirs.pk]), {
            'dismiss_reason': 'theirs',
        })

        self.client.force_login(self.t1)
        history = self.client.get(reverse('my_queue')).context['manual_history']
        self.assertEqual([t.pk for t in history], [mine.pk])

    def test_dismiss_closes_the_report_without_a_ticket(self):
        triage = self._claimed_report(self.t1)
        self.client.force_login(self.t1)
        self.client.post(reverse('dismiss_manual_triage', args=[triage.pk]), {
            'dismiss_reason': 'แจ้งผิด ไม่เกี่ยวกับความปลอดภัย',
        })
        triage.refresh_from_db()
        self.assertEqual(triage.decision, TriageRecord.DECISION_FP)
        self.assertIsNone(triage.ticket)
        self.assertIsNone(triage.claimed_by)
        self.assertIn('แจ้งผิด ไม่เกี่ยวกับความปลอดภัย', triage.notes)
        # Left the queue, entered history, and can no longer become a ticket.
        response = self.client.get(reverse('my_queue'))
        self.assertNotIn(triage, response.context['manual_queue'])
        self.assertIn(triage, list(response.context['manual_history']))
        self.assertFalse(_can_create_ticket_from_triage(triage, self.t1))

    def test_only_the_claimer_may_dismiss(self):
        triage = self._claimed_report(self.t1)
        self.client.force_login(self.other_t1)
        self.client.post(reverse('dismiss_manual_triage', args=[triage.pk]), {
            'dismiss_reason': 'not mine to close',
        })
        triage.refresh_from_db()
        self.assertEqual(triage.decision, '')
        self.assertEqual(triage.claimed_by, self.t1)

    def test_invalid_ticket_form_keeps_wazuh_alert_in_progress(self):
        alert = WazuhAlert.objects.create(
            opensearch_id='invalid-ticket-alert', timestamp=timezone.now(),
            rule_level=12, rule_description='Suspicious command execution',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=self.t1, claimed_at=timezone.now(),
        )
        self.client.login(username='manual_t1', password='testpass123')
        response = self.client.post(reverse('create_ticket'), {'wazuh_alert': alert.pk})
        self.assertEqual(response.status_code, 200)
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRIAGING)
        self.assertFalse(Ticket.objects.filter(wazuh_alert=alert).exists())


# ──────────────────────────────────────────────────────────────────────────── #
# 14. Superuser bypass                                                          #
# ──────────────────────────────────────────────────────────────────────────── #

class SuperuserAccessTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.superuser = User.objects.create_superuser(
            username='all_access_superuser', email='superuser@example.com',
            password='testpass123',
        )
        cls.system_admin = _make_user('superuser_target_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.t1 = _make_t1('superuser_t1')
        cls.t2 = _make_t2('superuser_t2')

    def setUp(self):
        self.client.force_login(self.superuser)

    def test_superuser_without_profile_sees_all_tickets(self):
        first = _make_ticket(issue_description='First ticket')
        second = _make_ticket(issue_description='Second ticket', assigned_admin=self.system_admin)
        self.assertFalse(hasattr(self.superuser, 'profile'))
        self.assertQuerySetEqual(
            Ticket.objects.visible_to(self.superuser).order_by('pk'), [first, second],
        )

    def test_superuser_can_access_core_pages(self):
        ticket = _make_ticket()
        urls = [
            reverse('home'), reverse('ticket_list'), reverse('create_ticket'),
            reverse('ticket_detail', args=[ticket.pk]), reverse('ticket_history'),
            reverse('triage_list'), reverse('create_triage'),
            reverse('system_owner_dashboard'),
        ]
        for url in urls:
            with self.subTest(url=url):
                self.assertEqual(self.client.get(url).status_code, 200)

    def test_superuser_can_perform_every_ticket_role_transition(self):
        # Emergency flag set so the ticket legally routes through PENDING_MANAGER
        # (severity alone no longer triggers the manager gate).
        ticket = _make_ticket(
            assigned_admin=self.system_admin, severity='Critical',
            classification=Ticket.CLASSIFICATION_INCIDENT, is_emergency=True,
            t1_route=Ticket.T1_ROUTE_ADMIN,
        )
        ticket.transition_to(Ticket.STATUS_PENDING_MGR_TRIAGE, self.superuser, 'as t1')
        ticket.transition_to(Ticket.STATUS_AWAITING_CONTAINMENT, self.superuser, 'as manager forward')
        ticket.containment_report = 'Contained by superuser.'
        ticket.transition_to(Ticket.STATUS_CONTAINMENT_REPORTED, self.superuser, 'as admin')
        ticket.transition_to(Ticket.STATUS_PENDING_MANAGER, self.superuser, 'verify')
        ticket.transition_to(Ticket.STATUS_APPROVED, self.superuser, 'as manager')
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.STATUS_APPROVED)
        self.assertEqual(ticket.verified_by, self.superuser)
        self.assertEqual(ticket.approved_by, self.superuser)

    def test_superuser_can_submit_containment_for_any_ticket(self):
        ticket = _make_ticket(
            assigned_admin=self.system_admin, status=Ticket.STATUS_AWAITING_CONTAINMENT,
            classification=Ticket.CLASSIFICATION_INCIDENT,
        )
        response = self.client.post(reverse('ticket_detail', args=[ticket.pk]), {
            'action': 'containment',
            'containment_report': 'Superuser containment report.',
            'note': 'Completed with all-role access.',
        })
        self.assertRedirects(response, reverse('ticket_detail', args=[ticket.pk]))
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.STATUS_CONTAINMENT_REPORTED)


# ──────────────────────────────────────────────────────────────────────────── #
# 15. Attachment download authorization                                         #
# ──────────────────────────────────────────────────────────────────────────── #

@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_test_media_'))
class AttachmentDownloadSecurityTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.soc_staff = _make_t1('att_soc')
        cls.admin_a = _make_user('att_admin_a', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin_b = _make_user('att_admin_b', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.ticket_a = _make_ticket(assigned_admin=cls.admin_a)
        cls.attachment = TicketAttachment.objects.create(
            ticket=cls.ticket_a,
            file=SimpleUploadedFile(
                'evidence.html', b'<script>alert(document.cookie)</script>',
                content_type='text/html',
            ),
            original_name='evidence.html', uploaded_by=cls.soc_staff,
        )

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)

    def _url(self):
        return reverse('download_attachment', args=[self.attachment.pk])

    def test_unauthenticated_redirected_to_login(self):
        response = self.client.get(self._url())
        self.assertEqual(response.status_code, 302)
        self.assertIn('/login/', response['Location'])

    def test_authorized_user_downloads_with_safe_headers(self):
        self.client.force_login(self.soc_staff)
        response = self.client.get(self._url())
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response['Content-Disposition'].startswith('attachment'))
        self.assertEqual(response['X-Content-Type-Options'], 'nosniff')

    def test_admin_cannot_download_attachment_on_unrelated_ticket(self):
        self.client.force_login(self.admin_b)
        self.assertEqual(self.client.get(self._url()).status_code, 404)


# ──────────────────────────────────────────────────────────────────────────── #
# 16. Attachment upload size limit                                              #
# ──────────────────────────────────────────────────────────────────────────── #

class AttachmentUploadLimitTest(TestCase):
    def test_oversize_file_rejected_by_form(self):
        with patch('apps.incidents.models.MAX_ATTACHMENT_SIZE', 10):
            form = AttachmentForm(
                data={'description': ''},
                files={'file': SimpleUploadedFile(
                    'big.bin', b'01234567890', content_type='application/octet-stream',
                )},
            )
            self.assertFalse(form.is_valid())
            self.assertIn('file', form.errors)

    def test_within_limit_file_accepted_by_form(self):
        form = AttachmentForm(
            data={'description': 'ok'},
            files={'file': SimpleUploadedFile('small.txt', b'hello', content_type='text/plain')},
        )
        self.assertTrue(form.is_valid(), msg=form.errors)


# ──────────────────────────────────────────────────────────────────────────── #
# 16b. Attachment upload type / content validation                              #
# ──────────────────────────────────────────────────────────────────────────── #

class AttachmentUploadTypeTest(TestCase):
    def test_disallowed_extension_rejected(self):
        """Active-web content (.html) is not on the evidence allowlist."""
        form = AttachmentForm(
            data={'description': ''},
            files={'file': SimpleUploadedFile(
                'evidence.html', b'<script>alert(1)</script>',
                content_type='text/html',
            )},
        )
        self.assertFalse(form.is_valid())
        self.assertIn('file', form.errors)

    def test_extensionless_file_rejected(self):
        form = AttachmentForm(
            data={'description': ''},
            files={'file': SimpleUploadedFile('noext', b'data')},
        )
        self.assertFalse(form.is_valid())
        self.assertIn('file', form.errors)

    def test_spoofed_image_content_rejected(self):
        """An allowed extension whose bytes don't match its type is refused."""
        form = AttachmentForm(
            data={'description': ''},
            files={'file': SimpleUploadedFile(
                'shot.png', b'<svg onload=alert(1)>', content_type='image/png',
            )},
        )
        self.assertFalse(form.is_valid())
        self.assertIn('file', form.errors)

    def test_valid_png_accepted(self):
        png = b'\x89PNG\r\n\x1a\n' + b'\x00' * 32
        form = AttachmentForm(
            data={'description': 'ok'},
            files={'file': SimpleUploadedFile('shot.png', png, content_type='image/png')},
        )
        self.assertTrue(form.is_valid(), msg=form.errors)

    def test_log_evidence_accepted(self):
        form = AttachmentForm(
            data={'description': 'ok'},
            files={'file': SimpleUploadedFile('firewall.log', b'deny 1.2.3.4')},
        )
        self.assertTrue(form.is_valid(), msg=form.errors)


# ──────────────────────────────────────────────────────────────────────────── #
@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_attachment_test_media_'))
class AttachmentWorkflowPermissionTest(TestCase):
    """Ticket-level uploads belong to the role currently handling the ticket."""

    @classmethod
    def setUpTestData(cls):
        cls.creator = _make_t1('attachment_creator')
        cls.other_t1 = _make_t1('attachment_other_t1')
        cls.t2 = _make_t2('attachment_t2')
        cls.admin = _make_user('attachment_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.owner = _make_user('attachment_owner', UserProfile.ROLE_SYSTEM_OWNER)

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)

    def _ticket(self, status=Ticket.STATUS_NEW, **kwargs):
        return _make_ticket(
            status=status, created_by=self.creator, assigned_admin=self.admin,
            system_owner=self.owner, **kwargs,
        )

    def _upload(self, user, ticket):
        self.client.force_login(user)
        return self.client.post(
            reverse('upload_attachment', args=[ticket.pk]),
            data={
                'file': SimpleUploadedFile('evidence.log', b'deny 192.0.2.1'),
                'description': 'Investigation notes',
            },
        )

    def test_creator_can_upload_while_ticket_is_in_tier1_care(self):
        ticket = self._ticket()
        response = self._upload(self.creator, ticket)
        self.assertRedirects(response, reverse('ticket_detail', args=[ticket.pk]))
        attachment = TicketAttachment.objects.get(ticket=ticket)
        self.assertEqual(attachment.uploaded_by, self.creator)
        self.assertEqual(attachment.description, 'Investigation notes')

    def test_visible_but_non_responsible_user_cannot_upload(self):
        ticket = self._ticket()
        self._upload(self.other_t1, ticket)
        self.assertFalse(TicketAttachment.objects.filter(ticket=ticket).exists())

    def test_assigned_admin_can_upload_during_containment(self):
        ticket = self._ticket(status=Ticket.STATUS_AWAITING_CONTAINMENT)
        self._upload(self.admin, ticket)
        self.assertTrue(TicketAttachment.objects.filter(
            ticket=ticket, uploaded_by=self.admin).exists())

    def test_system_owner_can_upload_while_owner_is_handling_ticket(self):
        ticket = self._ticket(status=Ticket.STATUS_AWAITING_OWNER)
        self._upload(self.owner, ticket)
        self.assertTrue(TicketAttachment.objects.filter(
            ticket=ticket, uploaded_by=self.owner).exists())

    def test_tier2_can_upload_during_tier2_review(self):
        ticket = self._ticket(status=Ticket.STATUS_ESCALATED_T2)
        self._upload(self.t2, ticket)
        self.assertTrue(TicketAttachment.objects.filter(
            ticket=ticket, uploaded_by=self.t2).exists())

    def test_closed_ticket_rejects_upload_server_side(self):
        ticket = self._ticket(status=Ticket.STATUS_APPROVED)
        self._upload(self.creator, ticket)
        self.assertFalse(TicketAttachment.objects.filter(ticket=ticket).exists())

    def test_closed_ticket_rejects_subtask_result_upload_server_side(self):
        ticket = self._ticket(status=Ticket.STATUS_APPROVED)
        subtask = TicketSubtask.objects.create(
            ticket=ticket,
            subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Collect logs',
            assigned_to=self.creator,
        )
        self.client.force_login(self.creator)
        self.client.post(
            reverse('update_subtask', args=[subtask.pk]),
            data={
                'status': TicketSubtask.STATUS_DONE,
                'result_notes': 'Complete',
                'result_file': SimpleUploadedFile('result.log', b'complete'),
            },
        )
        subtask.refresh_from_db()
        self.assertEqual(subtask.status, TicketSubtask.STATUS_OPEN)
        self.assertFalse(TicketAttachment.objects.filter(ticket=ticket).exists())

    def _update_subtask(self, user, subtask, filename='result.log'):
        self.client.force_login(user)
        return self.client.post(
            reverse('update_subtask', args=[subtask.pk]),
            data={
                'status': TicketSubtask.STATUS_DONE,
                'result_notes': 'Report attached',
                'result_file': SimpleUploadedFile(filename, b'findings'),
            },
        )

    def test_soc_staff_cannot_attach_a_file_through_the_subtask_route(self):
        # K16. update_subtask used to accept a file from anyone with is_soc,
        # letting a non-creator T1 put an attachment on a ticket that
        # _can_upload_ticket_attachment() would refuse them. The notes/status
        # update still goes through — only the file is rejected.
        ticket = self._ticket(status=Ticket.STATUS_T1_REVIEW)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Collect logs', assigned_to=self.creator,
        )
        self.assertFalse(_can_upload_ticket_attachment(ticket, self.other_t1))

        self._update_subtask(self.other_t1, subtask)

        self.assertFalse(TicketAttachment.objects.filter(ticket=ticket).exists())
        subtask.refresh_from_db()
        self.assertEqual(subtask.status, TicketSubtask.STATUS_DONE)

    def test_assignee_can_attach_a_result_even_when_the_ticket_is_elsewhere(self):
        # The documented exception: for a response request the court that
        # matters is the REQUEST, not the ticket's workflow status. The
        # ticket-level rule for PENDING_MGR_TRIAGE answers is_soc_manager,
        # which would wrongly refuse the assigned responder.
        forensic = _make_forensic('attachment_forensic')
        ticket = self._ticket(status=Ticket.STATUS_PENDING_MGR_TRIAGE)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=forensic,
        )
        self.assertFalse(_can_upload_ticket_attachment(ticket, forensic))

        self._update_subtask(forensic, subtask, filename='rca.log')

        self.assertTrue(TicketAttachment.objects.filter(
            ticket=ticket, subtask=subtask, uploaded_by=forensic).exists())

    def test_soc_manager_can_attach_a_result(self):
        mgr = _make_user('attachment_sub_mgr', UserProfile.ROLE_SOC_MANAGER)
        ticket = self._ticket(status=Ticket.STATUS_T1_REVIEW)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Collect logs', assigned_to=self.creator,
        )
        self._update_subtask(mgr, subtask)
        self.assertTrue(TicketAttachment.objects.filter(
            ticket=ticket, uploaded_by=mgr).exists())

    def test_response_assignee_can_delete_a_response_request_deliverable(self):
        forensic = _make_forensic('delete_response_forensic')
        ticket = self._ticket(status=Ticket.STATUS_PENDING_MGR_TRIAGE)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=forensic,
        )
        attachment = TicketAttachment.objects.create(
            ticket=ticket, subtask=subtask,
            file=SimpleUploadedFile('rca.log', b'findings'),
            original_name='rca.log', uploaded_by=self.creator,
        )

        self.assertTrue(_can_delete_ticket_attachment(ticket, attachment, forensic))
        self.client.force_login(forensic)
        detail = self.client.get(reverse('ticket_detail', args=[ticket.pk]))
        self.assertContains(detail, f'data-bs-target="#delete-att-{attachment.pk}"')
        response = self.client.post(
            reverse('delete_attachment', args=[attachment.pk]),
            {'reason': 'Superseded by the corrected RCA report'},
        )

        self.assertRedirects(response, reverse('ticket_detail', args=[ticket.pk]))
        self.assertFalse(TicketAttachment.objects.filter(pk=attachment.pk).exists())
        retained = TicketAttachment.all_objects.get(pk=attachment.pk)
        self.assertEqual(retained.deleted_by, forensic)

    def test_soc_manager_can_delete_a_response_request_deliverable(self):
        forensic = _make_forensic('delete_response_manager_forensic')
        manager = _make_user(
            'delete_response_manager', UserProfile.ROLE_SOC_MANAGER)
        ticket = self._ticket(status=Ticket.STATUS_PENDING_MGR_TRIAGE)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=forensic,
        )
        attachment = TicketAttachment.objects.create(
            ticket=ticket, subtask=subtask,
            file=SimpleUploadedFile('rca.log', b'findings'),
            original_name='rca.log', uploaded_by=self.creator,
        )

        self.client.force_login(manager)
        response = self.client.post(
            reverse('delete_attachment', args=[attachment.pk]),
            {'reason': 'Replaced with a redacted report'},
        )

        self.assertRedirects(response, reverse('ticket_detail', args=[ticket.pk]))
        self.assertFalse(TicketAttachment.objects.filter(pk=attachment.pk).exists())

    def test_unassigned_responder_cannot_delete_a_response_request_deliverable(self):
        assignee = _make_forensic('delete_response_assignee')
        other_responder = _make_forensic('delete_response_other')
        ticket = self._ticket(status=Ticket.STATUS_PENDING_MGR_TRIAGE)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=assignee,
        )
        attachment = TicketAttachment.objects.create(
            ticket=ticket, subtask=subtask,
            file=SimpleUploadedFile('rca.log', b'findings'),
            original_name='rca.log', uploaded_by=self.creator,
        )

        self.assertFalse(
            _can_delete_ticket_attachment(ticket, attachment, other_responder))
        self.client.force_login(other_responder)
        self.client.post(
            reverse('delete_attachment', args=[attachment.pk]),
            {'reason': 'Should not be accepted'},
        )

        self.assertTrue(TicketAttachment.objects.filter(pk=attachment.pk).exists())

    def test_closed_ticket_rejects_response_assignee_deletion(self):
        forensic = _make_forensic('delete_response_closed_forensic')
        ticket = self._ticket(status=Ticket.STATUS_APPROVED)
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=forensic,
        )
        attachment = TicketAttachment.objects.create(
            ticket=ticket, subtask=subtask,
            file=SimpleUploadedFile('rca.log', b'findings'),
            original_name='rca.log', uploaded_by=self.creator,
        )

        self.assertFalse(_can_delete_ticket_attachment(ticket, attachment, forensic))
        self.client.force_login(forensic)
        self.client.post(
            reverse('delete_attachment', args=[attachment.pk]),
            {'reason': 'Should not be accepted'},
        )

        self.assertTrue(TicketAttachment.objects.filter(pk=attachment.pk).exists())

    def test_delete_retains_attachment_and_records_ticket_audit(self):
        ticket = self._ticket()
        attachment = TicketAttachment.objects.create(
            ticket=ticket,
            file=SimpleUploadedFile('retain.log', b'retain this evidence'),
            original_name='retain.log',
            uploaded_by=self.creator,
        )
        file_name = attachment.file.name
        self.client.force_login(self.creator)
        response = self.client.post(
            reverse('delete_attachment', args=[attachment.pk]),
            {'reason': 'แนบผิดเคส'},
        )
        self.assertRedirects(response, reverse('ticket_detail', args=[ticket.pk]))
        self.assertFalse(TicketAttachment.objects.filter(pk=attachment.pk).exists())

        retained = TicketAttachment.all_objects.get(pk=attachment.pk)
        self.assertEqual(retained.deleted_by, self.creator)
        self.assertIsNotNone(retained.deleted_at)
        self.assertEqual(retained.deleted_reason, 'แนบผิดเคส')
        self.assertTrue(retained.file.storage.exists(file_name))
        self.assertTrue(TicketLog.objects.filter(
            ticket=ticket,
            author=self.creator,
            note='Attachment removed: retain.log — เหตุผล: แนบผิดเคส',
        ).exists())

    def test_soft_deleted_attachment_cannot_be_downloaded(self):
        ticket = self._ticket()
        attachment = TicketAttachment.objects.create(
            ticket=ticket,
            file=SimpleUploadedFile('removed.log', b'not visible'),
            original_name='removed.log',
            uploaded_by=self.creator,
            deleted_by=self.creator,
            deleted_at=timezone.now(),
        )
        self.client.force_login(self.creator)
        response = self.client.get(reverse('download_attachment', args=[attachment.pk]))
        self.assertEqual(response.status_code, 404)

    def test_detail_page_shows_attachment_accountability_metadata(self):
        ticket = self._ticket()
        attachment = TicketAttachment.objects.create(
            ticket=ticket,
            file=SimpleUploadedFile('metadata.log', b'uploaded'),
            original_name='metadata.log',
            description='Firewall export',
            uploaded_by=self.creator,
        )
        self.client.force_login(self.creator)
        response = self.client.get(reverse('ticket_detail', args=[ticket.pk]))
        self.assertContains(response, attachment.original_name)
        self.assertContains(response, 'Firewall export')
        self.assertContains(response, 'Uploaded by')
        self.assertContains(response, self.creator.username)


# 17. 'Unknown' severity (additive, human-assigned)                            #
# ──────────────────────────────────────────────────────────────────────────── #

class UnknownSeverityTest(TestCase):
    """
    'Unknown' is a human-assigned severity for cases the analyst cannot yet
    classify. It is selectable in the manual create + manual triage forms,
    absent from the automated Wazuh mapping, ranks lowest for queue ordering
    (severity never routes to the manager — only the emergency flag), and
    renders a distinct badge.
    """

    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('uk_t1')
        cls.t2    = _make_t2('uk_t2')
        cls.mgr   = _make_user('uk_mgr',   UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('uk_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    # ── Model / choices ─────────────────────────────────────────────────── #
    def test_unknown_is_a_model_choice(self):
        self.assertIn(('Unknown', 'Unknown'), Ticket.SEVERITY_CHOICES)

    def test_unknown_ranks_lowest(self):
        ranks = Ticket.SEVERITY_RANK
        self.assertEqual(ranks['Unknown'], 0)
        self.assertLess(ranks['Unknown'], min(
            ranks['Low'], ranks['Medium'], ranks['High'], ranks['Critical']
        ))

    # ── Availability: manual create + manual triage forms ───────────────── #
    def test_unknown_selectable_on_manual_create_form(self):
        values = [v for v, _ in TicketForm().fields['severity'].choices]
        self.assertIn('Unknown', values)

    def test_unknown_selectable_on_manual_triage_create_form(self):
        # A ticket opened from a manual TriageRecord uses the same TicketForm,
        # so 'Unknown' must be offered there too.
        triage = TriageRecord.objects.create(
            source=TriageRecord.SOURCE_PHONE, analyst=self.t1,
            alert_description='Caller reported odd activity, severity unclear.',
            decision=TriageRecord.DECISION_TP, notes='Cannot yet classify.',
        )
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'), {'triage_id': triage.pk})
        self.assertEqual(response.status_code, 200)
        values = [v for v, _ in response.context['form'].fields['severity'].choices]
        self.assertIn('Unknown', values)

    def test_unknown_accepted_when_creating_ticket(self):
        self.client.force_login(self.t1)
        response = self.client.post(reverse('create_ticket'), _ticket_post_data(
            severity='Unknown',
            classification=Ticket.CLASSIFICATION_INCIDENT,
            t1_route=TicketForm.ROUTE_ESCALATE_T2,
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(device_name='TEST-ENDPOINT-01')
        self.assertEqual(ticket.severity, 'Unknown')

    # ── Availability: NOT in automated Wazuh ingest mapping ─────────────── #
    def test_unknown_absent_from_wazuh_severity_mapping(self):
        from apps.wazuh_ingest.views import _severity_for_rule_level
        mapped = {_severity_for_rule_level(level) for level in range(0, 20)}
        self.assertNotIn('Unknown', mapped)
        self.assertTrue(mapped <= {'Critical', 'High', 'Medium', 'Low'})

    # ── Routing: emergency flag is the only path to the manager ─────────── #
    def _contained(self, is_emergency=False):
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT,
            assigned_admin=self.admin, status=Ticket.STATUS_CONTAINMENT_REPORTED,
            severity='Unknown', is_emergency=is_emergency, containment_report='done',
        )

    def test_unknown_without_emergency_does_not_require_manager(self):
        self.assertFalse(self._contained().requires_manager_verification)

    def test_unknown_without_emergency_t2_closes_directly(self):
        t = self._contained()
        t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'closed — no manager needed')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_unknown_without_emergency_cannot_route_to_manager(self):
        t = self._contained()
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'no need')

    def test_unknown_with_emergency_requires_manager(self):
        self.assertTrue(self._contained(is_emergency=True).requires_manager_verification)

    def test_unknown_with_emergency_routes_to_manager(self):
        t = self._contained(is_emergency=True)
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'blocked by emergency')
        t.transition_to(Ticket.STATUS_PENDING_MANAGER, self.t2, 'route to manager')
        t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'approved')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    # ── Display: distinct badge ─────────────────────────────────────────── #
    def test_unknown_badge_renders_distinctly(self):
        from django.template.loader import render_to_string
        html = render_to_string('incidents/_severity_badge.html', {'severity': 'Unknown'})
        self.assertIn('Unknown', html)
        self.assertIn('#6f42c1', html)  # distinct colour, not reused by any severity
        # Sanity: the Unknown badge must not borrow an existing severity colour.
        for other_colour in ('bg-danger', '#fd7e14', 'bg-warning', 'bg-success'):
            self.assertNotIn(other_colour, html)


# ──────────────────────────────────────────────────────────────────────────── #
# Threat-type cascade: detailed_issue → detailed_issue2                         #
# ──────────────────────────────────────────────────────────────────────────── #

class DetailedIssueCascadeTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('cascade_t1')

    def test_form_hides_legacy_source_flavoured_categories(self):
        """Only the 10 clean threat categories are offered; leftovers hidden."""
        form = TicketForm(user=self.t1)
        codes = [c for c, _ in form.fields['detailed_issue'].choices]
        self.assertIn('Malicious Logic', codes)
        self.assertNotIn('SIEM Other', codes)
        self.assertNotIn('TI IOC', codes)
        self.assertNotIn('External Other', codes)

    def test_mismatched_detailed_issue_pair_is_rejected(self):
        form = TicketForm(
            data=_ticket_post_data(
                detailed_issue='Malicious Logic',
                detailed_issue2='Port Scanning',  # belongs to Reconnaissance
            ),
            user=self.t1,
        )
        self.assertFalse(form.is_valid())
        self.assertIn('detailed_issue2', form.errors)

    def test_matching_detailed_issue_pair_is_accepted(self):
        form = TicketForm(
            data=_ticket_post_data(
                detailed_issue='Malicious Logic',
                detailed_issue2='Ransomware Behavior',
            ),
            user=self.t1,
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_create_form_prefills_parent_from_detailed_issue2(self):
        """A detailed_issue2 passed in the URL (e.g. from Wazuh) sets its parent."""
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'), {'detailed_issue2': 'Malware EDR'})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['form'].initial.get('detailed_issue'), 'Malicious Logic')
        self.assertEqual(response.context['form'].initial.get('detailed_issue2'), 'Malware EDR')

    def test_create_form_renders_cascade_wiring(self):
        """The cascade include renders: JSON payload + the ids the JS targets."""
        self.client.force_login(self.t1)
        html = self.client.get(reverse('create_ticket')).content.decode()
        self.assertIn('id="detailed-issue-cascade"', html)   # json_script payload
        self.assertIn('Malicious Logic', html)               # a hierarchy key is embedded
        self.assertIn('id="id_detailed_issue"', html)        # parent select
        self.assertIn('id="id_detailed_issue2"', html)       # child select


class TicketListOlaFilterTest(TestCase):
    """The ticket-list ?ola= filter buckets the active queue by time-to-deadline,
    sharing thresholds with the dashboard OLA chart (apps.incidents.ola)."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('ola_filter_t1')  # SOC staff → sees all tickets
        now = timezone.now()
        cls.overdue  = _make_ticket(ola_contain_deadline=now - timedelta(hours=1))
        cls.due_1h   = _make_ticket(ola_contain_deadline=now + timedelta(minutes=30))
        cls.due_4h   = _make_ticket(ola_contain_deadline=now + timedelta(hours=2))
        cls.on_track = _make_ticket(ola_contain_deadline=now + timedelta(hours=10))

    def _list(self, **params):
        self.client.force_login(self.t1)
        resp = self.client.get(reverse('ticket_list'), params)
        return {t.pk for t in resp.context['page_obj']}, resp

    def test_overdue_filter_returns_only_overdue(self):
        ids, resp = self._list(ola='overdue')
        self.assertEqual(ids, {self.overdue.pk})
        self.assertEqual(resp.context['ola_filter'], 'overdue')

    def test_due_1h_filter_returns_only_due_within_1h(self):
        ids, _ = self._list(ola='due_1h')
        self.assertEqual(ids, {self.due_1h.pk})

    def test_on_track_filter_returns_only_on_track(self):
        ids, _ = self._list(ola='on_track')
        self.assertEqual(ids, {self.on_track.pk})

    def test_no_filter_returns_all_active(self):
        ids, _ = self._list()
        self.assertEqual(
            ids, {self.overdue.pk, self.due_1h.pk, self.due_4h.pk, self.on_track.pk})

    def test_invalid_bucket_is_ignored(self):
        ids, resp = self._list(ola='bogus')
        self.assertEqual(resp.context['ola_filter'], '')
        self.assertEqual(len(ids), 4)


class OlaBadgeTest(TestCase):
    """The shared queue OLA pill (incidents/_ola_badge.html via ola.badge_for).

    The badge must agree with bucket_case/bucket_filter — a list sorted by the
    ORM bucket showing contradicting badges would be worse than no badge.
    """

    def setUp(self):
        self.now = timezone.now()

    def _badge(self, deadline, **kwargs):
        return ola_buckets.badge_for(deadline, now=self.now, **kwargs)

    def test_bucket_matches_the_orm_bucketing(self):
        cases = [
            (self.now - timedelta(hours=1), ola_buckets.OVERDUE),
            (self.now + timedelta(minutes=30), ola_buckets.DUE_1H),
            (self.now + timedelta(hours=2), ola_buckets.DUE_4H),
            (self.now + timedelta(hours=10), ola_buckets.ON_TRACK),
        ]
        for deadline, expected in cases:
            with self.subTest(bucket=expected):
                self.assertEqual(self._badge(deadline)['bucket'], expected)
                self.assertEqual(
                    ola_buckets.bucket_for(deadline, self.now), expected)

    def test_overdue_reads_as_elapsed_and_on_track_as_remaining(self):
        overdue = self._badge(self.now - timedelta(hours=2, minutes=5))
        self.assertTrue(overdue['overdue'])
        self.assertTrue(overdue['label'].startswith('เกิน'))

        remaining = self._badge(self.now + timedelta(minutes=45))
        self.assertFalse(remaining['overdue'])
        self.assertEqual(remaining['label'], 'เหลือ 45 น.')

    def test_nothing_to_show_without_a_deadline_or_once_done(self):
        # Medium/Low never get a contain deadline; finished work has no clock.
        self.assertIsNone(self._badge(None))
        self.assertIsNone(self._badge(self.now + timedelta(hours=1), done=True))

    def test_ticket_badge_uses_the_contain_deadline_and_stops_when_terminal(self):
        ticket = _make_ticket(
            ola_contain_deadline=timezone.now() + timedelta(minutes=30))
        self.assertEqual(ticket.ola_badge['bucket'], ola_buckets.DUE_1H)

        ticket.status = Ticket.STATUS_APPROVED
        self.assertIsNone(ticket.ola_badge)


class OlaPolicyTest(TestCase):
    """Per-severity OLA targets (triage + contain) applied by Ticket.save().

    Policy: Critical 30m/4h, High 2h/24h, Medium & Low 24h/none (notify-only),
    Unknown mirrors Critical.
    """

    def _make(self, severity):
        base = timezone.now()
        return _make_ticket(severity=severity, incident_datetime=base), base

    def test_critical_targets(self):
        t, base = self._make('Critical')
        self.assertEqual(t.ola_triage_deadline, base + timedelta(minutes=30))
        self.assertEqual(t.ola_contain_deadline, base + timedelta(hours=4))

    def test_high_targets(self):
        t, base = self._make('High')
        self.assertEqual(t.ola_triage_deadline, base + timedelta(hours=2))
        self.assertEqual(t.ola_contain_deadline, base + timedelta(hours=24))

    def test_medium_triage_only_no_contain(self):
        t, base = self._make('Medium')
        self.assertEqual(t.ola_triage_deadline, base + timedelta(hours=24))
        self.assertIsNone(t.ola_contain_deadline)

    def test_low_triage_only_no_contain(self):
        t, base = self._make('Low')
        self.assertEqual(t.ola_triage_deadline, base + timedelta(hours=24))
        self.assertIsNone(t.ola_contain_deadline)

    def test_unknown_mirrors_critical(self):
        t, base = self._make('Unknown')
        self.assertEqual(t.ola_triage_deadline, base + timedelta(minutes=30))
        self.assertEqual(t.ola_contain_deadline, base + timedelta(hours=4))

    def test_triage_breach_vs_contain_breach_are_independent(self):
        base = timezone.now() - timedelta(hours=10)   # long ago
        t = _make_ticket(severity='Critical', incident_datetime=base,
                         status=Ticket.STATUS_NEW)
        # created_at is ~now, well past triage (base+30m) and contain (base+4h).
        self.assertTrue(t.is_ola_triage_breached)   # raised late
        self.assertTrue(t.is_ola_contain_breached)  # active + past contain
        # Notification-only severities never register a contain breach.
        low = _make_ticket(severity='Low', incident_datetime=base,
                           status=Ticket.STATUS_NEW)
        self.assertFalse(low.is_ola_contain_breached)


# ──────────────────────────────────────────────────────────────────────────── #
# Project Incident (Case Bundling) — one incident fanned out to many tickets    #
# ──────────────────────────────────────────────────────────────────────────── #

def _pi_post_data(admin_a, admin_b, **overrides):
    """A valid create_project_incident POST payload with 2 target systems."""
    data = {
        # shared incident facts
        'title': 'Multi-system intrusion via public-facing app',
        'severity': 'High',
        'ncsa_severity': Ticket.NCSA_SEVERITY_SEVERE,
        'log_source': 'Wazuh',
        'issue_type': 'SIEM',
        'detailed_issue': 'Malicious Logic',
        'detailed_issue2': 'C2 Server',
        'issue_description': 'Attacker pivoted across several core systems.',
        'action_required': 'Isolate host and rotate credentials.',
        'action_precautions': 'Preserve volatile evidence before reimaging.',
        'spread_to_others': 'true',
        # target formset management form
        'target-TOTAL_FORMS': '2',
        'target-INITIAL_FORMS': '0',
        'target-MIN_NUM_FORMS': '2',
        'target-MAX_NUM_FORMS': '25',
        # target A
        'target-0-device_name': 'HR Portal',
        'target-0-ip_address': '192.0.2.11',
        'target-0-assigned_admin': str(admin_a.pk),
        'target-0-t1_route': Ticket.T1_ROUTE_ADMIN,
        # target B
        'target-1-device_name': 'AD Server',
        'target-1-ip_address': '192.0.2.12',
        'target-1-assigned_admin': str(admin_b.pk),
        'target-1-t1_route': Ticket.T1_ROUTE_ADMIN,
    }
    data.update(overrides)
    return data


class BundleSuffixHelperTest(TestCase):
    def test_excel_style_labels(self):
        self.assertEqual(bundle_suffix_for_index(0), 'A')
        self.assertEqual(bundle_suffix_for_index(1), 'B')
        self.assertEqual(bundle_suffix_for_index(25), 'Z')
        self.assertEqual(bundle_suffix_for_index(26), 'AA')


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_project_test_media_'))
class ProjectIncidentFanOutTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1      = _make_t1('pi_t1')
        cls.t2      = _make_t2('pi_t2')
        cls.manager = _make_user('pi_manager', UserProfile.ROLE_SOC_MANAGER)
        cls.admin_a = _make_user('pi_admin_a', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.admin_b = _make_user('pi_admin_b', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.owner_b = _make_user('pi_owner_b', UserProfile.ROLE_SYSTEM_OWNER)

    def test_fanout_creates_linked_member_tickets(self):
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        self.assertEqual(resp.status_code, 302)

        project = ProjectIncident.objects.get()
        members = list(project.members)
        self.assertEqual(len(members), 2)

        # Trackable, ordered ids: PI-YYMMDD-NN-A / -B
        self.assertTrue(project.project_code.startswith('PI-'))
        self.assertEqual([m.bundle_suffix for m in members], ['A', 'B'])
        self.assertEqual(members[0].bundle_ref, f'{project.project_code}-A')
        self.assertEqual(members[1].display_id, members[1].ticket_id)

        # Each member routed to its own admin, waiting in the SOC Manager
        # pre-containment review, as Incident.
        self.assertEqual({m.assigned_admin for m in members}, {self.admin_a, self.admin_b})
        for m in members:
            self.assertEqual(m.status, Ticket.STATUS_PENDING_MGR_TRIAGE)
            self.assertEqual(m.t1_route, Ticket.T1_ROUTE_ADMIN)
            self.assertEqual(m.classification, Ticket.CLASSIFICATION_INCIDENT)
            self.assertEqual(m.created_by, self.t1)
            # Shared incident facts copied onto every member.
            self.assertEqual(m.action_required, 'Isolate host and rotate credentials.')
            self.assertEqual(m.issue_description, 'Attacker pivoted across several core systems.')
            self.assertEqual(m.detailed_issue2, 'C2 Server')
        # Per-target facts differ.
        self.assertEqual({m.device_name for m in members}, {'HR Portal', 'AD Server'})

    def test_members_keep_independent_lifecycle(self):
        """Closing one member must not move the others (grouping only)."""
        self.client.login(username='pi_t1', password='testpass123')
        self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        project = ProjectIncident.objects.get()
        first, second = list(project.members)
        self.client.logout()
        self.client.login(username='pi_manager', password='testpass123')
        self.client.post(reverse('project_incident_detail', args=[project.pk]), {
            'action': 'project_mgr_forward',
            'emergency_assessment': 'normal',
            'decision_note': 'Shared assessment complete.',
        })
        first.refresh_from_db()
        _advance_to(first, Ticket.STATUS_APPROVED, self.t1, admin=self.admin_a, t2=self.t2)
        second.refresh_from_db()
        self.assertEqual(second.status, Ticket.STATUS_AWAITING_CONTAINMENT)
        self.assertEqual(project.open_member_count, 1)
        self.assertFalse(project.all_closed)

    def test_project_review_controls_all_member_routes_and_emergency(self):
        self.client.login(username='pi_t1', password='testpass123')
        data = _pi_post_data(
            self.admin_a, self.admin_b,
            **{
                'target-1-t1_route': Ticket.T1_ROUTE_OWNER,
                'target-1-system_owner': str(self.owner_b.pk),
                'evidence_files': SimpleUploadedFile(
                    'timeline.txt', b'project-wide evidence',
                    content_type='text/plain',
                ),
            },
        )
        self.client.post(reverse('create_project_incident'), data)
        project = ProjectIncident.objects.get()
        first, second = list(project.members)

        with self.assertRaises(ValidationError):
            first.transition_to(
                Ticket.STATUS_AWAITING_CONTAINMENT, self.manager, 'bypass attempt',
            )

        self.client.logout()
        self.client.login(username='pi_manager', password='testpass123')
        response = self.client.post(reverse('project_incident_detail', args=[project.pk]), {
            'action': 'project_mgr_forward',
            'emergency_assessment': 'emergency',
            'decision_note': 'One incident, shared urgent response.',
        })
        self.assertEqual(response.status_code, 302)

        project.refresh_from_db()
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertTrue(project.is_emergency)
        self.assertEqual(project.emergency_decided_by, self.manager)
        self.assertEqual(first.status, Ticket.STATUS_AWAITING_CONTAINMENT)
        self.assertEqual(second.status, Ticket.STATUS_AWAITING_OWNER)
        self.assertTrue(first.is_emergency)
        self.assertTrue(second.is_emergency)
        self.assertTrue(ProjectIncidentLog.objects.filter(project=project).exists())
        self.assertEqual(ProjectIncidentAttachment.objects.filter(project=project).count(), 1)

        self.client.logout()
        self.client.login(username='pi_admin_a', password='testpass123')
        attachment = project.attachments.get()
        download = self.client.get(
            reverse('download_project_attachment', args=[attachment.pk]),
        )
        self.assertEqual(download.status_code, 200)
        self.assertIn('attachment', download['Content-Disposition'])

    def test_project_reassessment_changes_all_active_members_together(self):
        self.client.login(username='pi_t1', password='testpass123')
        self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        project = ProjectIncident.objects.get()
        self.client.logout()
        self.client.login(username='pi_manager', password='testpass123')
        self.client.post(reverse('project_incident_detail', args=[project.pk]), {
            'action': 'project_mgr_forward',
            'emergency_assessment': 'emergency',
            'decision_note': 'Initial group assessment.',
        })
        response = self.client.post(reverse('project_incident_detail', args=[project.pk]), {
            'action': 'project_reassess_emergency',
            'emergency_value': '0',
            'emergency_reason': 'Scope reduced after validation.',
        })
        self.assertEqual(response.status_code, 302)
        project.refresh_from_db()
        self.assertFalse(project.is_emergency)
        self.assertFalse(project.members.filter(is_emergency=True).exists())

        member = project.members.first()
        with self.assertRaises(ValidationError):
            member.reassess_emergency(True, self.manager, 'individual exception')

    def test_fewer_than_two_targets_is_rejected(self):
        self.client.login(username='pi_t1', password='testpass123')
        data = _pi_post_data(self.admin_a, self.admin_b)
        # Blank out the second target row → only one valid target remains.
        data['target-1-device_name'] = ''
        data['target-1-ip_address'] = ''
        data['target-1-assigned_admin'] = ''
        resp = self.client.post(reverse('create_project_incident'), data)
        self.assertEqual(resp.status_code, 200)  # re-rendered with errors
        self.assertFalse(ProjectIncident.objects.exists())
        self.assertFalse(Ticket.objects.exists())

    def test_non_tier1_cannot_open_fanout_page(self):
        self.client.login(username='pi_t2', password='testpass123')
        resp = self.client.get(reverse('create_project_incident'))
        self.assertEqual(resp.status_code, 302)  # Tier 1 only

    def test_detail_page_lists_members_for_soc(self):
        self.client.login(username='pi_t1', password='testpass123')
        self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        project = ProjectIncident.objects.get()
        resp = self.client.get(reverse('project_incident_detail', args=[project.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, project.project_code)
        self.assertContains(resp, 'HR Portal')
        self.assertContains(resp, 'AD Server')

    def test_detail_page_scopes_members_to_system_admin(self):
        """A system admin only sees the member ticket assigned to them."""
        self.client.login(username='pi_t1', password='testpass123')
        self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        project = ProjectIncident.objects.get()
        self.client.logout()
        self.client.login(username='pi_admin_a', password='testpass123')
        resp = self.client.get(reverse('project_incident_detail', args=[project.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'HR Portal')      # assigned to admin_a
        self.assertNotContains(resp, 'AD Server')   # assigned to admin_b

    def test_ticket_lists_link_project_members_to_their_overview(self):
        self.client.login(username='pi_t1', password='testpass123')
        self.client.post(
            reverse('create_project_incident'),
            _pi_post_data(self.admin_a, self.admin_b),
        )
        project = ProjectIncident.objects.get()
        member = project.members.first()
        overview_url = reverse('project_incident_detail', args=[project.pk])

        self.client.logout()
        self.client.login(username='pi_manager', password='testpass123')
        active = self.client.get(reverse('ticket_list'))
        self.assertContains(active, f'Project · {member.bundle_ref}')
        self.assertContains(active, overview_url)

        member.status = Ticket.STATUS_APPROVED
        member.save(update_fields=('status', 'updated_at'))
        history = self.client.get(reverse('ticket_history'), {'all_time': '1'})
        self.assertContains(history, f'Project · {member.bundle_ref}')
        self.assertContains(history, overview_url)

    # ── Origin from a Wazuh alert (analyst-initiated, pre-filled) ──────── #

    def _claimed_alert(self, claimer):
        return WazuhAlert.objects.create(
            opensearch_id=f'os-pi-{claimer.username}-{timezone.now().timestamp()}',
            timestamp=timezone.now(), rule_level=13,
            rule_description='Coordinated intrusion across core systems',
            agent_name='DC-01',
            triage_status=WazuhAlert.TRIAGE_TRIAGING,
            claimed_by=claimer, claimed_at=timezone.now(),
        )

    def test_triage_action_routes_to_project_incident_keeping_claim(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.post(reverse('triage_action'), {
            'alert_id': alert.pk, 'action': 'create_project_incident',
            'note': 'multi-system', 'category': WazuhAlert.CATEGORY_MALWARE,
        })
        self.assertEqual(resp.status_code, 302)
        self.assertIn(reverse('create_project_incident'), resp.url)
        self.assertIn(f'wazuh_alert={alert.pk}', resp.url)
        # Alert stays claimed + triaging until the fan-out form is saved.
        alert.refresh_from_db()
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRIAGING)
        self.assertEqual(alert.claimed_by, self.t1)

    def test_get_prefills_from_alert(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.get(
            reverse('create_project_incident'), {'wazuh_alert': alert.pk},
        )
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Coordinated intrusion across core systems')
        self.assertContains(resp, f'name="wazuh_alert" value="{alert.pk}"')

    def test_create_form_embeds_shared_containment_guidance(self):
        self.client.login(username='pi_t1', password='testpass123')
        response = self.client.get(reverse('create_project_incident'))

        self.assertContains(response, 'id="insert-guidance-btn"')
        match = re.search(
            r'<script id="threat-guidance-data"[^>]*>(.*?)</script>',
            response.content.decode(), re.S,
        )
        self.assertIsNotNone(match)
        guidance = json.loads(match.group(1))
        self.assertIn('Malicious Logic', guidance)
        self.assertTrue(guidance['Malicious Logic']['action_required'].strip())
        self.assertTrue(guidance['Malicious Logic']['action_precautions'].strip())
        self.assertContains(response, 'id="guidance-note-data"')

    def test_get_rejects_alert_not_claimed_by_user(self):
        alert = self._claimed_alert(self.t2)  # claimed by someone else
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.get(
            reverse('create_project_incident'), {'wazuh_alert': alert.pk},
        )
        self.assertEqual(resp.status_code, 302)  # bounced back to the queue

    def test_fanout_from_alert_links_bundle_and_consumes_alert(self):
        alert = self._claimed_alert(self.t1)
        self.client.login(username='pi_t1', password='testpass123')
        data = _pi_post_data(self.admin_a, self.admin_b, wazuh_alert=str(alert.pk))
        resp = self.client.post(reverse('create_project_incident'), data)
        self.assertEqual(resp.status_code, 302)

        project = ProjectIncident.objects.get()
        alert.refresh_from_db()
        # Alert points at the whole bundle (option B), not a single ticket.
        self.assertEqual(alert.project_incident, project)
        self.assertFalse(hasattr(alert, 'ticket'))
        # Alert consumed → leaves the triage queue.
        self.assertEqual(alert.triage_status, WazuhAlert.TRIAGE_TRUE_POSITIVE)
        self.assertEqual(alert.triaged_by, self.t1)
        self.assertIsNone(alert.claimed_by)
        # Response time stamped on the first member.
        first = project.members.first()
        self.assertIsNotNone(first.alert_conversion_duration)
        # The bundle exposes its origin alert via the reverse relation.
        self.assertEqual(project.source_alerts.first(), alert)

    # ── Origin from a Manual Triage record ────────────────────────────── #

    def _claimed_triage(self, claimer):
        return TriageRecord.objects.create(
            source=TriageRecord.SOURCE_EMAIL,
            source_reference='REP-2026-777',
            analyst=claimer,
            alert_description='User reported ransomware note across two shared drives',
            claimed_by=claimer, claimed_at=timezone.now(),
        )

    def test_manual_triage_get_prefills(self):
        triage = self._claimed_triage(self.t1)
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.get(
            reverse('create_project_incident'), {'triage_id': triage.pk},
        )
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'ransomware note across two shared drives')
        self.assertContains(resp, f'name="triage_id" value="{triage.pk}"')

    def test_manual_triage_get_rejects_record_of_other_user(self):
        triage = self._claimed_triage(self.t2)  # claimed by someone else
        self.client.login(username='pi_t1', password='testpass123')
        resp = self.client.get(
            reverse('create_project_incident'), {'triage_id': triage.pk},
        )
        self.assertEqual(resp.status_code, 302)  # bounced to manual triage list

    def test_manual_triage_fanout_links_bundle_and_consumes_record(self):
        triage = self._claimed_triage(self.t1)
        self.client.login(username='pi_t1', password='testpass123')
        data = _pi_post_data(self.admin_a, self.admin_b, triage_id=str(triage.pk))
        resp = self.client.post(reverse('create_project_incident'), data)
        self.assertEqual(resp.status_code, 302)

        project = ProjectIncident.objects.get()
        triage.refresh_from_db()
        # Record points at the whole bundle, not a single ticket (option B).
        self.assertEqual(triage.project_incident, project)
        self.assertIsNone(triage.ticket_id)
        # Marked TP + unclaimed → leaves the manual triage queue.
        self.assertEqual(triage.decision, TriageRecord.DECISION_TP)
        self.assertIsNone(triage.claimed_by)
        self.assertEqual(project.source_triages.first(), triage)


# ──────────────────────────────────────────────────────────────────────────── #
# 17. Threat guidance (แทรกแนวทางมาตรฐาน) tests                                 #
# ──────────────────────────────────────────────────────────────────────────── #

class ThreatGuidanceTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('guidance_t1')

    @staticmethod
    def _embedded_json(response, script_id):
        match = re.search(
            rf'<script id="{script_id}"[^>]*>(.*?)</script>',
            response.content.decode(), re.S,
        )
        assert match, f'json_script {script_id!r} not found in page'
        return json.loads(match.group(1))

    def test_seed_covers_every_form_category(self):
        # Migration 0041 must leave one active row per clean threat category.
        for category in Ticket.DETAILED_ISSUE_HIERARCHY:
            guidance = ThreatGuidance.objects.get(detailed_issue=category)
            self.assertTrue(guidance.is_active)
            self.assertTrue(guidance.action_required.strip(), category)
            self.assertTrue(guidance.action_precautions.strip(), category)

    def test_create_form_embeds_guidance_data_and_button(self):
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'))

        self.assertContains(response, 'id="insert-guidance-btn"')
        data = self._embedded_json(response, 'threat-guidance-data')
        self.assertIn('Malicious Logic', data)
        self.assertIn(
            'Isolate เครื่องที่ติดมัลแวร์ออกจากเครือข่ายทันที',
            data['Malicious Logic']['action_required'],
        )
        self.assertIn(
            'ห้ามลบไฟล์ต้องสงสัยก่อนเก็บหลักฐาน',
            data['Malicious Logic']['action_precautions'],
        )
        note = self._embedded_json(response, 'guidance-note-data')
        self.assertIn('หมายเหตุ: สามารถประสานงานเรื่องเหตุละเมิดได้ดังนี้', note)
        self.assertIn('02-574-8209-10', note)

    def test_inactive_guidance_left_out_of_form_payload(self):
        ThreatGuidance.objects.filter(detailed_issue='Malicious Logic').update(is_active=False)
        self.client.force_login(self.t1)
        response = self.client.get(reverse('create_ticket'))
        data = self._embedded_json(response, 'threat-guidance-data')
        self.assertNotIn('Malicious Logic', data)
        self.assertIn('User Intrusion', data)  # others unaffected

    def test_containment_form_has_no_leftover_test_markup(self):
        # Regression: WIP markup (ทดสอบ1/fruit checkbox) must not resurface in
        # the System Admin Investigation form on the ticket detail page.
        admin_user = _make_user('guidance_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        ticket = _make_ticket(
            assigned_admin=admin_user,
            status=Ticket.STATUS_AWAITING_CONTAINMENT,
            classification=Ticket.CLASSIFICATION_INCIDENT,
        )
        self.client.force_login(admin_user)
        response = self.client.get(reverse('ticket_detail', args=[ticket.pk]))
        self.assertNotContains(response, 'ทดสอบ1')
        self.assertNotContains(response, 'name="fruit"')


# ──────────────────────────────────────────────────────────────────────────── #
# 18. Containment checklist tests                                               #
# ──────────────────────────────────────────────────────────────────────────── #

class ContainmentChecklistTest(TestCase):
    ACTION = '1) Isolate เครื่อง\n2) Block IoC\n3) Patch\nหมายเหตุ: ประสาน 02-574-8209-10'

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('cl_t1')
        cls.admin = _make_user('cl_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _awaiting_ticket(self, **kw):
        kw.setdefault('action_required', self.ACTION)
        kw.setdefault('status', Ticket.STATUS_AWAITING_CONTAINMENT)
        return _make_ticket(
            created_by=self.t1, assigned_admin=self.admin,
            classification=Ticket.CLASSIFICATION_INCIDENT, **kw,
        )

    def test_parse_items_vs_trailing(self):
        items, trailing = Ticket.parse_checklist_items('1) A\n- B\n• C\nnote line\n\n2. D')
        self.assertEqual(items, ['1) A', '- B', '• C', '2. D'])
        self.assertEqual(trailing, ['note line'])

    def test_display_restores_saved_states(self):
        t = self._awaiting_ticket()
        t.containment_checklist = [
            {'text': '1) Isolate เครื่อง', 'done': True},
            {'text': '2) Block IoC', 'done': False},
        ]
        items, trailing = t.containment_checklist_display()
        done = {i['text']: i['done'] for i in items}
        self.assertTrue(done['1) Isolate เครื่อง'])
        self.assertFalse(done['2) Block IoC'])
        self.assertFalse(done['3) Patch'])  # never saved → unchecked
        self.assertIn('หมายเหตุ', trailing)

    def test_rejection_loop_keeps_matching_ticks(self):
        t = self._awaiting_ticket()
        t.containment_checklist = [
            {'text': '1) Isolate เครื่อง', 'done': True},
            {'text': '2) Block IoC', 'done': True},
            {'text': '3) Patch', 'done': False},
        ]
        t.save()
        # Tier 1 rewords item 2 during the rejection loop.
        t.action_required = '1) Isolate เครื่อง\n2) Block IoC และ Domain\n3) Patch'
        done = {i['text']: i['done'] for i in t.containment_checklist_display()[0]}
        self.assertTrue(done['1) Isolate เครื่อง'])          # unchanged text → kept
        self.assertFalse(done['2) Block IoC และ Domain'])    # reworded → unticked
        self.assertFalse(done['3) Patch'])

    def test_submit_saves_checklist_and_logs_progress(self):
        t = self._awaiting_ticket()
        self.client.force_login(self.admin)
        self.client.post(reverse('ticket_detail', args=[t.pk]), {
            'action': 'containment',
            'containment_report': 'Contained and cleaned.',
            'checklist_done': ['0', '2'],  # items 1 and 3 done
        })
        t.refresh_from_db()
        self.assertEqual(t.status, Ticket.STATUS_CONTAINMENT_REPORTED)
        done = {c['text']: c['done'] for c in t.containment_checklist}
        self.assertEqual(done, {
            '1) Isolate เครื่อง': True,
            '2) Block IoC': False,
            '3) Patch': True,
        })
        last_note = TicketLog.objects.filter(ticket=t).order_by('-created_at').first().note
        self.assertIn('ดำเนินการแล้ว 2/3 รายการ', last_note)

    def test_admin_card_renders_editable_checklist(self):
        t = self._awaiting_ticket()
        self.client.force_login(self.admin)
        response = self.client.get(reverse('ticket_detail', args=[t.pk]))
        self.assertContains(response, 'name="checklist_done"')
        self.assertContains(response, '1) Isolate เครื่อง')

    def test_detail_shows_readonly_checklist_after_submit(self):
        t = self._awaiting_ticket(status=Ticket.STATUS_CONTAINMENT_REPORTED)
        t.containment_checklist = [
            {'text': '1) Isolate เครื่อง', 'done': True},
            {'text': '2) Block IoC', 'done': False},
            {'text': '3) Patch', 'done': False},
        ]
        t.save()
        self.client.force_login(self.t1)  # reviewer, not the admin
        response = self.client.get(reverse('ticket_detail', args=[t.pk]))
        self.assertNotContains(response, 'name="checklist_done"')  # read-only
        self.assertContains(response, '☑')  # a done item is ticked
        self.assertContains(response, '1) Isolate เครื่อง')

    def test_report_docx_renders_checklist_with_dejavu_ballots(self):
        t = self._awaiting_ticket()
        t.containment_checklist = [
            {'text': '1) Isolate เครื่อง', 'done': True},
            {'text': '2) Block IoC', 'done': False},
            {'text': '3) Patch', 'done': False},
        ]
        t.save()
        report = generate_ticket_report(t.pk)
        text = _docx_text(report.content)
        self.assertIn('☑ 1) Isolate เครื่อง', text)
        self.assertIn('☐ 2) Block IoC', text)
        # Ballot glyphs must sit in DejaVu Sans runs (TH Sarabun has no such
        # glyph) so they render instead of tofu.
        doc = Document(BytesIO(report.content))
        ballot_fonts = {
            r.font.name
            for p in _iter_paragraphs(doc) for r in p.runs
            if r.text in ('☑', '☐')
        }
        self.assertEqual(ballot_fonts, {'DejaVu Sans'})

    def test_preview_shows_containment_checklist(self):
        t = self._awaiting_ticket()
        t.containment_checklist = [
            {'text': '1) Isolate เครื่อง', 'done': True},
            {'text': '2) Block IoC', 'done': False},
            {'text': '3) Patch', 'done': False},
        ]
        t.save()
        self.client.force_login(self.t1)
        response = self.client.get(reverse('ticket_report_preview', args=[t.pk]))
        self.assertContains(response, '&#9745;</span>&#160;1) Isolate เครื่อง')
        self.assertContains(response, '&#9744;</span>&#160;2) Block IoC')

    def test_no_items_falls_back_to_plain_text(self):
        # action_required with no numbered/bulleted lines → no checklist.
        t = self._awaiting_ticket(action_required='ดำเนินการตามความเหมาะสม')
        items, trailing = t.containment_checklist_display()
        self.assertEqual(items, [])
        self.assertEqual(trailing, 'ดำเนินการตามความเหมาะสม')
        self.assertIsNone(__import__('apps.incidents.reports', fromlist=['_containment_checklist_row'])._containment_checklist_row(t))


# ──────────────────────────────────────────────────────────────────────────── #
# Response-team requests (Forensic / Red Team) — model + routing + gating       #
# ──────────────────────────────────────────────────────────────────────────── #

class UserDropdownLabelTest(TestCase):
    """User-selection dropdowns show a person's name, not their login name."""

    def test_all_user_dropdowns_prefer_full_name_with_username_fallback(self):
        named_user = User.objects.create_user(
            username='named-login', first_name='Ada', last_name='Lovelace',
        )
        unnamed_user = User.objects.create_user(username='login-only')
        fields = [
            TicketForm().fields['assigned_admin'],
            ProjectIncidentTargetForm().fields['assigned_admin'],
            AdminAssignmentForm().fields['assigned_admin'],
            # SubtaskForm has no user dropdown — legacy subtasks are unassigned
            # working notes (see test_legacy_subtask_form_has_no_assignee_field).
            ResponseRequestForm().fields['assigned_to'],
        ]

        for field in fields:
            self.assertEqual(field.label_from_instance(named_user), 'Ada Lovelace')
            self.assertEqual(field.label_from_instance(unnamed_user), 'login-only')


class ResponseRequestRoutingTest(TestCase):
    """Type → role routing and eligible-assignee resolution."""

    @classmethod
    def setUpTestData(cls):
        cls.forensic = _make_forensic('rt_forensic')
        cls.redteam  = _make_redteam_manager('rt_redteam')

    def test_va_pt_and_infra_sec_route_to_redteam_manager(self):
        self.assertEqual(
            TicketSubtask.role_for_type(TicketSubtask.TYPE_VA_PT),
            UserProfile.ROLE_REDTEAM_MANAGER,
        )
        self.assertEqual(
            TicketSubtask.role_for_type(TicketSubtask.TYPE_INFRA_SEC),
            UserProfile.ROLE_REDTEAM_MANAGER,
        )

    def test_legacy_subtask_form_has_no_assignee_field(self):
        # A legacy subtask notifies nobody, gates nothing and never reaches the
        # report, so an assignee could not summon anyone — it only looked like it
        # did. Removing the field also closes the old response-team leak (an
        # ordinary subtask assigned to a responder would have exposed the whole
        # ticket through visible_to) without relying on the picker's filtering.
        # Work that must actually reach a person goes through ResponseRequestForm.
        self.assertNotIn('assigned_to', SubtaskForm().fields)
        self.assertEqual(
            list(SubtaskForm().fields),
            ['subtask_type', 'title', 'description'],
        )

    def test_response_request_form_choices_derive_from_model(self):
        # DRY: form response-type choices must exactly mirror the model's
        # RESPONSE_TYPES (single source of truth), not a hand-maintained copy.
        form_codes = {code for code, _ in ResponseRequestForm.RESPONSE_TYPE_CHOICES}
        self.assertEqual(form_codes, set(TicketSubtask.RESPONSE_TYPES))

    def test_forensic_rca_routes_to_forensic(self):
        self.assertEqual(
            TicketSubtask.role_for_type(TicketSubtask.TYPE_FORENSIC_RCA),
            UserProfile.ROLE_FORENSIC,
        )

    def test_legacy_types_have_no_route(self):
        self.assertIsNone(TicketSubtask.role_for_type(TicketSubtask.TYPE_INVESTIGATION))
        self.assertIsNone(TicketSubtask.role_for_type(TicketSubtask.TYPE_COUNTERMEASURE))

    def test_eligible_assignees_filters_by_role(self):
        self.assertEqual(
            list(TicketSubtask.eligible_assignees(TicketSubtask.TYPE_FORENSIC_RCA)),
            [self.forensic],
        )
        self.assertEqual(
            list(TicketSubtask.eligible_assignees(TicketSubtask.TYPE_INFRA_SEC)),
            [self.redteam],
        )

    def test_eligible_assignees_excludes_inactive(self):
        self.forensic.is_active = False
        self.forensic.save(update_fields=['is_active'])
        self.assertEqual(
            list(TicketSubtask.eligible_assignees(TicketSubtask.TYPE_FORENSIC_RCA)), [],
        )

    def test_eligible_assignees_empty_for_legacy_type(self):
        self.assertEqual(
            list(TicketSubtask.eligible_assignees(TicketSubtask.TYPE_INVESTIGATION)), [],
        )

    def test_is_response_request_flag(self):
        t = _make_ticket()
        resp = TicketSubtask.objects.create(
            ticket=t, subtask_type=TicketSubtask.TYPE_VA_PT, title='x',
        )
        legacy = TicketSubtask.objects.create(
            ticket=t, subtask_type=TicketSubtask.TYPE_INVESTIGATION, title='y',
        )
        self.assertTrue(resp.is_response_request)
        self.assertFalse(legacy.is_response_request)

    def test_role_props(self):
        self.assertTrue(self.forensic.profile.is_forensic)
        self.assertTrue(self.forensic.profile.is_response_team)
        self.assertFalse(self.forensic.profile.is_soc)
        self.assertTrue(self.redteam.profile.is_redteam_manager)
        self.assertTrue(self.redteam.profile.is_response_team)
        self.assertFalse(self.redteam.profile.is_soc)


class ResponseRequestAssignmentValidationTest(TestCase):
    """K9 write-time half: TicketSubtask.clean() rejects a misrouted assignment.

    Covers the paths that go through full_clean() — ModelForms and the Django
    admin. objects.create() deliberately skips this; that gap is covered by the
    read-time tests in TicketVisibilityQuerysetTest.
    """

    @classmethod
    def setUpTestData(cls):
        cls.forensic = _make_forensic('av_forensic')
        cls.redteam  = _make_redteam_manager('av_redteam')
        cls.soc      = _make_t1('av_soc')
        cls.ticket   = _make_ticket()

    def _subtask(self, subtask_type, assigned_to):
        return TicketSubtask(
            ticket=self.ticket, subtask_type=subtask_type,
            title='x', assigned_to=assigned_to,
        )

    def test_forensic_rca_assigned_to_redteam_is_rejected(self):
        with self.assertRaises(ValidationError) as ctx:
            self._subtask(TicketSubtask.TYPE_FORENSIC_RCA, self.redteam).full_clean()
        self.assertIn('assigned_to', ctx.exception.error_dict)

    def test_va_pt_assigned_to_forensic_is_rejected(self):
        with self.assertRaises(ValidationError) as ctx:
            self._subtask(TicketSubtask.TYPE_VA_PT, self.forensic).full_clean()
        self.assertIn('assigned_to', ctx.exception.error_dict)

    def test_response_request_assigned_to_non_response_user_is_rejected(self):
        with self.assertRaises(ValidationError) as ctx:
            self._subtask(TicketSubtask.TYPE_INFRA_SEC, self.soc).full_clean()
        self.assertIn('assigned_to', ctx.exception.error_dict)

    def test_profileless_assignee_is_rejected(self):
        # Fail closed, matching visible_to()'s treatment of the profile-less
        # account that exists between admin creation and profile setup.
        stray = User.objects.create_user(username='av_stray', password='testpass123')
        with self.assertRaises(ValidationError) as ctx:
            self._subtask(TicketSubtask.TYPE_VA_PT, stray).full_clean()
        self.assertIn('assigned_to', ctx.exception.error_dict)

    def test_correctly_routed_assignments_validate(self):
        self._subtask(TicketSubtask.TYPE_FORENSIC_RCA, self.forensic).full_clean()
        self._subtask(TicketSubtask.TYPE_VA_PT, self.redteam).full_clean()
        self._subtask(TicketSubtask.TYPE_INFRA_SEC, self.redteam).full_clean()

    def test_unassigned_response_request_validates(self):
        # The spawn flow legitimately leaves assigned_to blank when a picker is
        # needed; clean() must not turn that into an error.
        self._subtask(TicketSubtask.TYPE_VA_PT, None).full_clean()

    def test_legacy_subtask_assignment_is_unconstrained(self):
        # Investigation/Countermeasure route nowhere and keep their freedom;
        # SubtaskForm is what keeps response-team users out of them.
        self._subtask(TicketSubtask.TYPE_INVESTIGATION, self.soc).full_clean()
        self._subtask(TicketSubtask.TYPE_COUNTERMEASURE, self.soc).full_clean()


class ResponseRequestQueueScopingTest(TestCase):
    """K9 read-time half, at the HTTP level."""

    @classmethod
    def setUpTestData(cls):
        cls.forensic = _make_forensic('qs_forensic')
        cls.redteam  = _make_redteam_manager('qs_redteam')
        cls.mgr      = _make_user('qs_mgr', UserProfile.ROLE_SOC_MANAGER)

        cls.ticket_rt  = _make_ticket()
        cls.ticket_bad = _make_ticket()

        cls.va_pt = TicketSubtask.objects.create(
            ticket=cls.ticket_rt, subtask_type=TicketSubtask.TYPE_VA_PT,
            title='Pentest the DMZ', assigned_to=cls.redteam,
        )
        # Misrouted on purpose (objects.create bypasses clean()) — this is the
        # row a seed or migration could leave behind.
        cls.misrouted = TicketSubtask.objects.create(
            ticket=cls.ticket_bad, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='Memory dump RCA', assigned_to=cls.redteam,
        )

    def test_queue_hides_a_request_of_another_teams_type(self):
        self.client.login(username='qs_redteam', password='testpass123')
        body = self.client.get(reverse('response_request_queue')).content.decode()
        self.assertIn('Pentest the DMZ', body)
        self.assertNotIn('Memory dump RCA', body)

    def test_misrouted_request_does_not_unlock_the_ticket(self):
        self.client.login(username='qs_redteam', password='testpass123')
        self.assertEqual(
            self.client.get(
                reverse('ticket_detail', kwargs={'pk': self.ticket_bad.pk})
            ).status_code, 404,
        )

    def test_correctly_routed_request_still_grants_access(self):
        self.client.login(username='qs_redteam', password='testpass123')
        self.assertEqual(
            self.client.get(
                reverse('ticket_detail', kwargs={'pk': self.ticket_rt.pk})
            ).status_code, 200,
        )

    def test_soc_overview_still_sees_every_type(self):
        # The all-team overview is unfiltered by design — scoping applies only
        # to a response-team member's own queue.
        self.client.login(username='qs_mgr', password='testpass123')
        body = self.client.get(reverse('response_request_queue')).content.decode()
        self.assertIn('Pentest the DMZ', body)
        self.assertIn('Memory dump RCA', body)


class ResponseRequestApprovalGateTest(TestCase):
    """An open response request blocks every path into APPROVED, but not the
    Event-close path."""

    @classmethod
    def setUpTestData(cls):
        cls.t1    = _make_t1('rg_t1')
        cls.t2    = _make_t2('rg_t2')
        cls.admin = _make_user('rg_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.mgr   = _make_user('rg_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.forensic = _make_forensic('rg_forensic')

    def _incident(self, **kwargs):
        kwargs.setdefault('assigned_admin', self.admin)
        return _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_INCIDENT, **kwargs
        )

    def _open_request(self, ticket, status=TicketSubtask.STATUS_OPEN):
        return TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=self.forensic, status=status,
        )

    def test_has_open_response_requests_property(self):
        t = self._incident()
        self.assertFalse(t.has_open_response_requests)
        st = self._open_request(t)
        self.assertTrue(t.has_open_response_requests)
        st.status = TicketSubtask.STATUS_DONE
        st.save(update_fields=['status'])
        self.assertFalse(t.has_open_response_requests)

    def test_legacy_subtask_does_not_block(self):
        t = self._incident()
        TicketSubtask.objects.create(
            ticket=t, subtask_type=TicketSubtask.TYPE_INVESTIGATION, title='dig',
            status=TicketSubtask.STATUS_OPEN,
        )
        self.assertFalse(t.has_open_response_requests)

    def test_open_request_blocks_manager_approval(self):
        t = self._incident(is_emergency=True)
        _advance_to(t, Ticket.STATUS_PENDING_MANAGER, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        self._open_request(t)
        t.refresh_from_db()
        self.assertFalse(t.can_transition_to(Ticket.STATUS_APPROVED))
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'close')

    def test_open_request_blocks_tier2_direct_close(self):
        # Non-emergency incident closes via Tier 2 (CONTAINMENT_REPORTED→APPROVED)
        # — the gate must catch this path too, not just manager approval.
        t = self._incident()
        _advance_to(t, Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        self._open_request(t)
        t.refresh_from_db()
        self.assertFalse(t.can_transition_to(Ticket.STATUS_APPROVED))
        with self.assertRaises(ValidationError):
            t.transition_to(Ticket.STATUS_APPROVED, self.t2, 'close')

    def test_done_request_allows_approval(self):
        t = self._incident(is_emergency=True)
        _advance_to(t, Ticket.STATUS_PENDING_MANAGER, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        self._open_request(t, status=TicketSubtask.STATUS_DONE)
        t.refresh_from_db()
        t.transition_to(Ticket.STATUS_APPROVED, self.mgr, 'close')
        self.assertEqual(t.status, Ticket.STATUS_APPROVED)

    def test_open_request_does_not_block_event_close(self):
        # Reclassify-to-Event mid-containment still closes; the open forensic
        # request simply outlives the closed ticket.
        t = self._incident()
        _advance_to(t, Ticket.STATUS_CONTAINMENT_REPORTED, self.t1, self.admin,
                    mgr=self.mgr, t2=self.t2)
        self._open_request(t)
        t.refresh_from_db()
        t.classification = Ticket.CLASSIFICATION_EVENT
        t.save(update_fields=['classification'])
        t.transition_to(Ticket.STATUS_CLOSED_EVENT, self.t2, 'reclassified as event')
        self.assertEqual(t.status, Ticket.STATUS_CLOSED_EVENT)


class ResponseRequestNotificationTest(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('rn_t1')
        cls.mgr = _make_user('rn_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.mgr.email = 'mgr@example.com'
        cls.mgr.save(update_fields=['email'])
        cls.forensic = _make_forensic('rn_forensic')
        cls.forensic.email = 'forensic@example.com'
        cls.forensic.save(update_fields=['email'])
        cls.ticket = _make_ticket(created_by=cls.t1)

    def setUp(self):
        mail.outbox = []

    def _request(self, assigned_to, **kwargs):
        return TicketSubtask.objects.create(
            ticket=self.ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='Collect memory image', assigned_to=assigned_to,
            created_by=self.mgr, **kwargs
        )

    def test_created_emails_assigned_responder(self):
        st = self._request(self.forensic)
        self.assertTrue(notify_response_request_created(st))
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, ['forensic@example.com'])
        self.assertIn(self.ticket.ticket_id, mail.outbox[0].subject)

    def test_created_skips_when_no_assignee(self):
        st = self._request(None)
        self.assertFalse(notify_response_request_created(st))
        self.assertEqual(len(mail.outbox), 0)

    def test_completed_emails_managers(self):
        st = self._request(self.forensic, status=TicketSubtask.STATUS_DONE,
                           result_notes='Root cause: phishing.')
        self.assertTrue(notify_response_request_completed(st))
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, ['mgr@example.com'])


# ── Template/markup regression tests (2026-07 UX audit) ──────────────────── #

class TemplateMarkupRegressionTest(TestCase):
    """
    Guards template-level fixes from the 2026-07 visual audit that are easy to
    silently reintroduce when editing markup.
    """

    @staticmethod
    def _source(template_name):
        from django.template.loader import get_template
        return get_template(template_name).template.source

    def test_no_template_comment_spans_multiple_lines(self):
        """Django strips {# #} line-by-line, so a comment whose closing #} is on
        a later line is not a comment at all — it renders to the page as literal
        text (and any markup inside it is parsed by the browser). Multi-line
        notes must use {% comment %}...{% endcomment %}.
        """
        import pathlib
        from django.conf import settings

        offenders = []
        root = pathlib.Path(settings.BASE_DIR) / 'templates'
        for path in sorted(root.rglob('*.html')):
            for lineno, line in enumerate(
                path.read_text(encoding='utf-8').splitlines(), start=1
            ):
                idx = line.find('{#')
                while idx != -1:
                    if '#}' not in line[idx + 2:]:
                        offenders.append(f'{path.relative_to(root)}:{lineno}')
                        break
                    idx = line.find('{#', line.find('#}', idx) + 2)

        self.assertEqual(offenders, [], f'unterminated {{# #}} comments: {offenders}')

    def test_t1_route_script_carries_csp_nonce(self):
        """
        Every inline <script> needs the nonce — script-src has no
        'unsafe-inline', so a bare tag is silently dropped and the Tier-1
        route toggle stops hiding the System Admin picker.
        """
        source = self._source('incidents/ticket_detail.html')
        bare = re.findall(r'<script(?![^>]*\bnonce=)[^>]*>', source)
        self.assertEqual(bare, [], f'inline script(s) without a CSP nonce: {bare}')

    def test_status_badge_can_suppress_duplicate_emergency_flag(self):
        """The detail hero renders its own light EMERGENCY badge."""
        partial = self._source('incidents/_status_badge.html')
        self.assertIn('not hide_emergency', partial)
        detail = self._source('incidents/ticket_detail.html')
        self.assertIn('hide_emergency=True', detail)

    def test_history_truncation_uses_inner_block(self):
        """
        text-truncate needs a block box: on a <td> under auto table layout the
        max-width is ignored and the full description renders.
        """
        source = self._source('incidents/ticket_history.html')
        self.assertNotIn('<td class="text-truncate"', source)
        self.assertIn('<div class="text-truncate"', source)

    def test_page_titles_use_em_dash_separator(self):
        templates = [
            'incidents/ticket_detail.html', 'incidents/my_queue.html',
            'incidents/triage_form.html', 'wazuh_ingest/triage_queue.html',
            'wazuh_ingest/escalation_queue.html', 'dashboard/dashboard.html',
        ]
        for name in templates:
            with self.subTest(template=name):
                title = re.search(
                    r'{% block title %}(.*?){% endblock %}', self._source(name))
                self.assertIsNotNone(title)
                self.assertNotIn(' - ', title.group(1))
                self.assertIn(' — ', title.group(1))

    def test_ticket_form_labels_are_associated(self):
        """
        Every .form-label either points at a control (for=) or is a group
        heading wired up with aria-labelledby — no orphan labels.
        """
        source = self._source('incidents/ticket_form.html')
        orphans = re.findall(
            r'<label class="form-label[^"]*"(?![^>]*\bfor=)[^>]*>', source)
        self.assertEqual(orphans, [], f'labels without for=: {orphans}')

    def test_ticket_form_radio_groups_are_labelled(self):
        source = self._source('incidents/ticket_form.html')
        for group_id in ['severity-group', 'ncsa-severity-group',
                         'spread-group', 'asset-type-group']:
            with self.subTest(group=group_id):
                self.assertIn(f'aria-labelledby="{group_id}-label"', source)
                self.assertIn(f'id="{group_id}-label"', source)

    def test_project_incident_form_has_local_draft_autosave(self):
        source = self._source('incidents/project_incident_form.html')
        self.assertIn('project_incident_form_draft', source)
        self.assertIn('targetCount:', source)
        self.assertIn('project:add-target', source)
        self.assertIn("field.type !== 'file'", source)
        self.assertIn("localStorage.removeItem(DRAFT_KEY)", source)


class LoginRedirectsAuthenticatedUserTest(TestCase):
    """
    An authenticated visitor hitting /login/ used to get the login form
    rendered inside the full app shell (sidebar, nav badges, user pill).
    """

    def test_authenticated_user_is_redirected_away_from_login(self):
        user = User.objects.create_user(username='already_in', password='pw')
        UserProfile.objects.create(user=user, role=UserProfile.ROLE_SOC_STAFF,
                                   tier=UserProfile.TIER_T1)
        self.client.force_login(user)
        resp = self.client.get(reverse('login'))
        self.assertEqual(resp.status_code, 302)

    def test_anonymous_user_still_gets_the_form(self):
        resp = self.client.get(reverse('login'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'name="username"')


class T1RouteToggleCspTest(TestCase):
    """
    The Tier-1 route toggle is an inline <script>. script-src has no
    'unsafe-inline', so without a nonce the browser drops it and the System
    Admin picker stays visible and required even on the Direct-to-Owner route.
    """

    def setUp(self):
        self.t1 = _make_t1('t1_route_user')
        self.ticket = _make_ticket(
            status=Ticket.STATUS_T1_REVIEW,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            created_by=self.t1,
        )
        self.client.force_login(self.t1)

    def test_route_toggle_script_nonce_matches_csp_header(self):
        resp = self.client.get(
            reverse('ticket_detail', args=[self.ticket.pk]))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode()

        # The panel carrying the toggle must actually be on the page.
        self.assertIn('id="t1-route-form"', html)

        script = re.search(
            r'<script([^>]*)>\(function\(\)\{var f=document\.getElementById\('
            r"'t1-route-form'\)", html)
        self.assertIsNotNone(script, 'route-toggle script not found')

        nonce = re.search(r'nonce="([^"]+)"', script.group(1))
        self.assertIsNotNone(nonce, 'route-toggle script has no nonce')

        csp = resp.headers.get('Content-Security-Policy', '')
        self.assertIn(f"'nonce-{nonce.group(1)}'", csp)

    def test_no_inline_script_on_the_page_is_missing_a_nonce(self):
        resp = self.client.get(
            reverse('ticket_detail', args=[self.ticket.pk]))
        html = resp.content.decode()
        bare = re.findall(r'<script(?![^>]*\bnonce=)(?![^>]*\bsrc=)[^>]*>', html)
        self.assertEqual(bare, [], f'inline script(s) without a nonce: {bare}')


# ──────────────────────────────────────────────────────────────────────────── #
# 24. Evidence staging — attachments survive a validation-error re-render       #
# ──────────────────────────────────────────────────────────────────────────── #

@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_staging_test_media_'))
class EvidenceStagingTest(TestCase):
    """A browser cannot repopulate <input type=file> after a page load, so the
    server has to hold uploaded evidence across a failed submit itself."""

    def setUp(self):
        self.t1 = _make_t1('staging_t1')
        self.client.login(username='staging_t1', password='testpass123')

    @staticmethod
    def _files(*names):
        return [
            SimpleUploadedFile(name, f'evidence for {name}'.encode())
            for name in names
        ]

    def test_multiple_files_in_one_post_all_attach(self):
        """The create-ticket evidence loop had no coverage at all before this."""
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_files=self._files('a.log', 'b.log', 'c.txt'),
        ))
        self.assertEqual(resp.status_code, 302)
        ticket = Ticket.objects.get()
        self.assertEqual(ticket.attachments.count(), 3)
        self.assertEqual(
            sorted(a.original_name for a in ticket.attachments.all()),
            ['a.log', 'b.log', 'c.txt'],
        )
        # Nothing left behind in staging once the ticket exists.
        self.assertEqual(StagedAttachment.objects.count(), 0)

    def test_files_survive_a_validation_error(self):
        """Regression: evidence used to be discarded when the form came back."""
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='',  # required — forces a re-render
            evidence_files=self._files('keepme.log', 'keeptoo.txt'),
        ))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(TicketAttachment.objects.count(), 0)

        staged = StagedAttachment.objects.filter(uploaded_by=self.t1)
        self.assertEqual(staged.count(), 2)

        # And the analyst can see they are being held.
        html = resp.content.decode()
        self.assertIn('keepme.log', html)
        self.assertIn('keeptoo.txt', html)
        token = staged.first().token
        self.assertIn(f'name="evidence_token" value="{token}"', html)

    def test_resubmit_adopts_staged_files_without_reattaching(self):
        """The whole point: fix the field, submit again, keep the evidence."""
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='',
            evidence_files=self._files('first.log', 'second.log'),
        ))
        token = StagedAttachment.objects.first().token

        # Second submit is valid and carries NO files — only the token.
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=token,
        ))
        self.assertEqual(resp.status_code, 302)

        ticket = Ticket.objects.get()
        self.assertEqual(
            sorted(a.original_name for a in ticket.attachments.all()),
            ['first.log', 'second.log'],
        )
        self.assertEqual(StagedAttachment.objects.count(), 0)

    def test_resubmit_with_extra_files_accumulates(self):
        """Adding more files on the retry tops up rather than replaces."""
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('one.log'),
        ))
        token = StagedAttachment.objects.first().token

        self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=token, evidence_files=self._files('two.log'),
        ))
        ticket = Ticket.objects.get()
        self.assertEqual(
            sorted(a.original_name for a in ticket.attachments.all()),
            ['one.log', 'two.log'],
        )

    def test_invalid_file_is_rejected_by_name(self):
        """A rejected upload must be named, never silently dropped."""
        resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_files=[
                SimpleUploadedFile('good.log', b'fine'),
                SimpleUploadedFile('bad.html', b'<script>alert(1)</script>'),
            ],
        ))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(Ticket.objects.count(), 0)
        self.assertContains(resp, 'bad.html')
        # The valid one is still held, so the analyst only re-picks the bad one.
        self.assertEqual(
            [s.original_name for s in StagedAttachment.objects.all()], ['good.log'],
        )

    def test_batch_over_total_cap_is_refused_with_a_message(self):
        """Guards the bare nginx 413 that reads as 'my files vanished'."""
        with patch('apps.incidents.staging.MAX_ATTACHMENT_BATCH_SIZE', 12):
            resp = self.client.post(reverse('create_ticket'), _ticket_post_data(
                evidence_files=[
                    SimpleUploadedFile('a.log', b'0123456789'),
                    SimpleUploadedFile('b.log', b'0123456789'),
                ],
            ))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'b.log')
        self.assertEqual(StagedAttachment.objects.count(), 1)

    def test_staged_files_are_scoped_to_their_uploader(self):
        """The token travels in a hidden field and a query string, so it is not
        a secret — the uploader check is what actually protects the evidence."""
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('mine.log'),
        ))
        staged = StagedAttachment.objects.get()

        _make_t1('staging_intruder')
        self.client.logout()
        self.client.login(username='staging_intruder', password='testpass123')

        # Cannot discard it...
        resp = self.client.post(
            reverse('discard_staged_attachment', args=[staged.pk]))
        self.assertEqual(resp.status_code, 404)
        self.assertTrue(StagedAttachment.objects.filter(pk=staged.pk).exists())

        # ...and cannot adopt it onto a ticket of their own.
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=staged.token,
        ))
        self.assertEqual(Ticket.objects.get().attachments.count(), 0)
        self.assertTrue(StagedAttachment.objects.filter(pk=staged.pk).exists())

    def test_owner_can_discard_a_staged_file(self):
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('oops.log', 'keep.log'),
        ))
        unwanted = StagedAttachment.objects.get(original_name='oops.log')
        stored = unwanted.file.path

        resp = self.client.post(
            reverse('discard_staged_attachment', args=[unwanted.pk]))
        self.assertEqual(resp.status_code, 204)
        # Discard is now recoverable: the row and its bytes survive so the
        # analyst can undo a misclick. purge_staged_attachments clears them.
        unwanted.refresh_from_db()
        self.assertIsNotNone(unwanted.discarded_at)
        self.assertTrue(Path(stored).exists())

        # The other one is untouched and still adoptable; the discarded one is
        # excluded from the picker and never reaches the ticket.
        token = StagedAttachment.objects.filter(
            discarded_at__isnull=True).get().token
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=token))
        self.assertEqual(
            [a.original_name for a in Ticket.objects.get().attachments.all()],
            ['keep.log'],
        )

    def test_discard_requires_post(self):
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('x.log'),
        ))
        staged = StagedAttachment.objects.get()
        resp = self.client.get(
            reverse('discard_staged_attachment', args=[staged.pk]))
        self.assertEqual(resp.status_code, 405)
        self.assertTrue(StagedAttachment.objects.filter(pk=staged.pk).exists())

    def test_case_mode_switch_carries_the_evidence_token(self):
        """The single-to-multi toggle is a full page load; without the token in
        the link the staged evidence would be stranded."""
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('shared.log'),
        ))
        token = StagedAttachment.objects.get().token

        resp = self.client.get(
            reverse('create_project_incident'), {'evidence_token': token})
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'shared.log')
        self.assertContains(resp, f'name="evidence_token" value="{token}"')

    def test_purge_command_removes_only_expired_staging(self):
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('stale.log'),
        ))
        ticket = _make_ticket(created_by=self.t1)
        TicketAttachment.objects.create(
            ticket=ticket, file=SimpleUploadedFile('real.log', b'evidence'),
            original_name='real.log', uploaded_by=self.t1,
        )
        stale_path = StagedAttachment.objects.get().file.path

        call_command('purge_staged_attachments', '--hours', '0')

        self.assertEqual(StagedAttachment.objects.count(), 0)
        self.assertFalse(Path(stale_path).exists())
        # Real evidence is never in scope for the purge.
        self.assertEqual(TicketAttachment.objects.count(), 1)

    def test_purge_command_spares_fresh_staging(self):
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('fresh.log'),
        ))
        call_command('purge_staged_attachments', '--hours', '24')
        self.assertEqual(StagedAttachment.objects.count(), 1)

    def test_purge_command_sweeps_row_less_files(self):
        """A file is written before its row commits, so a crash in between can
        leave one behind that the row-driven purge would never see."""
        orphan_dir = Path(settings.MEDIA_ROOT) / 'staged_attachments' / ('a' * 32)
        orphan_dir.mkdir(parents=True, exist_ok=True)
        orphan = orphan_dir / 'ghost.log'
        orphan.write_bytes(b'no row points at me')

        # A staged file that DOES have a row must survive a purge that spares
        # its age, proving the sweep goes by row-existence and not by path.
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='', evidence_files=self._files('tracked.log'),
        ))
        tracked = Path(StagedAttachment.objects.get().file.path)

        call_command('purge_staged_attachments', '--hours', '0')

        self.assertFalse(orphan.exists())
        self.assertFalse(orphan_dir.exists(), 'empty token dir should be tidied')
        # The tracked one went too (--hours 0), but via its row, not the sweep.
        self.assertEqual(StagedAttachment.objects.count(), 0)
        self.assertFalse(tracked.exists())

    def test_purge_sweep_spares_recent_row_less_files(self):
        """A file being written by a request in flight must not be yanked."""
        orphan_dir = Path(settings.MEDIA_ROOT) / 'staged_attachments' / ('b' * 32)
        orphan_dir.mkdir(parents=True, exist_ok=True)
        inflight = orphan_dir / 'writing-right-now.log'
        inflight.write_bytes(b'still uploading')

        call_command('purge_staged_attachments', '--hours', '24')

        self.assertTrue(inflight.exists())


# ──────────────────────────────────────────────────────────────────────────── #
# 25. Multi-file upload on the ticket detail page                               #
# ──────────────────────────────────────────────────────────────────────────── #

@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_multiupload_test_media_'))
class MultiFileAttachmentUploadTest(TestCase):
    def setUp(self):
        self.t1 = _make_t1('multi_t1')
        self.ticket = _make_ticket(
            created_by=self.t1, assigned_to=self.t1, status=Ticket.STATUS_NEW,
        )
        self.client.login(username='multi_t1', password='testpass123')

    def test_several_files_upload_in_one_post(self):
        resp = self.client.post(
            reverse('upload_attachment', args=[self.ticket.pk]),
            {
                'description': 'batch evidence',
                'file': [
                    SimpleUploadedFile('one.log', b'first'),
                    SimpleUploadedFile('two.log', b'second'),
                    SimpleUploadedFile('three.txt', b'third'),
                ],
            },
        )
        self.assertEqual(resp.status_code, 302)
        attachments = self.ticket.attachments.all()
        self.assertEqual(attachments.count(), 3)
        # The single description applies to the whole batch.
        self.assertTrue(all(a.description == 'batch evidence' for a in attachments))
        self.assertEqual(
            sorted(a.original_name for a in attachments),
            ['one.log', 'three.txt', 'two.log'],
        )

    def test_a_single_file_still_works(self):
        resp = self.client.post(
            reverse('upload_attachment', args=[self.ticket.pk]),
            {'description': '', 'file': SimpleUploadedFile('solo.log', b'only')},
        )
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(self.ticket.attachments.count(), 1)

    def test_one_bad_file_rejects_the_batch_and_names_it(self):
        resp = self.client.post(
            reverse('upload_attachment', args=[self.ticket.pk]),
            {
                'description': '',
                'file': [
                    SimpleUploadedFile('fine.log', b'ok'),
                    SimpleUploadedFile('evil.html', b'<script>alert(1)</script>'),
                ],
            },
            follow=True,
        )
        self.assertEqual(self.ticket.attachments.count(), 0)
        self.assertContains(resp, 'html')

    def test_multiple_form_accepts_a_list(self):
        form = AttachmentForm(
            data={'description': 'ok'},
            files={'file': [
                SimpleUploadedFile('a.log', b'one'),
                SimpleUploadedFile('b.log', b'two'),
            ]},
        )
        self.assertTrue(form.is_valid(), msg=form.errors)
        self.assertEqual(len(form.cleaned_data['file']), 2)


# ──────────────────────────────────────────────────────────────────────────── #
# 26. Incident report is a SOC deliverable — role access matrix                  #
# ──────────────────────────────────────────────────────────────────────────── #

class ReportAccessRoleTest(TestCase):
    """The report carries NT branding and the sign-off block and is aimed
    outward, so the parties named *in* a case must not be able to mint one."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('rpt_t1')
        cls.t2 = _make_t2('rpt_t2')
        cls.manager = _make_user('rpt_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('rpt_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.owner = _make_user('rpt_owner', UserProfile.ROLE_SYSTEM_OWNER)
        cls.forensic = _make_forensic('rpt_forensic')
        cls.redteam = _make_redteam_manager('rpt_redteam')
        cls.root = User.objects.create_superuser('rpt_root', password='testpass123')

        cls.ticket = _make_ticket(
            created_by=cls.t1,
            assigned_admin=cls.admin,
            system_owner=cls.owner,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            remediation_summary='Unauthorized service removed.',
            containment_report='Host isolated.',
        )
        # Give the response-team roles a real reason to see this ticket, so the
        # refusal below is about the report and not about visibility.
        TicketSubtask.objects.create(
            ticket=cls.ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=cls.forensic,
        )
        TicketSubtask.objects.create(
            ticket=cls.ticket, subtask_type=TicketSubtask.TYPE_VA_PT,
            title='Pentest', assigned_to=cls.redteam,
        )

    def _try_all_three(self, user):
        """(preview, docx, pdf) status codes for this user."""
        self.client.force_login(user)
        preview = self.client.get(
            reverse('ticket_report_preview', args=[self.ticket.pk])).status_code
        docx = self.client.post(
            reverse('ticket_report_docx', args=[self.ticket.pk])).status_code
        pdf = self.client.post(
            reverse('ticket_report_pdf', args=[self.ticket.pk])).status_code
        return preview, docx, pdf

    def test_soc_tier1_can_reach_the_report(self):
        preview, docx, _pdf = self._try_all_three(self.t1)
        self.assertEqual(preview, 200)
        self.assertEqual(docx, 200)

    def test_soc_tier2_can_reach_the_report(self):
        preview, docx, _pdf = self._try_all_three(self.t2)
        self.assertEqual(preview, 200)
        self.assertEqual(docx, 200)

    def test_soc_manager_can_reach_the_report(self):
        preview, docx, _pdf = self._try_all_three(self.manager)
        self.assertEqual(preview, 200)
        self.assertEqual(docx, 200)

    def test_superuser_can_reach_the_report(self):
        preview, docx, _pdf = self._try_all_three(self.root)
        self.assertEqual(preview, 200)
        self.assertEqual(docx, 200)

    def test_assigned_system_admin_is_refused(self):
        """The regression for this item: the admin can see the ticket in full,
        and previously could export its report."""
        self.assertIn(
            self.ticket, Ticket.objects.visible_to(self.admin),
            'fixture broken — the admin should still see the ticket itself',
        )
        self.assertEqual(self._try_all_three(self.admin), (404, 404, 404))

    def test_system_owner_is_refused(self):
        self.assertIn(self.ticket, Ticket.objects.visible_to(self.owner))
        self.assertEqual(self._try_all_three(self.owner), (404, 404, 404))

    def test_forensic_analyst_is_refused(self):
        self.assertIn(self.ticket, Ticket.objects.visible_to(self.forensic))
        self.assertEqual(self._try_all_three(self.forensic), (404, 404, 404))

    def test_redteam_manager_is_refused(self):
        self.assertIn(self.ticket, Ticket.objects.visible_to(self.redteam))
        self.assertEqual(self._try_all_three(self.redteam), (404, 404, 404))

    def test_refused_export_leaves_report_provenance_untouched(self):
        """Export writes report_generated_by / sha256 / the stale-report
        watermark. A blocked role must not be able to move any of it."""
        self.client.force_login(self.t1)
        self.client.post(reverse('ticket_report_docx', args=[self.ticket.pk]))
        self.ticket.refresh_from_db()
        soc_generator = self.ticket.report_generated_by
        soc_digest = self.ticket.report_sha256
        soc_watermark = self.ticket.report_ticket_updated_at
        self.assertEqual(soc_generator, self.t1)

        self.client.force_login(self.admin)
        self.client.post(reverse('ticket_report_docx', args=[self.ticket.pk]))
        self.client.post(reverse('ticket_report_pdf', args=[self.ticket.pk]))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.report_generated_by, soc_generator)
        self.assertEqual(self.ticket.report_sha256, soc_digest)
        self.assertEqual(self.ticket.report_ticket_updated_at, soc_watermark)

    def test_report_buttons_hidden_from_the_assigned_admin(self):
        self.client.force_login(self.admin)
        html = self.client.get(
            reverse('ticket_detail', args=[self.ticket.pk])).content.decode()
        self.assertNotIn('Report Preview', html)
        self.assertNotIn('Report DOCX', html)
        self.assertNotIn('Report PDF', html)
        # The admin still has their own ticket page.
        self.assertIn('กลับรายการ Ticket', html)

    def test_report_buttons_shown_to_soc(self):
        self.client.force_login(self.t1)
        html = self.client.get(
            reverse('ticket_detail', args=[self.ticket.pk])).content.decode()
        self.assertIn('Report Preview', html)
        self.assertIn('Report DOCX', html)
        self.assertIn('Report PDF', html)


# ──────────────────────────────────────────────────────────────────────────── #
# 27. Section 8 no longer prints the static remediation checklist               #
# ──────────────────────────────────────────────────────────────────────────── #

class ReportSectionEightTest(TestCase):
    """The 15 boilerplate items were never ticked by anything; section 6 holds
    the real, data-driven containment checklist."""

    # A sample of the removed items, in both scripts used by the form.
    REMOVED = [
        'ติดตั้ง Sysmon',
        'ติดตั้ง Agent Wazuh',
        'Dump memory ของเครื่อง Server',
        'Harden Configuration',
        'รวบรวม Event Logs เพื่อส่งต่อ ปปกก.',
    ]

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('sec8_t1')
        cls.ticket = _make_ticket(
            created_by=cls.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            action_required='1) Block IoC\n2) Isolate the host',
            remediation_summary='Unauthorized service removed.',
            containment_report='Host isolated and C2 destination blocked.',
        )

    def test_section_eight_keeps_its_title_and_data_rows(self):
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('ticket_report_preview', args=[self.ticket.pk]))
        self.assertContains(response, 'สรุปผลการดำเนินการแก้ไข')
        self.assertContains(response, 'ผลการตรวจสอบ / Investigation Findings')
        self.assertContains(response, 'มาตรการควบคุม / Countermeasure')
        self.assertContains(response, 'Unauthorized service removed.')
        self.assertContains(response, 'Host isolated and C2 destination blocked.')

    def test_preview_no_longer_prints_the_boilerplate(self):
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('ticket_report_preview', args=[self.ticket.pk]))
        for item in self.REMOVED:
            self.assertNotContains(response, item)

    def test_committed_docx_template_no_longer_carries_the_boilerplate(self):
        """Guards a forgotten template rebuild. The checklist lived in the .docx
        as literal text, not as a {{placeholder}}, so
        test_build_script_matches_committed_template — which compares only
        placeholder sets — would never notice a stale file."""
        doc = Document(str(REPORT_TEMPLATE_PATH))
        text = '\n'.join(p.text for p in _iter_paragraphs(doc))
        for item in self.REMOVED:
            self.assertNotIn(item, text)

    def test_generated_docx_no_longer_carries_the_boilerplate(self):
        report = generate_ticket_report(self.ticket.pk, generated_by=self.t1)
        text = _docx_text(report.content)
        for item in self.REMOVED:
            self.assertNotIn(item, text)
        # Section 8's real content is still rendered.
        self.assertIn('Unauthorized service removed.', text)
        self.assertIn('Host isolated and C2 destination blocked.', text)

    def test_section_six_dynamic_checklist_still_works(self):
        """The removal must not touch the checklist that is actually driven by
        ticket data."""
        sections = build_ticket_report_sections(
            build_ticket_report_context(self.ticket), self.ticket)
        section_six = next(s for s in sections if s['number'] == '6')
        self.assertEqual(section_six['rows'][0]['type'], 'containment_checklist')
        self.assertEqual(
            [i['text'] for i in section_six['rows'][0]['items']],
            ['1) Block IoC', '2) Isolate the host'],
        )

    def test_no_section_uses_the_removed_static_checklist_row(self):
        sections = build_ticket_report_sections(
            build_ticket_report_context(self.ticket), self.ticket)
        types = {row.get('type') for s in sections for row in s['rows']}
        self.assertNotIn('checklist', types)


# ──────────────────────────────────────────────────────────────────────────── #
# 28. Report matches the NT paper form's layout                                 #
# ──────────────────────────────────────────────────────────────────────────── #

class ReportNTFormLayoutTest(TestCase):
    """Alignment with the official NT incident-report form the SOC files
    alongside this report."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('ntform_t1')
        cls.ticket = _make_ticket(
            created_by=cls.t1,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            severity='High',
            ncsa_severity=Ticket.NCSA_SEVERITY_SEVERE,
            # 2 July 2026 = 2 ก.ค. 2569 → "2 / ก.ค. / 69"
            incident_datetime=timezone.make_aware(datetime(2026, 7, 2, 11, 0)),
            ioc_user='DOMAIN\\svc_backup',
            destination_ip='203.0.113.50',
        )

    def _sections(self):
        return build_ticket_report_sections(
            build_ticket_report_context(self.ticket), self.ticket)

    def _section(self, number):
        return next(s for s in self._sections() if s['number'] == number)

    # ── Thai Buddhist dates, section 1 only ──────────────────────────── #

    def test_section_one_dates_use_the_thai_form_style(self):
        rows = {r['label']: r['value'] for r in self._section('1')['rows']
                if r.get('type') == 'kv'}
        self.assertEqual(rows['1.2 วันที่ เวลา ที่พบเหตุ'], '2 / ก.ค. / 69  11:00 น.')
        # Still the same single field in both rows — unchanged by design.
        self.assertEqual(rows['1.3 วันที่ เวลา ที่เกิดเหตุ'],
                         rows['1.2 วันที่ เวลา ที่พบเหตุ'])

    def test_generated_at_stays_gregorian(self):
        """Only the section 1 date rows switched; the meta line must stay
        sortable and unambiguous."""
        ctx = build_ticket_report_context(self.ticket)
        self.assertRegex(ctx['generated_at'], r'^\d{2}/\d{2}/\d{4} \d{2}:\d{2}$')

    def test_thai_date_helper_handles_every_month_and_blank(self):
        from apps.incidents.reports import _format_dt_thai
        self.assertEqual(_format_dt_thai(None), '-')
        expected = ['ม.ค.', 'ก.พ.', 'มี.ค.', 'เม.ย.', 'พ.ค.', 'มิ.ย.',
                    'ก.ค.', 'ส.ค.', 'ก.ย.', 'ต.ค.', 'พ.ย.', 'ธ.ค.']
        for month, abbr in enumerate(expected, start=1):
            rendered = _format_dt_thai(
                timezone.make_aware(datetime(2026, month, 5, 9, 30)))
            self.assertEqual(rendered, f'5 / {abbr} / 69  09:30 น.')

    def test_buddhist_year_rolls_over_correctly(self):
        from apps.incidents.reports import _format_dt_thai
        # 2057 CE = 2600 BE → two-digit year "00", not "600".
        self.assertEqual(
            _format_dt_thai(timezone.make_aware(datetime(2057, 1, 1, 0, 0))),
            '1 / ม.ค. / 00  00:00 น.',
        )

    # ── Section 1 numbering ──────────────────────────────────────────── #

    def test_section_one_rows_are_numbered_sequentially(self):
        labels = [r['label'] for r in self._section('1')['rows']]
        self.assertEqual(len(labels), 20)
        for index, label in enumerate(labels, start=1):
            self.assertTrue(
                label.startswith(f'1.{index} '),
                f'row {index} is {label!r}, expected a "1.{index} " prefix',
            )

    def test_only_section_one_is_numbered(self):
        """The paper form numbers section 1 only — sections 3 and 4 are plain."""
        for number in ('3', '4'):
            for row in self._section(number)['rows']:
                if row.get('type') in ('kv', 'checks'):
                    self.assertFalse(
                        re.match(r'^\d+\.\d+ ', row['label']),
                        f'section {number} row {row["label"]!r} should not be numbered',
                    )

    # ── Checkbox ordering ────────────────────────────────────────────── #

    def test_checkbox_options_run_low_to_high(self):
        rows = {r['label']: r for r in self._section('1')['rows']
                if r.get('type') == 'checks'}
        sev = [o['label'] for o in rows['1.6 ระดับความรุนแรง (อ้างอิงตามระบบ SIEM)']['options']]
        self.assertEqual(sev, ['Low', 'Medium', 'High', 'Critical'])

        ncsa = [o['label'] for o in rows['1.9 ระดับความรุนแรง (อ้างอิงตาม สกมช.)']['options']]
        self.assertEqual(ncsa, ['ไม่ร้ายแรง', 'ร้ายแรง', 'วิกฤต'])

        kind = [o['label'] for o in rows['1.5 ประเภท: event หรือ incident']['options']]
        self.assertEqual(kind, ['Incident', 'Event'])

    def test_reordering_did_not_change_which_option_is_ticked(self):
        rows = {r['label']: r for r in self._section('1')['rows']
                if r.get('type') == 'checks'}
        ticked = lambda label: [  # noqa: E731
            o['label'] for o in rows[label]['options'] if o['checked']]
        self.assertEqual(ticked('1.6 ระดับความรุนแรง (อ้างอิงตามระบบ SIEM)'), ['High'])
        self.assertEqual(ticked('1.9 ระดับความรุนแรง (อ้างอิงตาม สกมช.)'), ['ร้ายแรง'])
        self.assertEqual(ticked('1.5 ประเภท: event หรือ incident'), ['Incident'])

    # ── Section 4 User row ───────────────────────────────────────────── #

    def test_section_four_carries_the_user_row(self):
        rows = {r['label']: r['value'] for r in self._section('4')['rows']}
        self.assertEqual(rows['User'], 'DOMAIN\\svc_backup')
        # Placed after IP, as on the paper form.
        labels = [r['label'] for r in self._section('4')['rows']]
        self.assertEqual(labels[-2:], ['IP', 'User'])

    def test_blank_user_renders_as_a_dash_not_an_empty_cell(self):
        self.ticket.ioc_user = ''
        rows = {r['label']: r['value'] for r in self._section('4')['rows']}
        self.assertEqual(rows['User'], '-')

    def test_user_reaches_the_generated_docx(self):
        report = generate_ticket_report(self.ticket.pk, generated_by=self.t1)
        text = _docx_text(report.content)
        self.assertIn('DOMAIN\\svc_backup', text)
        self.assertNotIn('{{ioc_user}}', text)

    def test_ioc_user_is_editable_on_the_create_form(self):
        self.client.force_login(self.t1)
        html = self.client.get(reverse('create_ticket')).content.decode()
        self.assertIn('name="ioc_user"', html)

    # ── Banner, appendix clause, coordination note ───────────────────── #

    def test_banner_names_the_containment_form(self):
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('ticket_report_preview', args=[self.ticket.pk]))
        self.assertContains(response, 'INCIDENT REPORT: Containment')

        doc = Document(str(REPORT_TEMPLATE_PATH))
        text = '\n'.join(p.text for p in _iter_paragraphs(doc))
        self.assertIn('INCIDENT REPORT: Containment', text)

    def test_appendix_clause_heading_appears_in_both_outputs(self):
        """It was in the DOCX but missing from HTML/PDF, so the two disagreed."""
        clause = 'ข้อ ๑ การจำแนกหมวดหมู่ของภัยคุกคามทางไซเบอร์'
        self.client.force_login(self.t1)
        response = self.client.get(
            reverse('ticket_report_preview', args=[self.ticket.pk]))
        self.assertContains(response, clause)

        doc = Document(str(REPORT_TEMPLATE_PATH))
        text = '\n'.join(p.text for p in _iter_paragraphs(doc))
        self.assertIn(clause, text)

    def test_coordination_note_lists_the_hardening_contact(self):
        from apps.incidents.report_content import GUIDANCE_COORDINATION_NOTE
        self.assertIn('Hardening', GUIDANCE_COORDINATION_NOTE)
        self.assertIn('02-575-6883', GUIDANCE_COORDINATION_NOTE)
        # The pre-existing วปกก. contact must survive the addition.
        self.assertIn('02-574-8186', GUIDANCE_COORDINATION_NOTE)


# ──────────────────────────────────────────────────────────────────────────── #
# 29. Evidence deletion safeguards                                              #
# ──────────────────────────────────────────────────────────────────────────── #

@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_delete_test_media_'))
class AttachmentDeletionSafeguardTest(TestCase):
    """Removing evidence must be deliberate, attributable, reversible, and
    impossible once a case is closed."""

    @classmethod
    def setUpTestData(cls):
        cls.uploader = _make_t1('del_uploader')
        cls.other_t1 = _make_t1('del_other_t1')
        cls.t2 = _make_t2('del_t2')
        cls.manager = _make_user('del_manager', UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('del_admin', UserProfile.ROLE_SYSTEM_ADMIN)
        cls.root = User.objects.create_superuser('del_root', password='testpass123')

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)

    def _ticket(self, status=Ticket.STATUS_NEW):
        return _make_ticket(
            status=status, created_by=self.uploader, assigned_admin=self.admin,
        )

    def _attachment(self, ticket, name='evidence.log'):
        return TicketAttachment.objects.create(
            ticket=ticket,
            file=SimpleUploadedFile(name, b'evidence bytes'),
            original_name=name,
            uploaded_by=self.uploader,
        )

    def _delete(self, user, attachment, reason='แนบผิดเคส'):
        self.client.force_login(user)
        data = {'reason': reason} if reason is not None else {}
        return self.client.post(
            reverse('delete_attachment', args=[attachment.pk]), data)

    # ── Who may delete ───────────────────────────────────────────────── #

    def test_uploader_may_delete_their_own_evidence(self):
        att = self._attachment(self._ticket())
        self._delete(self.uploader, att)
        self.assertFalse(TicketAttachment.objects.filter(pk=att.pk).exists())

    def test_soc_manager_may_delete(self):
        att = self._attachment(self._ticket())
        self._delete(self.manager, att)
        self.assertFalse(TicketAttachment.objects.filter(pk=att.pk).exists())

    def test_plain_soc_staff_who_did_not_upload_is_refused(self):
        """Regression: 'any SOC member' used to be enough, and SOC sees every
        ticket — so an uninvolved analyst could remove another team's evidence."""
        att = self._attachment(self._ticket())
        for user in (self.other_t1, self.t2):
            with self.subTest(user=user.username):
                self._delete(user, att)
                self.assertTrue(
                    TicketAttachment.objects.filter(pk=att.pk).exists(),
                    f'{user.username} should not be able to delete this',
                )

    def test_refused_delete_writes_no_timeline_row(self):
        ticket = self._ticket()
        att = self._attachment(ticket)
        self._delete(self.other_t1, att)
        self.assertFalse(
            TicketLog.objects.filter(ticket=ticket, author=self.other_t1).exists())

    # ── Closed cases are frozen ──────────────────────────────────────── #

    def test_deletion_refused_on_a_closed_ticket_for_every_role(self):
        """The sharp edge: a closed case is when a quiet deletion would never
        be noticed. Terminal check runs ahead of the superuser bypass."""
        for status in (Ticket.STATUS_APPROVED, Ticket.STATUS_CLOSED_EVENT):
            for user in (self.uploader, self.manager, self.root):
                with self.subTest(status=status, user=user.username):
                    att = self._attachment(self._ticket(status=status))
                    self._delete(user, att)
                    self.assertTrue(
                        TicketAttachment.objects.filter(pk=att.pk).exists(),
                        f'{user.username} deleted evidence on a {status} ticket',
                    )

    # ── A reason is mandatory ────────────────────────────────────────── #

    def test_delete_without_a_reason_is_refused(self):
        att = self._attachment(self._ticket())
        for reason in (None, '', '   '):
            with self.subTest(reason=reason):
                self._delete(self.uploader, att, reason=reason)
                self.assertTrue(
                    TicketAttachment.objects.filter(pk=att.pk).exists())

    def test_reason_is_stored_and_logged(self):
        ticket = self._ticket()
        att = self._attachment(ticket)
        self._delete(self.uploader, att, reason='ไฟล์ซ้ำกับที่แนบไว้แล้ว')
        att.refresh_from_db()
        self.assertEqual(att.deleted_reason, 'ไฟล์ซ้ำกับที่แนบไว้แล้ว')
        self.assertTrue(TicketLog.objects.filter(
            ticket=ticket, note__contains='ไฟล์ซ้ำกับที่แนบไว้แล้ว').exists())

    def test_get_does_not_delete(self):
        att = self._attachment(self._ticket())
        self.client.force_login(self.uploader)
        resp = self.client.get(reverse('delete_attachment', args=[att.pk]))
        self.assertEqual(resp.status_code, 405)
        self.assertTrue(TicketAttachment.objects.filter(pk=att.pk).exists())

    # ── Restore ──────────────────────────────────────────────────────── #

    def test_soc_manager_can_restore_deleted_evidence(self):
        ticket = self._ticket()
        att = self._attachment(ticket)
        self._delete(self.uploader, att)

        self.client.force_login(self.manager)
        resp = self.client.post(reverse('restore_attachment', args=[att.pk]))
        self.assertRedirects(resp, reverse('ticket_detail', args=[ticket.pk]))

        att.refresh_from_db()
        self.assertIsNone(att.deleted_at)
        self.assertEqual(att.deleted_reason, '')
        self.assertTrue(TicketAttachment.objects.filter(pk=att.pk).exists())
        # And it downloads again.
        self.assertEqual(
            self.client.get(reverse('download_attachment', args=[att.pk])).status_code,
            200,
        )

    def test_non_manager_cannot_restore(self):
        att = self._attachment(self._ticket())
        self._delete(self.uploader, att)
        for user in (self.uploader, self.t2, self.admin):
            with self.subTest(user=user.username):
                self.client.force_login(user)
                self.client.post(reverse('restore_attachment', args=[att.pk]))
                self.assertFalse(
                    TicketAttachment.objects.filter(pk=att.pk).exists())

    def test_restore_works_on_a_closed_ticket_even_though_delete_does_not(self):
        """Deliberate asymmetry: refusing removal on a closed case protects the
        evidence set; refusing recovery would only make a mistake permanent."""
        ticket = self._ticket()
        att = self._attachment(ticket)
        self._delete(self.uploader, att)

        Ticket.objects.filter(pk=ticket.pk).update(status=Ticket.STATUS_APPROVED)
        self.client.force_login(self.manager)
        self.client.post(reverse('restore_attachment', args=[att.pk]))
        self.assertTrue(TicketAttachment.objects.filter(pk=att.pk).exists())

    def test_deleted_evidence_is_listed_for_managers_only(self):
        ticket = self._ticket()
        att = self._attachment(ticket, name='removed-evidence.log')
        self._delete(self.uploader, att, reason='แนบผิดเคส')
        restore_url = reverse('restore_attachment', args=[att.pk])

        self.client.force_login(self.manager)
        html = self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode()
        self.assertIn('ไฟล์ที่ถูกลบ', html)
        self.assertIn(restore_url, html)
        self.assertIn('แนบผิดเคส', html)

        # The uploader still sees the deletion in the timeline audit note —
        # that is the point of it — but gets no recovery block or control.
        self.client.force_login(self.uploader)
        html = self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode()
        self.assertNotIn('ไฟล์ที่ถูกลบ', html)
        self.assertNotIn(restore_url, html)

    def test_delete_control_hidden_once_the_ticket_closes(self):
        ticket = self._ticket(status=Ticket.STATUS_APPROVED)
        self._attachment(ticket)
        self.client.force_login(self.manager)
        html = self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode()
        self.assertNotIn(
            reverse('delete_attachment', args=[ticket.attachments.get().pk]), html)


# ──────────────────────────────────────────────────────────────────────────── #
# 30. Staged discard is recoverable                                             #
# ──────────────────────────────────────────────────────────────────────────── #

@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix='soc_discard_test_media_'))
class StagedDiscardRecoveryTest(TestCase):
    def setUp(self):
        self.t1 = _make_t1('discard_t1')
        self.client.login(username='discard_t1', password='testpass123')
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            device_name='',
            evidence_files=[SimpleUploadedFile('oops.log', b'staged bytes')],
        ))
        self.staged = StagedAttachment.objects.get()

    def _discard(self):
        return self.client.post(
            reverse('discard_staged_attachment', args=[self.staged.pk]))

    def _restore(self):
        return self.client.post(
            reverse('restore_staged_attachment', args=[self.staged.pk]))

    def test_discard_keeps_the_bytes(self):
        stored = self.staged.file.path
        self.assertEqual(self._discard().status_code, 204)
        self.staged.refresh_from_db()
        self.assertIsNotNone(self.staged.discarded_at)
        self.assertTrue(Path(stored).exists())

    def test_discarded_file_is_hidden_from_the_picker(self):
        self._discard()
        resp = self.client.get(
            reverse('create_ticket'), {'evidence_token': self.staged.token})
        self.assertNotContains(resp, 'oops.log')

    def test_discarded_file_is_never_adopted_onto_the_ticket(self):
        self._discard()
        self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=self.staged.token))
        self.assertEqual(Ticket.objects.get().attachments.count(), 0)

    def test_undo_brings_it_back_and_it_adopts_normally(self):
        self._discard()
        self.assertEqual(self._restore().status_code, 204)
        self.staged.refresh_from_db()
        self.assertIsNone(self.staged.discarded_at)

        self.client.post(reverse('create_ticket'), _ticket_post_data(
            evidence_token=self.staged.token))
        self.assertEqual(
            [a.original_name for a in Ticket.objects.get().attachments.all()],
            ['oops.log'],
        )

    def test_discarded_rows_do_not_count_toward_the_batch_cap(self):
        from apps.incidents.staging import staged_total_size
        before = staged_total_size(self.t1, self.staged.token)
        self.assertGreater(before, 0)
        self._discard()
        self.assertEqual(staged_total_size(self.t1, self.staged.token), 0)

    def test_purge_removes_discarded_rows_and_their_files(self):
        stored = self.staged.file.path
        self._discard()
        call_command('purge_staged_attachments', '--hours', '0')
        self.assertFalse(StagedAttachment.objects.filter(pk=self.staged.pk).exists())
        self.assertFalse(Path(stored).exists())

    def test_discard_and_restore_are_owner_scoped(self):
        _make_t1('discard_intruder')
        self.client.logout()
        self.client.login(username='discard_intruder', password='testpass123')
        self.assertEqual(self._discard().status_code, 404)
        self.staged.refresh_from_db()
        self.assertIsNone(self.staged.discarded_at)

        self.staged.discarded_at = timezone.now()
        self.staged.save(update_fields=('discarded_at',))
        self.assertEqual(self._restore().status_code, 404)


# ──────────────────────────────────────────────────────────────────────────── #
# 31. Timeline-note edits are versioned                                         #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketLogRevisionTest(TestCase):
    """The timeline is the audit trail, and edit_log lets the author or any SOC
    manager rewrite it — so a rewrite must not be able to erase what was there."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('logrev_t1')
        cls.manager = _make_user('logrev_manager', UserProfile.ROLE_SOC_MANAGER)
        cls.ticket = _make_ticket(created_by=cls.t1)

    def _log(self, note='Attachment removed: secret.log — เหตุผล: ไม่เกี่ยวข้อง'):
        return TicketLog.objects.create(
            ticket=self.ticket, note=note,
            status_at_time=self.ticket.status, author=self.t1,
        )

    def test_editing_banks_the_previous_text(self):
        log = self._log()
        self.client.force_login(self.t1)
        self.client.post(reverse('edit_log', args=[log.pk]), {'note': 'rewritten'})

        log.refresh_from_db()
        self.assertEqual(log.note, 'rewritten')
        revision = log.revisions.get()
        self.assertEqual(
            revision.previous_note,
            'Attachment removed: secret.log — เหตุผล: ไม่เกี่ยวข้อง',
        )
        self.assertEqual(revision.edited_by, self.t1)

    def test_a_deletion_record_survives_being_rewritten(self):
        """The point of the whole change: whoever removed the evidence cannot
        quietly remove the record of having done so."""
        log = self._log()
        self.client.force_login(self.t1)
        self.client.post(reverse('edit_log', args=[log.pk]), {'note': 'routine note'})

        log.refresh_from_db()
        self.assertNotIn('secret.log', log.note)
        self.assertIn('secret.log', log.revisions.get().previous_note)

    def test_repeated_edits_accumulate(self):
        log = self._log(note='first')
        self.client.force_login(self.manager)
        for note in ('second', 'third'):
            self.client.post(reverse('edit_log', args=[log.pk]), {'note': note})
        self.assertEqual(
            [r.previous_note for r in log.revisions.all()], ['second', 'first'])

    def test_unchanged_note_creates_no_revision(self):
        log = self._log(note='unchanged')
        self.client.force_login(self.t1)
        self.client.post(reverse('edit_log', args=[log.pk]), {'note': 'unchanged'})
        self.assertEqual(log.revisions.count(), 0)

    def test_was_edited_flag_and_badge(self):
        log = self._log()
        self.assertFalse(log.was_edited)
        self.client.force_login(self.t1)
        self.client.post(reverse('edit_log', args=[log.pk]), {'note': 'changed'})
        self.assertTrue(TicketLog.objects.get(pk=log.pk).was_edited)

        html = self.client.get(
            reverse('ticket_detail', args=[self.ticket.pk])).content.decode()
        self.assertIn('แก้ไขแล้ว', html)


# ──────────────────────────────────────────────────────────────────────────── #
# 32. Field-level change history (Phase 2)                                      #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketFieldHistoryTest(TestCase):
    """TicketLog says what happened; this says what changed. Without it every
    content edit was destructive and unrecoverable."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('hist_t1')
        cls.t2 = _make_t2('hist_t2')
        cls.manager = _make_user('hist_mgr', UserProfile.ROLE_SOC_MANAGER)

    def test_snapshot_and_record_captures_only_what_moved(self):
        ticket = _make_ticket(
            created_by=self.t1, device_name='OLD-HOST', ip_address='10.0.0.1')
        before = history.snapshot(ticket)
        ticket.device_name = 'NEW-HOST'
        ticket.save()
        rows = history.record_changes(ticket, before, self.t1, source='edit')

        self.assertEqual(len(rows), 1)
        change = ticket.field_changes.get()
        self.assertEqual(change.field_name, 'device_name')
        self.assertEqual(change.old_value, 'OLD-HOST')
        self.assertEqual(change.new_value, 'NEW-HOST')
        self.assertEqual(change.changed_by, self.t1)
        self.assertEqual(change.source, 'edit')

    def test_no_change_records_nothing(self):
        ticket = _make_ticket(created_by=self.t1, device_name='SAME')
        before = history.snapshot(ticket)
        history.record_changes(ticket, before, self.t1)
        self.assertEqual(ticket.field_changes.count(), 0)

    def test_choice_fields_are_stored_as_labels(self):
        ticket = _make_ticket(
            created_by=self.t1, classification=Ticket.CLASSIFICATION_EVENT)
        before = history.snapshot(ticket)
        ticket.classification = Ticket.CLASSIFICATION_INCIDENT
        ticket.save()
        history.record_changes(ticket, before, self.t1)
        change = ticket.field_changes.get(field_name='classification')
        # Reads as the label a human sees, not the stored code.
        self.assertNotEqual(change.new_value, Ticket.CLASSIFICATION_INCIDENT)
        self.assertTrue(change.new_value)

    def test_booleans_render_readably(self):
        ticket = _make_ticket(created_by=self.t1, spread_to_others=False)
        before = history.snapshot(ticket)
        ticket.spread_to_others = True
        ticket.save()
        history.record_changes(ticket, before, self.t1)
        self.assertEqual(
            ticket.field_changes.get(field_name='spread_to_others').new_value, 'ใช่')

    def test_tier2_review_records_its_rewrites(self):
        """Tier 2 overwrites ~25 fields in one submit; the old values used to
        vanish with only a prose note left behind."""
        ticket = _make_ticket(
            created_by=self.t1, status=Ticket.STATUS_ESCALATED_T2,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            device_name='ORIGINAL-HOST', ip_address='10.0.0.9',
            issue_description='original description',
            log_source='Wazuh', issue_type='SIEM',
            detailed_issue='Investigating', detailed_issue2='Investigating Other',
            severity='High', ncsa_severity=Ticket.NCSA_SEVERITY_SEVERE,
            incident_datetime=timezone.now(),
        )
        self.client.force_login(self.t2)
        self.client.post(reverse('ticket_detail', args=[ticket.pk]), {
            'action': 't2_review',
            'status': Ticket.STATUS_T1_REVIEW,
            'classification': Ticket.CLASSIFICATION_INCIDENT,
            'incident_name': '', 'severity': 'High',
            'ncsa_severity': Ticket.NCSA_SEVERITY_SEVERE,
            'incident_datetime': timezone.localtime().strftime('%Y-%m-%dT%H:%M'),
            'log_source': 'Wazuh', 'issue_type': 'SIEM',
            'detailed_issue': 'Investigating',
            'detailed_issue2': 'Investigating Other',
            'device_name': 'CORRECTED-HOST',
            'issue_description': 'corrected description',
            'ip_address': '10.0.0.9',
            'decision_note': 'returning to T1',
        })
        changed = {c.field_name: c for c in ticket.field_changes.all()}
        self.assertIn('device_name', changed)
        self.assertEqual(changed['device_name'].old_value, 'ORIGINAL-HOST')
        self.assertEqual(changed['device_name'].new_value, 'CORRECTED-HOST')
        self.assertEqual(changed['device_name'].source, 't2_review')

    def test_subtask_result_overwrite_is_recorded(self):
        """Result notes could be replaced by any SOC member with no audit at
        all — a forensic analyst's findings could vanish silently."""
        ticket = _make_ticket(created_by=self.t1)
        forensic = _make_forensic('hist_forensic')
        subtask = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='RCA', assigned_to=forensic, result_notes='original findings',
        )
        self.client.force_login(self.manager)
        self.client.post(reverse('update_subtask', args=[subtask.pk]), {
            'status': subtask.status, 'result_notes': 'overwritten findings',
        })
        change = ticket.field_changes.get(subtask=subtask)
        self.assertEqual(change.old_value, 'original findings')
        self.assertEqual(change.new_value, 'overwritten findings')
        self.assertEqual(change.changed_by, self.manager)

    def test_response_tasks_track_each_status_change_per_responder(self):
        """Forensic and Red Team work streams keep independent status clocks
        and audit records; a note-only edit must not reset either clock."""
        ticket = _make_ticket(created_by=self.t1)
        forensic = _make_forensic('status_forensic')
        redteam = _make_redteam_manager('status_redteam')
        forensic_task = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_FORENSIC_RCA,
            title='Forensic RCA', assigned_to=forensic,
        )
        redteam_task = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_VA_PT,
            title='VA assessment', assigned_to=redteam,
        )
        forensic_initial_stamp = forensic_task.status_changed_at
        redteam_initial_stamp = redteam_task.status_changed_at
        self.assertIsNotNone(forensic_initial_stamp)
        self.assertIsNotNone(redteam_initial_stamp)

        self.client.force_login(forensic)
        self.client.post(reverse('update_subtask', args=[forensic_task.pk]), {
            'status': TicketSubtask.STATUS_IN_PROGRESS,
            'result_notes': '',
        })
        forensic_task.refresh_from_db()
        redteam_task.refresh_from_db()
        forensic_status_change = ticket.field_changes.get(
            subtask=forensic_task, field_name='status')
        self.assertGreater(forensic_task.status_changed_at, forensic_initial_stamp)
        self.assertEqual(redteam_task.status_changed_at, redteam_initial_stamp)
        self.assertEqual(forensic_status_change.changed_by, forensic)
        self.assertEqual(forensic_status_change.old_value, dict(
            TicketSubtask.STATUS_CHOICES)[TicketSubtask.STATUS_OPEN])
        self.assertEqual(forensic_status_change.new_value, dict(
            TicketSubtask.STATUS_CHOICES)[TicketSubtask.STATUS_IN_PROGRESS])

        # A result-note correction leaves the current-status timestamp intact.
        forensic_status_stamp = forensic_task.status_changed_at
        self.client.post(reverse('update_subtask', args=[forensic_task.pk]), {
            'status': TicketSubtask.STATUS_IN_PROGRESS,
            'result_notes': 'Initial triage complete',
        })
        forensic_task.refresh_from_db()
        self.assertEqual(forensic_task.status_changed_at, forensic_status_stamp)
        self.assertEqual(ticket.field_changes.filter(
            subtask=forensic_task, field_name='status').count(), 1)

        self.client.force_login(redteam)
        self.client.post(reverse('update_subtask', args=[redteam_task.pk]), {
            'status': TicketSubtask.STATUS_DONE,
            'result_notes': 'Assessment delivered',
        })
        redteam_task.refresh_from_db()
        redteam_status_change = ticket.field_changes.get(
            subtask=redteam_task, field_name='status')
        self.assertGreater(redteam_task.status_changed_at, redteam_initial_stamp)
        self.assertEqual(redteam_status_change.changed_by, redteam)

        self.client.force_login(self.manager)
        response = self.client.get(reverse('ticket_detail', args=[ticket.pk]))
        html = response.content.decode()
        self.assertIn('สร้างเมื่อ', html)
        self.assertIn('สถานะตั้งแต่', html)
        self.assertIn('ประวัติสถานะ', html)
        # The compact history is opt-in and excludes result-note edits.
        self.assertNotIn('class="subtask-status-history" open', html)
        rendered_forensic_task = next(
            task for task in response.context['subtasks']
            if task.pk == forensic_task.pk)
        rendered_redteam_task = next(
            task for task in response.context['subtasks']
            if task.pk == redteam_task.pk)
        self.assertEqual(
            [change.field_name for change in rendered_forensic_task.field_changes.all()],
            ['status'],
        )
        self.assertEqual(
            [change.field_name for change in rendered_redteam_task.field_changes.all()],
            ['status'],
        )

    def test_model_stamps_status_changed_at_outside_the_update_view(self):
        # The stamp belongs to the model, not to update_subtask: admin edits,
        # seeds and data migrations move a subtask's status too, and a
        # view-only stamp leaves every one of them with a stale per-task age.
        ticket = _make_ticket(created_by=self.t1)
        task = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Collect logs',
        )
        seeded = task.status_changed_at
        self.assertIsNotNone(seeded)

        # Reload so the instance carries its stored status, as any other write
        # path would.
        task = TicketSubtask.objects.get(pk=task.pk)
        task.result_notes = 'partial findings'
        task.save()
        task.refresh_from_db()
        self.assertEqual(
            task.status_changed_at, seeded, 'a note edit must not bump it')

        task = TicketSubtask.objects.get(pk=task.pk)
        task.status = TicketSubtask.STATUS_DONE
        task.save()
        task.refresh_from_db()
        self.assertGreater(task.status_changed_at, seeded)

    def test_status_stamp_survives_an_update_fields_save(self):
        # update_fields would otherwise drop the freshly computed column.
        ticket = _make_ticket(created_by=self.t1)
        task = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_COUNTERMEASURE,
            title='Block IP',
        )
        seeded = task.status_changed_at

        task = TicketSubtask.objects.get(pk=task.pk)
        task.status = TicketSubtask.STATUS_IN_PROGRESS
        task.save(update_fields=['status'])
        task.refresh_from_db()
        self.assertGreater(task.status_changed_at, seeded)

    def test_subtask_status_rows_stay_out_of_the_ticket_change_table(self):
        # They have their own per-task history; repeating them here only burns
        # the 50-row cap and pushes real ticket edits out of view. Subtask NOTE
        # changes have no other home, so those stay.
        ticket = _make_ticket(created_by=self.t1, device_name='DILUTE-HOST')
        task = TicketSubtask.objects.create(
            ticket=ticket, subtask_type=TicketSubtask.TYPE_INVESTIGATION,
            title='Collect logs', assigned_to=self.t1,
        )
        history.record_subtask_status_change(
            task, TicketSubtask.STATUS_OPEN,
            TicketSubtask.STATUS_DONE, self.t1)
        history.record_subtask_change(task, '', 'findings attached', self.t1)
        before = history.snapshot(ticket)
        ticket.device_name = 'DILUTE-HOST-2'
        ticket.save()
        history.record_changes(ticket, before, self.t1, source='edit')

        self.client.force_login(self.t1)
        rows = self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).context['field_changes']

        pairs = {(row.field_name, row.subtask_id) for row in rows}
        self.assertIn(('device_name', None), pairs)
        self.assertIn(('result_notes', task.pk), pairs)
        self.assertNotIn(('status', task.pk), pairs)

    def test_history_is_shown_on_the_ticket_page(self):
        ticket = _make_ticket(created_by=self.t1, device_name='SHOWN-HOST')
        before = history.snapshot(ticket)
        ticket.device_name = 'CHANGED-HOST'
        ticket.save()
        history.record_changes(ticket, before, self.t1, source='edit')

        self.client.force_login(self.t1)
        html = self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode()
        self.assertIn('ประวัติการแก้ไขข้อมูล', html)
        self.assertIn('SHOWN-HOST', html)
        self.assertIn('CHANGED-HOST', html)


# ──────────────────────────────────────────────────────────────────────────── #
# 33. Ticket edit view (Phase 3)                                                #
# ──────────────────────────────────────────────────────────────────────────── #

class TicketEditViewTest(TestCase):
    """Before this there was no ticket edit route at all: Tier 1's content could
    only ever be corrected by Tier 2 during ESCALATED_T2."""

    @classmethod
    def setUpTestData(cls):
        cls.creator = _make_t1('edit_creator')
        cls.other_t1 = _make_t1('edit_other_t1')
        cls.t2 = _make_t2('edit_t2')
        cls.t2_other = _make_t2('edit_t2_other')
        cls.manager = _make_user('edit_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('edit_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _ticket(self, status=Ticket.STATUS_NEW):
        return _make_ticket(
            status=status, created_by=self.creator, assigned_admin=self.admin,
            classification=Ticket.CLASSIFICATION_INCIDENT,
            device_name='ORIGINAL-HOST', ip_address='192.0.2.1',
            issue_description='original text',
            log_source='Wazuh', issue_type='SIEM',
            detailed_issue='Investigating', detailed_issue2='Investigating Other',
            severity='High', ncsa_severity=Ticket.NCSA_SEVERITY_SEVERE,
            incident_datetime=timezone.now(),
        )

    def _payload(self, ticket, **overrides):
        data = {
            'classification': ticket.classification,
            'incident_name': ticket.incident_name or '',
            'severity': 'High',
            'ncsa_severity': Ticket.NCSA_SEVERITY_SEVERE,
            'incident_datetime': timezone.localtime().strftime('%Y-%m-%dT%H:%M'),
            'log_source': 'Wazuh', 'issue_type': 'SIEM',
            'detailed_issue': 'Investigating',
            'detailed_issue2': 'Investigating Other',
            'device_name': ticket.device_name,
            'issue_description': ticket.issue_description,
            'ip_address': ticket.ip_address,
            'reason': 'พิมพ์ผิด',
        }
        data.update(overrides)
        return data

    # ── Access ───────────────────────────────────────────────────────── #

    def test_creator_may_edit_while_the_ticket_is_untouched(self):
        ticket = self._ticket()
        self.client.force_login(self.creator)
        self.assertEqual(
            self.client.get(reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_creator_keeps_access_after_the_ticket_moves_because_they_are_soc(self):
        """"Creator while untouched, then SOC" — and Tier 1 *is* SOC, so the
        creator keeps access either way. The rule bites on non-SOC roles; the
        NEW-and-creator clause only matters for a Tier 1 who did not create it.
        """
        ticket = self._ticket(status=Ticket.STATUS_ESCALATED_T2)
        self.client.force_login(self.creator)
        self.assertEqual(
            self.client.get(reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_another_tier1_may_also_edit(self):
        """Not a loophole — SOC members are trusted with each other's tickets
        and every change is attributed. Asserted so the boundary is explicit."""
        ticket = self._ticket()
        self.client.force_login(self.other_t1)
        self.assertEqual(
            self.client.get(reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_soc_may_edit_after_the_ticket_has_moved(self):
        ticket = self._ticket(status=Ticket.STATUS_ESCALATED_T2)
        for user in (self.t2, self.manager):
            with self.subTest(user=user.username):
                self.client.force_login(user)
                self.assertEqual(
                    self.client.get(
                        reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_system_admin_may_not_edit(self):
        ticket = self._ticket(status=Ticket.STATUS_AWAITING_CONTAINMENT)
        self.client.force_login(self.admin)
        resp = self.client.get(reverse('edit_ticket', args=[ticket.pk]))
        self.assertRedirects(resp, reverse('ticket_detail', args=[ticket.pk]))

    # ── Tier 2 claim ─────────────────────────────────────────────────── #

    def _claimed_ticket(self, claimer):
        ticket = self._ticket(status=Ticket.STATUS_ESCALATED_T2)
        ticket.t2_claimed_by = claimer
        ticket.t2_claimed_at = timezone.now()
        ticket.save(update_fields=['t2_claimed_by', 't2_claimed_at'])
        return ticket

    def test_a_claim_blocks_a_second_tier2_from_editing(self):
        """The claim covers content, not just the status move. transition_to
        already refuses this analyst; the edit surface must agree, or they can
        rewrite what the claimer is reading."""
        ticket = self._claimed_ticket(self.t2)
        self.client.force_login(self.t2_other)
        resp = self.client.get(reverse('edit_ticket', args=[ticket.pk]))
        self.assertRedirects(resp, reverse('ticket_detail', args=[ticket.pk]))

    def test_a_claim_does_not_block_the_claimer(self):
        ticket = self._claimed_ticket(self.t2)
        self.client.force_login(self.t2)
        self.assertEqual(
            self.client.get(reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_a_claim_does_not_block_other_roles(self):
        """Only Tier 2 queues on the Tier 2 claim — a manager or the opening
        analyst still corrects content while a Tier 2 holds it."""
        ticket = self._claimed_ticket(self.t2)
        for user in (self.manager, self.creator):
            with self.subTest(user=user.username):
                self.client.force_login(user)
                self.assertEqual(
                    self.client.get(
                        reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    def test_an_unclaimed_ticket_stays_open_to_any_tier2(self):
        ticket = self._ticket(status=Ticket.STATUS_ESCALATED_T2)
        self.client.force_login(self.t2_other)
        self.assertEqual(
            self.client.get(reverse('edit_ticket', args=[ticket.pk])).status_code, 200)

    # ── Out-of-court warning ─────────────────────────────────────────── #

    def test_no_warning_when_the_editor_holds_the_court(self):
        ticket = self._ticket()  # NEW → the creator's court
        self.client.force_login(self.creator)
        resp = self.client.get(reverse('edit_ticket', args=[ticket.pk]))
        self.assertFalse(resp.context['out_of_court'])

    def test_warning_names_the_holder_when_out_of_court(self):
        """Editing a ticket parked with someone else is allowed but flagged."""
        ticket = self._ticket(status=Ticket.STATUS_AWAITING_CONTAINMENT)
        self.client.force_login(self.creator)
        resp = self.client.get(reverse('edit_ticket', args=[ticket.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.context['out_of_court'])
        self.assertIn('System Admin', resp.context['court_holder'])
        self.assertIn(self.admin.username, resp.context['court_holder'])
        # And it actually reaches the page — a context flag the template never
        # reads would warn nobody.
        body = resp.content.decode()
        self.assertIn('alert-warning', body)
        self.assertIn(self.admin.username, body)

    def test_court_holder_is_none_for_a_terminal_ticket(self):
        ticket = self._ticket(status=Ticket.STATUS_APPROVED)
        self.assertIsNone(ticket.court_holder_label)

    def test_closed_tickets_are_frozen(self):
        for status in (Ticket.STATUS_APPROVED, Ticket.STATUS_CLOSED_EVENT):
            ticket = self._ticket(status=status)
            for user in (self.creator, self.manager):
                with self.subTest(status=status, user=user.username):
                    self.client.force_login(user)
                    resp = self.client.post(
                        reverse('edit_ticket', args=[ticket.pk]),
                        self._payload(ticket, device_name='HACKED'))
                    self.assertRedirects(
                        resp, reverse('ticket_detail', args=[ticket.pk]))
                    ticket.refresh_from_db()
                    self.assertEqual(ticket.device_name, 'ORIGINAL-HOST')

    # ── Behaviour ────────────────────────────────────────────────────── #

    def test_edit_saves_and_records_the_change(self):
        ticket = self._ticket()
        self.client.force_login(self.creator)
        self.client.post(
            reverse('edit_ticket', args=[ticket.pk]),
            self._payload(ticket, device_name='FIXED-HOST'))

        ticket.refresh_from_db()
        self.assertEqual(ticket.device_name, 'FIXED-HOST')
        change = ticket.field_changes.get(field_name='device_name')
        self.assertEqual(change.old_value, 'ORIGINAL-HOST')
        self.assertEqual(change.source, 'edit')

    def test_edit_requires_a_reason(self):
        ticket = self._ticket()
        self.client.force_login(self.creator)
        self.client.post(
            reverse('edit_ticket', args=[ticket.pk]),
            self._payload(ticket, device_name='NO-REASON', reason=''))
        ticket.refresh_from_db()
        self.assertEqual(ticket.device_name, 'ORIGINAL-HOST')

    def test_edit_writes_a_timeline_entry_naming_the_fields(self):
        ticket = self._ticket()
        self.client.force_login(self.creator)
        self.client.post(
            reverse('edit_ticket', args=[ticket.pk]),
            self._payload(ticket, device_name='LOGGED-HOST', reason='ระบุชื่อเครื่องผิด'))
        log = TicketLog.objects.filter(ticket=ticket).latest('created_at')
        self.assertIn('แก้ไขข้อมูลเคส', log.note)
        self.assertIn('ระบุชื่อเครื่องผิด', log.note)

    def test_edit_cannot_move_the_ticket_through_the_workflow(self):
        """It is a correction surface, not a workflow one."""
        ticket = self._ticket()
        self.client.force_login(self.creator)
        self.client.post(
            reverse('edit_ticket', args=[ticket.pk]),
            self._payload(ticket, status=Ticket.STATUS_APPROVED,
                          t1_route=Ticket.T1_ROUTE_OWNER))
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.STATUS_NEW)

    def test_edit_button_visibility_follows_permission(self):
        ticket = self._ticket()
        edit_url = reverse('edit_ticket', args=[ticket.pk])

        self.client.force_login(self.creator)
        self.assertIn(edit_url, self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode())

        self.client.force_login(self.admin)
        self.assertNotIn(edit_url, self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode())


# ──────────────────────────────────────────────────────────────────────────── #
# 34. Manager step-back (Phase 4)                                               #
# ──────────────────────────────────────────────────────────────────────────── #

class ManagerStepBackTest(TestCase):
    """The manager's own forward at PENDING_MGR_TRIAGE had no way back, and
    approval was approve-or-nothing. Closure stays terminal."""

    @classmethod
    def setUpTestData(cls):
        cls.t1 = _make_t1('back_t1')
        cls.t2 = _make_t2('back_t2')
        cls.manager = _make_user('back_mgr', UserProfile.ROLE_SOC_MANAGER)
        cls.admin = _make_user('back_admin', UserProfile.ROLE_SYSTEM_ADMIN)

    def _ticket(self, status, route=Ticket.T1_ROUTE_ADMIN):
        return _make_ticket(
            status=status, created_by=self.t1, assigned_admin=self.admin,
            classification=Ticket.CLASSIFICATION_INCIDENT, t1_route=route,
        )

    def _step_back(self, user, ticket, reason='มอบหมายผิดผู้ดูแลระบบ'):
        self.client.force_login(user)
        return self.client.post(reverse('ticket_detail', args=[ticket.pk]), {
            'action': 'step_back', 'step_back_reason': reason,
        })

    def test_manager_can_pull_a_forwarded_ticket_back_to_triage(self):
        for status in (Ticket.STATUS_AWAITING_CONTAINMENT, Ticket.STATUS_AWAITING_OWNER):
            with self.subTest(status=status):
                ticket = self._ticket(status)
                self._step_back(self.manager, ticket)
                ticket.refresh_from_db()
                self.assertEqual(ticket.status, Ticket.STATUS_PENDING_MGR_TRIAGE)

    def test_pending_manager_steps_back_along_the_lane_it_came_from(self):
        admin_lane = self._ticket(Ticket.STATUS_PENDING_MANAGER,
                                  route=Ticket.T1_ROUTE_ADMIN)
        self._step_back(self.manager, admin_lane)
        admin_lane.refresh_from_db()
        self.assertEqual(admin_lane.status, Ticket.STATUS_CONTAINMENT_REPORTED)

        owner_lane = self._ticket(Ticket.STATUS_PENDING_MANAGER,
                                  route=Ticket.T1_ROUTE_OWNER)
        self._step_back(self.manager, owner_lane)
        owner_lane.refresh_from_db()
        self.assertEqual(owner_lane.status, Ticket.STATUS_PENDING_T2_REVIEW)

    def test_closed_tickets_stay_closed(self):
        """The guarantee the whole feature is designed around."""
        for status in (Ticket.STATUS_APPROVED, Ticket.STATUS_CLOSED_EVENT):
            with self.subTest(status=status):
                ticket = self._ticket(status)
                self._step_back(self.manager, ticket)
                ticket.refresh_from_db()
                self.assertEqual(ticket.status, status)

    def test_only_a_soc_manager_may_step_back(self):
        for user in (self.t1, self.t2, self.admin):
            with self.subTest(user=user.username):
                ticket = self._ticket(Ticket.STATUS_AWAITING_CONTAINMENT)
                self._step_back(user, ticket)
                ticket.refresh_from_db()
                self.assertEqual(ticket.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    def test_a_reason_is_required(self):
        ticket = self._ticket(Ticket.STATUS_AWAITING_CONTAINMENT)
        self._step_back(self.manager, ticket, reason='   ')
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.STATUS_AWAITING_CONTAINMENT)

    def test_step_back_is_logged_with_both_statuses_and_the_reason(self):
        ticket = self._ticket(Ticket.STATUS_AWAITING_CONTAINMENT)
        self._step_back(self.manager, ticket, reason='ต้องประเมินความรุนแรงใหม่')
        log = TicketLog.objects.filter(ticket=ticket).latest('created_at')
        self.assertIn('ย้อนขั้นตอน', log.note)
        self.assertIn('ต้องประเมินความรุนแรงใหม่', log.note)
        self.assertEqual(log.author, self.manager)

    def test_step_back_does_not_rewrite_who_decided_what(self):
        """Sign-off and emergency stamps name the first decider forever."""
        ticket = self._ticket(Ticket.STATUS_PENDING_MANAGER)
        ticket.verified_by = self.t2
        ticket.verified_at = timezone.now()
        ticket.emergency_decided_by = self.manager
        ticket.emergency_decided_at = timezone.now()
        ticket.save()

        self._step_back(self.manager, ticket)
        ticket.refresh_from_db()
        self.assertEqual(ticket.verified_by, self.t2)
        self.assertEqual(ticket.emergency_decided_by, self.manager)

    def test_statuses_without_a_defined_target_are_refused(self):
        ticket = self._ticket(Ticket.STATUS_NEW)
        self.assertIsNone(ticket.step_back_target())
        self.assertFalse(ticket.can_step_back(self.manager))
        self._step_back(self.manager, ticket)
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.STATUS_NEW)

    def test_project_members_must_use_the_project_page(self):
        project = ProjectIncident.objects.create(
            title='bundle', created_by=self.t1)
        ticket = self._ticket(Ticket.STATUS_AWAITING_CONTAINMENT)
        ticket.project_incident = project
        ticket.save(update_fields=['project_incident'])
        self.assertFalse(ticket.can_step_back(self.manager))

    def test_control_is_only_offered_to_managers(self):
        ticket = self._ticket(Ticket.STATUS_AWAITING_CONTAINMENT)
        self.client.force_login(self.manager)
        self.assertIn('ย้อนขั้นตอน', self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode())

        self.client.force_login(self.t2)
        self.assertNotIn('ย้อนขั้นตอน', self.client.get(
            reverse('ticket_detail', args=[ticket.pk])).content.decode())
